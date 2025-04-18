import asyncio
import base64
import datetime
import json
import os
import re
import sys
import threading
import time

from io import BytesIO
from qasync import QEventLoop, asyncSlot

def load_settings():
    """Load application settings from JSON file or use defaults."""
    try:
        if os.path.exists('settings.json'):
            with open('settings.json', 'r') as f:
                saved_settings = json.load(f)
                # Merge saved settings with default settings to handle new keys
                merged_settings = DEFAULT_SETTINGS.copy()
                merged_settings.update(saved_settings)
                return merged_settings
    except Exception as e:
        print(f"Error loading settings: {e}")
    return DEFAULT_SETTINGS.copy()

# Load settings before using them
settings = load_settings()

from PyQt5.QtWidgets import (
    QApplication, QWidget, QPushButton, QVBoxLayout, QLabel, QTextEdit, QLineEdit,
    QFileDialog, QMessageBox, QHBoxLayout, QSystemTrayIcon, QMenu, QAction, QComboBox,
    QSplitter, QTabWidget, QDialog, QFormLayout, QSpinBox, QCheckBox, QRadioButton,
    QButtonGroup, QToolButton, QSizePolicy, QShortcut, QProgressBar, QScrollArea,
    QFrame, QGroupBox, QGridLayout, QColorDialog, QFontDialog
)
from PyQt5.QtGui import (
    QPalette, QColor, QFont, QIcon, QPixmap, QKeySequence, QTextCharFormat, 
    QSyntaxHighlighter, QTextCursor, QFontMetrics, QPainter, QTextFormat, 
    QTextBlockFormat, QBrush, QTextDocument
)
from PyQt5.QtCore import (
    Qt, QTimer, pyqtSignal, QObject, QSize, QRect, QUrl, QThread, 
    QRegExp, QEvent, QPoint, QMargins
)

from playwright.async_api import async_playwright
import groq
from docx import Document
from fpdf import FPDF
from wordcloud import WordCloud
from PIL import Image
import keyboard
from markdown import markdown
import webbrowser

# --------------------------
# Constants and Settings
# --------------------------
VERSION = "2.0.0"
APP_NAME = "MeetingAssistant+"
DEFAULT_FONT_SIZE = 14
DEFAULT_FONT_FAMILY = "SF Pro Text"
DEFAULT_LINE_SPACING = 150  # percentage
DEFAULT_WORD_SPACING = 1.2  # multiplier
DEFAULT_POLL_INTERVAL = 5   # seconds
DEFAULT_CODE_WORD = "blossom"

# --------------------------
# Default model configuration
# --------------------------
DEFAULT_MODEL = "llama3-70b-8192"
AVAILABLE_MODELS = [
    "llama3-70b-8192",
    "llama3-8b-8192",
    "mixtral-8x7b-32768",
    "gemma-7b-it"
]

# --------------------------
# Theme definitions
# --------------------------
THEMES = {
    "dark": {
        "window": "#292A2D",
        "windowText": "#FFFFFF",
        "base": "#292A2D",
        "alternateBase": "#3C3E44",
        "text": "#FFFFFF",
        "button": "#3C3E44",
        "buttonText": "#FFFFFF",
        "highlight": "#4DB6AC",
        "highlightedText": "#000000",
        "link": "#88C0D0",
        "brightText": "#FF5555",
        "borderColor": "#4D4D4D",
        "gradientStart": "#292A2D",
        "gradientEnd": "#1E1F21"
    },
    "light": {
        "window": "#F5F6F7",
        "windowText": "#333333",
        "base": "#FFFFFF",
        "alternateBase": "#F0F0F0",
        "text": "#333333",
        "button": "#E4E4E4",
        "buttonText": "#333333",
        "highlight": "#4DB6AC",
        "highlightedText": "#FFFFFF",
        "link": "#0077CC",
        "brightText": "#FF5555",
        "borderColor": "#DDDDDD",
        "gradientStart": "#F5F6F7",
        "gradientEnd": "#E6E7E8"
    },
    "navy": {
        "window": "#0D101E",
        "windowText": "#E6F1FF",
        "base": "#0D101E",
        "alternateBase": "#172A45",
        "text": "#E6F1FF", 
        "button": "#172A45",
        "buttonText": "#E6F1FF",
        "highlight": "#64FFDA",
        "highlightedText": "#0A192F",
        "link": "#64FFDA",
        "brightText": "#FF5555",
        "borderColor": "#253656",
        "gradientStart": "#0D101E",
        "gradientEnd": "#0A0D16"
    },
    "dracula": {
        "window": "#282A36",
        "windowText": "#F8F8F2",
        "base": "#282A36",
        "alternateBase": "#44475A",
        "text": "#F8F8F2",
        "button": "#44475A",
        "buttonText": "#F8F8F2",
        "highlight": "#FF79C6",
        "highlightedText": "#282A36",
        "link": "#8BE9FD",
        "brightText": "#FF5555",
        "borderColor": "#383A48",
        "gradientStart": "#282A36",
        "gradientEnd": "#21222C"
    },
    "macOS": {
        "window": "#EBEBEB",
        "windowText": "#333333",
        "base": "#FFFFFF",
        "alternateBase": "#F7F7F7",
        "text": "#333333",
        "button": "#F2F2F2",
        "buttonText": "#333333",
        "highlight": "#007AFF",
        "highlightedText": "#FFFFFF",
        "link": "#007AFF",
        "brightText": "#FF3B30",
        "borderColor": "#DDDDDD",
        "gradientStart": "#F7F7F7",
        "gradientEnd": "#EBEBEB"
    },
    "macOSDark": {
        "window": "#323232",
        "windowText": "#FFFFFF",
        "base": "#2A2A2A",
        "alternateBase": "#3A3A3A",
        "text": "#FFFFFF",
        "button": "#4A4A4A",
        "buttonText": "#FFFFFF",
        "highlight": "#0A84FF",
        "highlightedText": "#FFFFFF",
        "link": "#0A84FF",
        "brightText": "#FF453A",
        "borderColor": "#494949",
        "gradientStart": "#323232",
        "gradientEnd": "#2A2A2A"
    }
}
# --------------------------
# Default Settings
# --------------------------
DEFAULT_SETTINGS = {
    "active_prompt": "Job Interview Assistant",
    "prompts": [{
        "name": "Job Interview Assistant",
        "prompt": """You are an intelligent assistant attending a professional job interview. Analyze the conversation in real time and help the user prepare to answer if someone poses a question to them.

Instructions:
- Detect if another speaker asks a question directed at the user (e.g., using "you", "your experience", or direct names).
- When a question is asked, generate a contextual, helpful response or insight the user could consider sharing.
- Constantly monitor the discussion to summarize the current topic and key takeaways.
- Present this information clearly and concisely.
- If appropriate, include action points or things the user should pay attention to.

Output format:
SUGGESTED RESPONSE: [Your suggested response if a question was asked]

KEY INSIGHTS: [Important insights about the conversation]

TOPIC SUMMARY: [Brief summary of the current topic]
"""
    }],
    "summary_line_interval": 10,
    "summary_time_interval": 120,
    "poll_interval": DEFAULT_POLL_INTERVAL,
    "theme": "macOS",  # Default to macOS theme
    "start_in_background": False,
    "font_size": DEFAULT_FONT_SIZE,
    "font_family": "SF Pro Text",  # Mac system font
    "line_spacing": DEFAULT_LINE_SPACING,
    "word_spacing": DEFAULT_WORD_SPACING,
    "model": DEFAULT_MODEL,
    "keyboard_shortcut": "ctrl+i",
    "code_word": DEFAULT_CODE_WORD,
    "app_icon_path": "",
    "generate_icon_path": "",
    "settings_icon_path": "",
    "anonymize_transcript": False,
    "spellcheck_enabled": False,
    "meeting_timer_enabled": True,
    "meeting_reminder_interval": 15,  # minutes
    "use_rounded_corners": True,  # Enable rounded corners for Mac-like appearance
    "use_animations": True,  # Enable animations for modern feel
    "use_transparency": True,  # Enable transparency effects
}

# --------------------------
# Embedded Icons (Base64)
# --------------------------
DEFAULT_APP_ICON = "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAACXBIWXMAAAsTAAALEwEAmpwYAAAF7ElEQVR4nO2bW2xURRjHf93dttBSKHcRqKWCBYJGQSAqIheuJhAUEzQqEBJvD5oYE33AB0x8UI2aGBOFqBiJGkWJd6IYRdAIhHspFGhpuRQo0ALttrB7fPhmOWfbc2b2zDlnd9v9JZPuzsx833z/uXxnZs+BEkooIXfiLuAssKzYisQ0eUCXeZxX0L7HAH3ALGDmIPQTKxwEHgTuNs9HB6HPVJBrBMwDngOmF0gXSYeRwMPAEuP0AqA3i+I6L2yqkRHgllFXcQcQ5sQ4RoCqAtCNDXxGgBoZSdIxAvpyJYjC7H8d2DpA5UKwB7jRJa8qPK+64pXc1AEMNpTbTbXAKeAPYHSK+DGIRGKwv0B8cBcQxEpgB/AzsNwlbw7wK/ArMNkl/1HEgLeBWYNhwO8h+VsTyNsd/rNIIKXCXqSdVAc5fAuwE7ijAPqNAK4DXgXWACnlRFUMFRqBy4BLgU7gR2BzwD1vApYCdcBxYD0wyuRNBC4CJgDHgF+Adqtsa8QE+wwVYiZsb91dttLdiSSCXcAmcz3NlB8FWgJ12AEsNG30ADPDGDCB4CJoRrA9zJ8EbqXzaYE8gDrgHiTiROmpAxYjgxi52NEdks9hynzgCZO3FVgBXBR4T52DdcAjwHeIDw6xCIpgW3kVw3jgNaAZ2Ae8iIwEhc1Ie00HXou0tREY53CfPUo2Bziq5MuMbKzJdL8iE2CHvRfQRrJttIQa1GZGxgzjfDuwzl4EVVlKPoKsx78Bs4EPzPU/yKTGxjZkONt4C6hR8vcAU03+bF8D3LAN+BE4gXS8FnkR37jkzUXmlOUW2WrEsAsRJ04j79fLkPoUwOeAlcA1yIgKrQPijGZgF/CJtg44iqxvNYH3dADrzPVnJv0Vn8rLwhhKZuIU7p6HA9WIfaFQw0dV9qnwfJtqrZvzBXLgNcAFyiQnaFBylRwxUkbCO3EYcKFyPxJJa4P0eD1izx7EvjlINHnXp/IBSAy02GXuDYFaXuOULLhz9ci8P0MHMnQPmvRF4G/kKA0y1aZZZQ4gS/RQZHRsBFab6y5gbcRoaVJlLlXypQm07HqoICPqZauOUWGjKWO0KfdpDqrrjkbmKltVGRs3K3mbkqm2txNQptx3tpKvV+6rkqXxdK4L+FDJg7EO+T5/ZJFrU7BXydrjf7vM47PIn9dAp/lgNDrhwGLvN8sYoLuuziFvV2U06O2wRnUdavQGj2sXNCjZYSWfoGRdNicxfpJVRof4euBdjyfhYk8DQuqALuC3WjnqAqkD7lP1eBswT8k+C2ijzWybSdW/y+i9JoDWGXURpFH0BuU+pA7Q+CJL3iLgTav8UPP3D0Lr6YpYOTVqVfucwA7rjWy3ufZZgnW7TQ5o40Ml6/CpfCiUGHkn/sRzFYiiNTJN6hFJBU5G7k6glUBHR+tHDjw97PxOcaH1VCCnVLuR1fJu4GXgc+T0OMtxGj1hehDYH5Kvc8E4TuGI+dM3X+tN5bOtkMnS0Hm6XsJG+pR9sZKpQA6+gnE56mwqoEbIcVgceNr0eQyZcJXlCDeQI/ORSu6dCzqATnnbPLcCwbOzdmBn6KGJjlw9/5Rl5dwE/B/wkWWmX7dXdvAQ+x1/eGAv0NfMXpXWxYAHgctDGk0gXlFldJi8zaSXIUO0FznYiaxGNLKpugVYinw7OI0cRn6G7HbrkGW7ATl9qkO+kB83tE5ZX1UgI2wRsmG5ErkvcnARDVvBwTLgIJJdjXxAvAmZS/o+UO/IOjtYjtQlv5jn95FtOhHaMk1g/GG7c3lD3Uw9ZBOzCNk5xlrNvJFnPbKiQBFRr8r2RXTKAKnz0/NpQJo0T3ehj9TqMnz7jjJCnx8D0iNDIXM2GeALnWs4R0bAfPIwZ58jo+chz4+LiWGpIzww3DxODs+nF+8YGaXI8hzhqELOyXVedLRGxidA6/9EvIQBxL/n+ibGRp+7nQAAAABJRU5ErkJggg=="
DEFAULT_SETTINGS_ICON = "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAACXBIWXMAAAsTAAALEwEAmpwYAAACnUlEQVR4nO2bS2sUQRDHf5OsJkGNEo0XUTxFiaCCeFE8KH4BRT2IB8WzFw9+A0G/gF/AgyjiK4ioRHzFq5hE8yDmgVHjBpNoMrO7HgZmDZr0TPfszux0wT9gDzM1Vf3vjZ6q7h0oKSkpyTe7gSHgFjAN/AG+A++Bh8B5oBfolPY7jee79DMs99j3TAMPgHNAt/RbCLqAu8An0ZQmPAeGgIPAtgb9twOH5Zlxea5RvwmgXyYrNxwB3ko0FWEZuAn0NTHGXuCWPKvJ/wg0aU0uRsElYMlDWf8DvAJGgUvAZWAM+OThPekzJWNsijMiC6eROCLCJjwEl4DRmISPgV6HMfeK0JrId8DF5GRuzHTF9Cv0D+AEsNHBz03ASeB7RPYrcCSR0C42AS8VuVlgh7YfEbmasGiw7Bb/d+Bn5d8+/dWciqBcRf5Jn4F9KoLP0wS+KPIDGvwBReYl2kxJL/apCT6QKWA4TeBbRX5Z6TvVBE9KuC+0t1qlvUuR+Qps1RBbh98x/o3OFN5rTOaFBr9XkaFegK+Q6XrLHyEHmBNOiW81yfkc+8OZ6prkUG0C9uy9VzszHxS5L9pcPUw+tZTp+1/v4bTlBVQzKgAAAABJRU5ErkJggg=="
DEFAULT_GENERATE_ICON = "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAACXBIWXMAAAsTAAALEwEAmpwYAAAD5klEQVR4nO2aXWjOYRjGr4+Z2Ij5KJGU+WimJhpSalIzopkPMR8H5IDIgTjAyYQDB1JKFGuKFCkHREopB8qBpChfOdqBKEVTTDN/13q73q3t/f6e93mfj3d7r+rX+3zc93U/7/s893PfzwvEiBFjDGAmgCYAhNkAGwCOARgXsm2jADQAOALAFbLzAI4DGBuBfLl4fABwCMDoUAwA8ES2nT5APYCxA2jbEAD1AJ6atr0DUDXQzlcYdATAkAFmwBAAhwF8MvJ9AbATgJ3rDbCdJQDqADzP0tl+AM8AbMl1wBHYRpnudFnwAVSGRn4vTgX0BMAyAG7I9jkA7gJ4F6ANnYYOb+dTQZDn6XxU+3o3gAJD9gxj+Swp9GQeMAR4GIA8pdz1nv93BYBbAOZ7rp+V4SuiREILtS3KApKnLPZuAMOMHWyXz3BFrXdE+f+2CWBNFPLFAG4DOKdsm+lhRyWAz552fQJQ5lPhZgVxH4AxWqUAjAZwV9kmrzfcxzSjjbX6OmvhWUYfHwMoA1Cr5LpGJm+RUu4Q45nfA0ARWUQBLPV0MgOgTcntM9ZYpZTnGoYEtM0qyoxlLAPghJbUcZzrAE4aJe/vGsf5Oy+OiYcAtuFl5c0TaHh6+VWYRQLPyPr8J54xq1DaSUDDm4PtMnZJW7lk+sKvAAUyZTCAxwKd7A7pA+qU62Sx/2Cwy+cw7QYsYLXQ5LMSVGbJbwlb0Hag3pE9hEXsEJr4AEm+S/aeM4JWJf1gGVOFJr0oQNJFSlkZtFOg2U+cHO0Ep0gBWiGFBHQA+CWEKlHihlDWVzFJSOGjEkzuCGWlEolENKEX04UUXipB3RgllZTwqsVRQsIrIYWpSrAEySW3lOOjAuWEWCkxS0jhsxIsFCStVsrfU/q5YoDQQgnmCCn0KMG8J0qbD0LCXKGEV4rQR0G+LEJ+XlGeGSZ0FHK1W5Iy5Z/KfrdHwRkGUQpgohJsUlaUvqQM6+QipGzTd4CuIMQGsaFDSFmvpLSRFraaYdCCzJNkIfZHwQkhXRHCdBLQqTBiI4ASIX2rAC/tWwbgp4BOhkmsE9L12UyIpJt3RSwAUCugl8pzx4bLNhH6C0AWCGol9LID6mfYL0S4UeDeXl4a3PGcPgXFX8uCGzY3GzjOv78D+J0JTsgdVSFRDuCcZwNTkWtQFBmZnZ904eRMbpwRdtjA7RFWNtnQWQmPL0pA9lVXKCezoSDbmGPz1X1hFWRRkMOO3aSHsRLtDZJlTHM7e4OkWWqSPtvQGYwRI0aMf/wGCaZOUO2BGa0AAAAASUVORK5CYII="

# --------------------------
# Helper Functions
# --------------------------

def load_icon_from_path_or_default(path, default_base64):
    """Load an icon from a file path or fall back to the default base64 icon."""
    if path and os.path.exists(path):
        try:
            return QIcon(path)
        except Exception as e:
            print(f"Error loading icon from {path}: {e}")
            
    # Fall back to default
    pix = QPixmap()
    pix.loadFromData(base64.b64decode(default_base64))
    icon = QIcon()
    icon.addPixmap(pix)
    return icon

def create_styled_button(text, icon=None, tooltip=None):
    """Create a styled button with optional icon and tooltip."""
    btn = QPushButton(text)
    if icon:
        btn.setIcon(icon)
    if tooltip:
        btn.setToolTip(tooltip)
    btn.setMinimumHeight(36)
    btn.setStyleSheet("""
        QPushButton {
            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #5E5E5E, stop:1 #444444);
            border: 1px solid #555;
            border-radius: 6px;
            padding: 5px 15px;
            color: white;
            font-weight: 500;
        }
        QPushButton:hover {
            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #656565, stop:1 #4F4F4F);
        }
        QPushButton:pressed {
            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #444444, stop:1 #5E5E5E);
            padding-top: 6px;
            padding-bottom: 4px;
        }
        QPushButton:disabled {
            background-color: #3A3A3A;
            color: #7F7F7F;
        }
    """)
    return btn

def create_icon_button(icon, tooltip=None, size=32):
    """Create a button with just an icon."""
    btn = QToolButton()
    btn.setIcon(icon)
    if tooltip:
        btn.setToolTip(tooltip)
    btn.setIconSize(QSize(size, size))
    btn.setMinimumSize(QSize(size + 12, size + 12))
    btn.setStyleSheet("""
        QToolButton {
            background-color: transparent;
            border: 1px solid #555;
            border-radius: 6px;
            padding: 3px;
        }
        QToolButton:hover {
            background-color: rgba(255, 255, 255, 0.1);
        }
        QToolButton:pressed {
            background-color: rgba(0, 0, 0, 0.1);
            padding-top: 4px;
            padding-bottom: 2px;
        }
    """)
    return btn
# --------------------------
# Word Cloud Generator
# --------------------------

def generate_wordcloud(text):
    wc = WordCloud(width=800, height=400, background_color='black', colormap='Blues').generate(text)
    buffer = BytesIO()
    wc.to_image().save(buffer, format="PNG")
    return buffer.getvalue()

# --------------------------
# Name Anonymizer
# --------------------------

def anonymize_names(text, known_names=None):
    """Replace detected names with anonymous identifiers."""
    if not known_names:
        known_names = []
        
    # Simple pattern to detect potential names (starting with capital letter)
    name_pattern = r'\b[A-Z][a-z]+\b'
    
    # Find all potential names
    potential_names = set(re.findall(name_pattern, text))
    
    # Add known names
    all_names = potential_names.union(set(known_names))
    
    # Create a mapping of names to anonymous identifiers
    name_mapping = {}
    for i, name in enumerate(all_names):
        name_mapping[name] = f"Person-{i+1}"
        
    # Replace names in text
    anonymized_text = text
    for name, anon_id in name_mapping.items():
        anonymized_text = re.sub(r'\b' + name + r'\b', anon_id, anonymized_text)
        
    return anonymized_text, name_mapping

# --------------------------
# Simple Spell Checker
# --------------------------

def correct_spelling(text):
    """Simple spell checking and correction."""
    # Common misspellings and corrections
    corrections = {
        "teh": "the",
        "taht": "that",
        "thier": "their",
        "waht": "what",
        "dont": "don't",
        "cant": "can't",
        "wont": "won't",
        "im": "I'm",
        "youre": "you're",
        "theyre": "they're",
        "thats": "that's",
        "isnt": "isn't",
        "arent": "aren't",
        "hasnt": "hasn't",
        "havent": "haven't",
        "doesnt": "doesn't",
        "didnt": "didn't",
        "wasnt": "wasn't",
        "werent": "weren't",
        "wouldnt": "wouldn't",
        "shouldnt": "shouldn't",
        "couldnt": "couldn't",
        "alot": "a lot",
        "alright": "all right",
        "aswell": "as well",
        "becuase": "because",
        "definately": "definitely",
        "wierd": "weird",
        "recieve": "receive",
        "seperate": "separate",
        "untill": "until",
        "tommorow": "tomorrow",
        "tommorrow": "tomorrow",
        "alway": "always",
        "ofcourse": "of course",
        "ofcource": "of course",
        "allready": "already",
        "allright": "all right",
        "accross": "across",
        "beleive": "believe",
        "comunicate": "communicate",
        "greatful": "grateful",
        "dependance": "dependence",
        "excelent": "excellent",
        "existance": "existence",
        "foreward": "forward",
        "grammer": "grammar",
        "gaurd": "guard",
        "hight": "height",
        "humerous": "humorous",
        "independant": "independent",
        "liason": "liaison",
        # --------------------------
def generate_wordcloud(text):
    wc = WordCloud(width=800, height=400, background_color='black', colormap='Blues')
    wc.generate(text)  # Generate the word cloud from the text
    buffer = BytesIO()  # Create a buffer to store the image
    wc.to_image().save(buffer, format="PNG")  # Save the image to the buffer
    return buffer.getvalue()  # Return the buffer's contents
        # --------------------------
        # Name Anonymizer
        # --------------------------
        
        def anonymize_names(text, known_names=None):
            """Replace detected names with anonymous identifiers."""
            if not known_names:
                known_names = []
        
            # Simple pattern to detect potential names (starting with capital letter)
            name_pattern = r'\b[A-Z][a-z]+\b'
        
            # Find all potential names
            potential_names = set(re.findall(name_pattern, text))
        
            # Add known names
            all_names = potential_names.union(set(known_names))
        
            # Create a mapping of names to anonymous identifiers
            name_mapping = {}
            for i, name in enumerate(all_names):
                name_mapping[name] = f"Person-{i+1}"
        
            # Replace names in text
            anonymized_text = text
            for name, anon_id in name_mapping.items():
                anonymized_text = re.sub(r'\b' + name + r'\b', anon_id, anonymized_text)
        
            return anonymized_text, name_mapping
        
        # --------------------------
        # Simple Spell Checker
        # --------------------------
        
        def correct_spelling(text):
            """Simple spell checking and correction."""
            # Common misspellings and corrections
            corrections = {
                "teh": "the",
                "taht": "that",
                "thier": "their",
                "waht": "what",
                "dont": "don't",
                "cant": "can't",
                "wont": "won't",
                "im": "I'm",
                "youre": "you're",
                "theyre": "they're",
                "thats": "that's",
                "isnt": "isn't",
                "arent": "aren't",
                "hasnt": "hasn't",
                "havent": "haven't",
                "doesnt": "doesn't",
                "didnt": "didn't",
                "wasnt": "wasn't",
                "werent": "weren't",
                "wouldnt": "wouldn't",
                "shouldnt": "shouldn't",
                "couldnt": "couldn't",
                "alot": "a lot",
                "alright": "all right",
                "aswell": "as well",
                "becuase": "because",
                "definately": "definitely",
                "wierd": "weird",
                "recieve": "receive",
                "seperate": "separate",
                "untill": "until",
                "tommorow": "tomorrow",
                "tommorrow": "tomorrow",
                "alway": "always",
                "ofcourse": "of course",
                "ofcource": "of course",
                "allready": "already",
                "allright": "all right",
                "accross": "across",
                "beleive": "believe",
                "comunicate": "communicate",
                "greatful": "grateful",
                "dependance": "dependence",
                "excelent": "excellent",
                "existance": "existence",
                "foreward": "forward",
                "grammer": "grammar",
                "gaurd": "guard",
                "hight": "height",
                "humerous": "humorous",
                "independant": "independent",
                "liason": "liaison",
                "millenium": "millennium",
                "momento": "memento",
                "neccessary": "necessary",
                "occured": "occurred",
                "occurance": "occurrence",
                "peice": "piece",
                "preform": "perform",
                "prefered": "preferred",
                "probly": "probably",
                "reccommend": "recommend",
                "refered": "referred",
                "rember": "remember",
                "remeber": "remember",
                "sieze": "seize",
                "sucess": "success",
                "suprise": "surprise",
                "tomatos": "tomatoes",
                "vaccum": "vacuum",
                "writeing": "writing"
            }
        
            # Find words in text and replace them
            corrected_text = text
            words = re.findall(r'\b\w+\b', text.lower())
            for word in words:
                if word in corrections:
                    # Replace word with correction, preserving case
                    pattern = re.compile(r'\b' + word + r'\b', re.IGNORECASE)
                    match = pattern.search(corrected_text)
                    while match:
                        original = match.group(0)
                        correction = corrections[word.lower()]
        
                        # Preserve case
                        if original.islower():
                            replacement = correction.lower()
                        elif original.isupper():
                            replacement = correction.upper()
                        elif original[0].isupper():
                            replacement = correction.capitalize()
                        else:
                            replacement = correction
        
                        corrected_text = corrected_text[:match.start()] + replacement + corrected_text[match.end():]
                        match = pattern.search(corrected_text, match.start() + len(replacement))
        
            return corrected_text
        
        # --------------------------
        # Signal Bridge for Thread-Safe UI Updates
        # --------------------------
        class SignalBridge(QObject):
            append_transcript = pyqtSignal(str)
            append_suggested_response = pyqtSignal(str)
            append_insights = pyqtSignal(str)
            update_status = pyqtSignal(str)
            update_progress = pyqtSignal(int)
            highlight_transcript = pyqtSignal(str, str)  # text, color
            update_wordcloud = pyqtSignal(str)
            update_meeting_timer = pyqtSignal(str)
            detected_key_phrases = pyqtSignal(list)
            detected_meeting_title = pyqtSignal(str)
        
        # --------------------------
        # Custom Text Highlighter
        # --------------------------
        class TranscriptHighlighter(QSyntaxHighlighter):
            def __init__(self, parent=None):
                super().__init__(parent)
                self.highlight_patterns = []
        
            def add_highlight_pattern(self, pattern, color):
                """Add a new pattern to highlight."""
                self.highlight_patterns.append((QRegExp(pattern, Qt.CaseInsensitive), QColor(color)))
                self.rehighlight()
        
            def clear_highlight_patterns(self):
                """Clear all highlight patterns."""
                self.highlight_patterns = []
                self.rehighlight()
        
            def highlightBlock(self, text):
                for pattern, color in self.highlight_patterns:
                    expression = QRegExp(pattern)
                    index = expression.indexIn(text)
                    while index >= 0:
                        length = expression.matchedLength()
                        fmt = QTextCharFormat()
                        fmt.setBackground(color)
                        self.setFormat(index, length, fmt)
                        index = expression.indexIn(text, index + length)
        
        # --------------------------
        # Custom Text Browser with Highlighting
        # --------------------------
        class CustomTextEdit(QTextEdit):
            def __init__(self, parent=None):
                super().__init__(parent)
                self.highlighter = TranscriptHighlighter(self.document())
                self.setReadOnly(True)
                # Set default formatting
                self.document().setDefaultFont(QFont(DEFAULT_FONT_FAMILY, DEFAULT_FONT_SIZE))
        
            def highlight_text(self, pattern, color="#FFFF00"):
                """Highlight text matching the pattern."""
                self.highlighter.add_highlight_pattern(pattern, color)
        
            def clear_highlights(self):
                """Clear all highlights."""
                self.highlighter.clear_highlight_patterns()
        
            def append_formatted(self, text, format_specs=None):
                """Append text with custom formatting."""
                cursor = self.textCursor()
                cursor.movePosition(cursor.End)
        
                if format_specs:
                    # Apply specified formatting
                    text_format = QTextCharFormat()
                    if "font_size" in format_specs:
                        text_format.setFontPointSize(format_specs["font_size"])
                    if "bold" in format_specs and format_specs["bold"]:
                        text_format.setFontWeight(QFont.Bold)
                    if "color" in format_specs:
                        text_format.setForeground(QBrush(QColor(format_specs["color"])))
                    if "font_family" in format_specs:
                        text_format.setFontFamily(format_specs["font_family"])
        
                    cursor.insertText(text, text_format)
                else:
                    # Default formatting
                    cursor.insertText(text)
        
                # Add a newline
                cursor.insertBlock()
        
                # Ensure visible
                self.setTextCursor(cursor)
                self.ensureCursorVisible()
        
        # --------------------------
        # WordCloud Window
        # --------------------------
        class WordCloudWindow(QWidget):
            def __init__(self, parent=None):
                super().__init__(parent)
                self.setWindowTitle("Live Word Cloud")
                self.setGeometry(300, 300, 800, 500)
                self.label = QLabel(self)
                self.label.setAlignment(Qt.AlignCenter)
        
                # Add auto-refresh option
                self.auto_refresh = QCheckBox("Auto-refresh (every 30 seconds)")
                self.auto_refresh.setChecked(True)
        
                # Add refresh button
                self.refresh_btn = QPushButton("Refresh Now")
        
                # Export button
                self.export_btn = QPushButton("Export Word Cloud")
        
                # Layout
                control_layout = QHBoxLayout()
                control_layout.addWidget(self.auto_refresh)
                control_layout.addWidget(self.refresh_btn)
                control_layout.addWidget(self.export_btn)
        
                layout = QVBoxLayout()
                layout.addWidget(self.label)
                layout.addLayout(control_layout)
        
                self.setLayout(layout)
        
                # Timer for auto-refresh
                self.timer = QTimer()
                self.timer.timeout.connect(self.auto_refresh_wordcloud)
                self.timer.start(30000)  # 30 seconds
        
            def update_wordcloud(self, text):
                """Update the word cloud with new text."""
                if not text:
                    return
        
                image_data = generate_wordcloud(text)
                pixmap = QPixmap()
                pixmap.loadFromData(image_data)
                self.label.setPixmap(pixmap.scaled(self.label.size(), Qt.KeepAspectRatio, Qt.SmoothTransformation))
        
            def auto_refresh_wordcloud(self):
                """Called by timer to refresh the word cloud."""
                if self.auto_refresh.isChecked() and self.isVisible():
                    # Signal to parent to refresh
                    self.parent().refresh_wordcloud()
        
            def export_wordcloud(self, file_path):
                """Export the current word cloud to an image file."""
                if not self.label.pixmap():
                    return False
        
                return self.label.pixmap().save(file_path)# --------------------------
        # Word Cloud Generator
        # --------------------------
        
        def generate_wordcloud(text):
            wc = WordCloud(width=800, height=400, background_color='black', colormap='Blues').generate(text)
            buffer = BytesIO()
            wc.to_image().save(buffer, format="PNG")
            return buffer.getvalue()
        
        # --------------------------
        # Name Anonymizer
        # --------------------------
        
        def anonymize_names(text, known_names=None):
            """Replace detected names with anonymous identifiers."""
            if not known_names:
                known_names = []
        
            # Simple pattern to detect potential names (starting with capital letter)
            name_pattern = r'\b[A-Z][a-z]+\b'
        
            # Find all potential names
            potential_names = set(re.findall(name_pattern, text))
        
            # Add known names
            all_names = potential_names.union(set(known_names))
        
            # Create a mapping of names to anonymous identifiers
            name_mapping = {}
            for i, name in enumerate(all_names):
                name_mapping[name] = f"Person-{i+1}"
        
            # Replace names in text
            anonymized_text = text
            for name, anon_id in name_mapping.items():
                anonymized_text = re.sub(r'\b' + name + r'\b', anon_id, anonymized_text)
        
            return anonymized_text, name_mapping
        
        # --------------------------
        # Simple Spell Checker
        # --------------------------
        
        def correct_spelling(text):
            """Simple spell checking and correction."""
            # Common misspellings and corrections
            corrections = {
                "teh": "the",
                "taht": "that",
                "thier": "their",
                "waht": "what",
                "dont": "don't",
                "cant": "can't",
                "wont": "won't",
                "im": "I'm",
                "youre": "you're",
                "theyre": "they're",
                "thats": "that's",
                "isnt": "isn't",
                "arent": "aren't",
                "hasnt": "hasn't",
                "havent": "haven't",
                "doesnt": "doesn't",
                "didnt": "didn't",
                "wasnt": "wasn't",
                "werent": "weren't",
                "wouldnt": "wouldn't",
                "shouldnt": "shouldn't",
                "couldnt": "couldn't",
                "alot": "a lot",
                "alright": "all right",
                "aswell": "as well",
                "becuase": "because",
                "definately": "definitely",
                "wierd": "weird",
                "recieve": "receive",
                "seperate": "separate",
                "untill": "until",
                "tommorow": "tomorrow",
                "tommorrow": "tomorrow",
                "alway": "always",
                "ofcourse": "of course",
                "ofcource": "of course",
                "allready": "already",
                "allright": "all right",
                "accross": "across",
                "beleive": "believe",
                "comunicate": "communicate",
                "greatful": "grateful",
                "dependance": "dependence",
                "excelent": "excellent",
                "existance": "existence",
                "foreward": "forward",
                "grammer": "grammar",
                "gaurd": "guard",
                "hight": "height",
                "humerous": "humorous",
                "independant": "independent",
                "liason": "liaison",
                
                # --------------------------
                # Settings Dialog
                # --------------------------
                class SettingsDialog(QDialog):
                        def __init__(self, parent=None, settings=None):
                                super().__init__(parent)
                                self.settings = settings if settings else {}
                                self.setWindowTitle("Settings")
                                self.setMinimumWidth(650)
                                self.setMinimumHeight(550)
                
                                # Apply macOS-style window
                                self.setWindowFlags(Qt.Dialog | Qt.WindowCloseButtonHint | Qt.WindowTitleHint)
                                self.setStyleSheet("""
                                        QDialog {
                                                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                                                                                                    stop:0 #F7F7F7, 
                                                                                                    stop:1 #E6E6E6);
                                                border-radius: 10px;
                                        }
                                """)
                
                                self.setup_ui()
                                self.load_settings()
                
                        def setup_ui(self):
                                layout = QVBoxLayout()
                                layout.setContentsMargins(20, 20, 20, 20)
                                layout.setSpacing(15)
                
                                # Create modern tabs with styling
                                tabs = QTabWidget()
                                tabs.setStyleSheet("""
                                        QTabWidget::pane {
                                                border: 1px solid #C0C0C0;
                                                border-radius: 6px;
                                                background-color: white;
                                                padding: 10px;
                                        }
                                        QTabBar::tab {
                                                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                                                                                                    stop:0 #F2F2F2, 
                                                                                                    stop:1 #E0E0E0);
                                                border: 1px solid #C0C0C0;
                                                border-bottom: none;
                                                border-top-left-radius: 6px;
                                                border-top-right-radius: 6px;
                                                padding: 8px 16px;
                                                margin-right: 2px;
                                                color: #404040;
                                                font-weight: 500;
                                        }
                                        QTabBar::tab:selected {
                                                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                                                                                                    stop:0 #FFFFFF, 
                                                                                                    stop:1 #F0F0F0);
                                                border-bottom-color: #FFFFFF;
                                                font-weight: bold;
                                        }
                                        QTabBar::tab:hover {
                                                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                                                                                                    stop:0 #FFFFFF, 
                                                                                                    stop:1 #F5F5F5);
                                        }
                                """)
                
                                # General Settings Tab with modern styling
                                general_tab = QWidget()
                                general_layout = QFormLayout()
                                general_layout.setVerticalSpacing(12)
                                general_layout.setHorizontalSpacing(15)
                                general_layout.setContentsMargins(15, 15, 15, 15)
                
                                # Theme selector with visual preview
                                theme_group = QGroupBox("Appearance")
                                theme_group.setStyleSheet("""
                                        QGroupBox {
                                                font-weight: bold;
                                                border: 1px solid #C0C0C0;
                                                border-radius: 6px;
                                                margin-top: 12px;
                                                padding-top: 15px;
                                        }
                                        QGroupBox::title {
                                                subcontrol-origin: margin;
                                                left: 10px;
                                                padding: 0 5px 0 5px;
                                        }
                                """)
                
                                theme_layout = QVBoxLayout()
                                theme_layout.setContentsMargins(10, 10, 10, 10)
                
                                theme_label = QLabel("Theme:")
                                theme_label.setStyleSheet("font-weight: bold;")
                
                                self.theme_combo = QComboBox()
                                self.theme_combo.addItems(list(THEMES.keys()))
                                self.theme_combo.setStyleSheet("""
                                        QComboBox {
                                                border: 1px solid #C0C0C0;
                                                border-radius: 4px;
                                                padding: 5px 10px;
                                                background: white;
                                                min-width: 200px;
                                        }
                                        QComboBox::drop-down {
                                                subcontrol-origin: padding;
                                                subcontrol-position: top right;
                                                width: 20px;
                                                border-left: 1px solid #C0C0C0;
                                                border-top-right-radius: 4px;
                                                border-bottom-right-radius: 4px;
                                        }
                                """)
                
                                # Theme preview swatches
                                theme_previews = QHBoxLayout()
                
                                for theme_name in THEMES.keys():
                                        preview = QFrame()
                                        preview.setFixedSize(40, 40)
                                        preview.setStyleSheet(f"""
                                                QFrame {{
                                                        background-color: {THEMES[theme_name]['window']};
                                                        border: 1px solid #C0C0C0;
                                                        border-radius: 6px;
                                                }}
                                        """)
                
                                        preview_layout = QVBoxLayout(preview)
                                        preview_layout.setContentsMargins(2, 2, 2, 2)
                
                                        accent = QFrame(preview)
                                        accent.setFixedSize(15, 15)
                                        accent.setStyleSheet(f"""
                                                QFrame {{
                                                        background-color: {THEMES[theme_name]['highlight']};
                                                        border-radius: 7px;
                                                }}
                                        """)
                
                                        preview_layout.addWidget(accent, 0, Qt.AlignCenter)
                
                                        theme_previews.addWidget(preview)
                
                                theme_layout.addWidget(theme_label)
                                theme_layout.addWidget(self.theme_combo)
                                theme_layout.addLayout(theme_previews)
                                theme_group.setLayout(theme_layout)
        # Start in background
                startup_group = QGroupBox("Startup")
                startup_group.setStyleSheet("""
                    QGroupBox {
                        font-weight: bold;
                        border: 1px solid #C0C0C0;
                        border-radius: 6px;
                        margin-top: 12px;
                        padding-top: 15px;
                    }
                    QGroupBox::title {
                        subcontrol-origin: margin;
                        left: 10px;
                        padding: 0 5px 0 5px;
                    }
                """)
                
                startup_layout = QVBoxLayout()
                self.start_bg_check = QCheckBox("Start application in background (system tray)")
                self.start_bg_check.setStyleSheet("""
                    QCheckBox {
                        spacing: 8px;
                    }
                    QCheckBox::indicator {
                        width: 18px;
                        height: 18px;
                        border: 1px solid #C0C0C0;
                        border-radius: 4px;
                    }
                    QCheckBox::indicator:checked {
                        background-color: #007AFF;
                        border: 1px solid #007AFF;
                        image: url(data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxOCIgaGVpZ2h0PSIxOCIgdmlld0JveD0iMCAwIDE4IDE4Ij48cGF0aCBmaWxsPSIjRkZGRkZGIiBkPSJNNi45NTI4MTYxLDEyLjc2MzYzNjQgTDMuMTEzMTI5ODcsOC45MjM5NTgyNCBMMiw5Ljk5NzgwMjc0IEw2Ljk1MjgxNjEsMTQuOTUwNjI2OSBMMTYuMTc2Nzc2OCw1LjcyNjY2NjE4IEwxNS4wNDI5MzMxLDQuNTkyODIxNDQgTDYuOTUyODE2MSwxMi43NjM2MzY0IFoiLz48L3N2Zz4=);
                    }
                    QCheckBox::indicator:unchecked:hover {
                        border: 1px solid #007AFF;
                    }
                """)
                
                startup_layout.addWidget(self.start_bg_check)
                startup_group.setLayout(startup_layout)
                
                # Special visual effects
                effects_group = QGroupBox("Visual Effects")
                effects_group.setStyleSheet("""
                    QGroupBox {
                        font-weight: bold;
                        border: 1px solid #C0C0C0;
                        border-radius: 6px;
                        margin-top: 12px;
                        padding-top: 15px;
                    }
                    QGroupBox::title {
                        subcontrol-origin: margin;
                        left: 10px;
                        padding: 0 5px 0 5px;
                    }
                """)
                
                effects_layout = QVBoxLayout()
                
                self.rounded_corners_check = QCheckBox("Use rounded corners")
                self.animations_check = QCheckBox("Enable animations")
                self.transparency_check = QCheckBox("Enable transparency effects")
                
                for checkbox in [self.rounded_corners_check, self.animations_check, self.transparency_check]:
                    checkbox.setStyleSheet("""
                        QCheckBox {
                            spacing: 8px;
                        }
                        QCheckBox::indicator {
                            width: 18px;
                            height: 18px;
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                        }
                        QCheckBox::indicator:checked {
                            background-color: #007AFF;
                            border: 1px solid #007AFF;
                            image: url(data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxOCIgaGVpZ2h0PSIxOCIgdmlld0JveD0iMCAwIDE4IDE4Ij48cGF0aCBmaWxsPSIjRkZGRkZGIiBkPSJNNi45NTI4MTYxLDEyLjc2MzYzNjQgTDMuMTEzMTI5ODcsOC45MjM5NTgyNCBMMiw5Ljk5NzgwMjc0IEw2Ljk1MjgxNjEsMTQuOTUwNjI2OSBMMTYuMTc2Nzc2OCw1LjcyNjY2NjE4IEwxNS4wNDI5MzMxLDQuNTkyODIxNDQgTDYuOTUyODE2MSwxMi43NjM2MzY0IFoiLz48L3N2Zz4=);
                        }
                        QCheckBox::indicator:unchecked:hover {
                            border: 1px solid #007AFF;
                        }
                    """)
                
                    effects_layout.addWidget(checkbox)
                
                effects_group.setLayout(effects_layout)
                
                # Keyboard shortcut
                shortcuts_group = QGroupBox("Keyboard Shortcuts")
                shortcuts_group.setStyleSheet("""
                    QGroupBox {
                        font-weight: bold;
                        border: 1px solid #C0C0C0;
                        border-radius: 6px;
                        margin-top: 12px;
                        padding-top: 15px;
                    }
                    QGroupBox::title {
                        subcontrol-origin: margin;
                        left: 10px;
                        padding: 0 5px 0 5px;
                    }
                """)
                
                shortcuts_layout = QFormLayout()
                shortcuts_layout.setVerticalSpacing(10)
                
                self.shortcut_edit = QLineEdit()
                self.shortcut_edit.setStyleSheet("""
                    QLineEdit {
                        border: 1px solid #C0C0C0;
                        border-radius: 4px;
                        padding: 6px 10px;
                        background: white;
                        min-width: 120px;
                    }
                    QLineEdit:focus {
                        border: 1px solid #007AFF;
                    }
                """)
                shortcuts_layout.addRow("Analysis shortcut:", self.shortcut_edit)
                
                # Code word
                self.code_word_edit = QLineEdit()
                self.code_word_edit.setStyleSheet("""
                    QLineEdit {
                        border: 1px solid #C0C0C0;
                        border-radius: 4px;
                        padding: 6px 10px;
                        background: white;
                        min-width: 120px;
                    }
                    QLineEdit:focus {
                        border: 1px solid #007AFF;
                    }
                """)
                shortcuts_layout.addRow("Trigger code word:", self.code_word_edit)
                
                shortcuts_group.setLayout(shortcuts_layout)
                
                # Font settings
                font_group = QGroupBox("Font Settings")
                font_group.setStyleSheet("""
                    QGroupBox {
                        font-weight: bold;
                        border: 1px solid #C0C0C0;
                        border-radius: 6px;
                        margin-top: 12px;
                        padding-top: 15px;
                    }
                    QGroupBox::title {
                        subcontrol-origin: margin;
                        left: 10px;
                        padding: 0 5px 0 5px;
                    }
                """)
                
                font_layout = QFormLayout()
                font_layout.setVerticalSpacing(10)
                
                self.font_family_combo = QComboBox()
                self.font_family_combo.addItems([
                    "SF Pro Text", "SF Pro Display", "Arial", "Verdana", "Roboto", 
                    "Times New Roman", "Helvetica Neue", "Courier New", "Consolas"
                ])
            self.font_family_combo.setStyleSheet("""
                        QComboBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 5px 10px;
                            background: white;
                            min-width: 200px;
                        }
                        QComboBox::drop-down {
                            subcontrol-origin: padding;
                            subcontrol-position: top right;
                            width: 20px;
                            border-left: 1px solid #C0C0C0;
                            border-top-right-radius: 4px;
                            border-bottom-right-radius: 4px;
                        }
                    """)
                    font_layout.addRow("Font family:", self.font_family_combo)
                
                    # Font size slider
                    self.font_size_spin = QSpinBox()
                    self.font_size_spin.setRange(8, 24)
                    self.font_size_spin.setStyleSheet("""
                        QSpinBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 5px 10px;
                            background: white;
                            min-width: 80px;
                        }
                        QSpinBox::up-button, QSpinBox::down-button {
                            border-left: 1px solid #C0C0C0;
                            width: 20px;
                            height: 12px;
                        }
                        QSpinBox::up-button {
                            border-top-right-radius: 4px;
                        }
                        QSpinBox::down-button {
                            border-bottom-right-radius: 4px;
                        }
                    """)
                
                    font_size_layout = QHBoxLayout()
                    font_size_layout.addWidget(self.font_size_spin)
                
                    # Preview label
                    font_preview = QLabel("Preview Text")
                    font_preview.setObjectName("font_preview")
                    font_preview.setAlignment(Qt.AlignCenter)
                    font_preview.setMinimumHeight(40)
                    font_preview.setStyleSheet("""
                        QLabel {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            background-color: white;
                            padding: 5px;
                        }
                    """)
                
                    # Update preview when font settings change
                    def update_preview():
                        font = QFont(self.font_family_combo.currentText(), self.font_size_spin.value())
                        font_preview.setFont(font)
                
                    self.font_family_combo.currentTextChanged.connect(update_preview)
                    self.font_size_spin.valueChanged.connect(update_preview)
                
                    font_size_layout.addWidget(font_preview)
                    font_layout.addRow("Font size:", font_size_layout)
                
                    self.line_spacing_spin = QSpinBox()
                    self.line_spacing_spin.setRange(100, 300)
                    self.line_spacing_spin.setSingleStep(10)
                    self.line_spacing_spin.setStyleSheet("""
                        QSpinBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 5px 10px;
                            background: white;
                            min-width: 80px;
                        }
                    """)
                    font_layout.addRow("Line spacing (%):", self.line_spacing_spin)
                
                    self.word_spacing_spin = QSpinBox()
                    self.word_spacing_spin.setRange(80, 200)
                    self.word_spacing_spin.setSingleStep(5)
                    self.word_spacing_spin.setStyleSheet("""
                        QSpinBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 5px 10px;
                            background: white;
                            min-width: 80px;
                        }
                    """)
                    font_layout.addRow("Word spacing (%):", self.word_spacing_spin)
                
                    # Custom font button
                    self.custom_font_btn = QPushButton("Choose Custom Font")
                    self.custom_font_btn.clicked.connect(self.choose_custom_font)
                    self.custom_font_btn.setStyleSheet("""
                        QPushButton {
                            background-color: #F2F2F2;
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 6px 12px;
                            font-weight: 500;
                        }
                        QPushButton:hover {
                            background-color: #E5E5E5;
                        }
                        QPushButton:pressed {
                            background-color: #D0D0D0;
                        }
                    """)
                    font_layout.addRow("", self.custom_font_btn)
                
                    font_group.setLayout(font_layout)
                
                    # Icon paths
                    icon_group = QGroupBox("Application Icons")
                    icon_group.setStyleSheet("""
                        QGroupBox {
                            font-weight: bold;
                            border: 1px solid #C0C0C0;
                            border-radius: 6px;
                            margin-top: 12px;
                            padding-top: 15px;
                        }
                        QGroupBox::title {
                            subcontrol-origin: margin;
                            left: 10px;
                            padding: 0 5px 0 5px;
                        }
                    """)
                
                    icon_layout = QFormLayout()
                    icon_layout.setVerticalSpacing(10)
                
                    icon_fields = [
                        ("App Icon:", "app_icon_edit"),
                        ("Generate Icon:", "generate_icon_edit"),
                        ("Settings Icon:", "settings_icon_edit")
                    ]
                
                    for label_text, field_name in icon_fields:
                        setattr(self, field_name, QLineEdit())
                        line_edit = getattr(self, field_name)
                        line_edit.setStyleSheet("""
                            QLineEdit {
                                border: 1px solid #C0C0C0;
                                border-radius: 4px;
                                padding: 6px 10px;
                                background: white;
                            }
                            QLineEdit:focus {
                                border: 1px solid #007AFF;
                            }
                        """)
                
                        browse_btn = QPushButton("Browse...")
                        browse_btn.clicked.connect(lambda checked=False, le=line_edit: self.browse_icon(le))
                        browse_btn.setStyleSheet("""
                            QPushButton {
                                background-color: #F2F2F2;
                                border: 1px solid #C0C0C0;
                                border-radius: 4px;
                                padding: 6px 12px;
                                font-weight: 500;
                            }
                            QPushButton:hover {
                                background-color: #E5E5E5;
                            }
                            QPushButton:pressed {
                                background-color: #D0D0D0;
                            }
                        """)
                
                        field_layout = QHBoxLayout()
                        field_layout.addWidget(line_edit)
                        field_layout.addWidget(browse_btn)
                        icon_layout.addRow(label_text, field_layout)
                
                    icon_group.setLayout(icon_layout)
                
                    # Add all these components to the General tab
                    general_layout.addRow(theme_group)
                    general_layout.addRow(startup_group)
                    general_layout.addRow(effects_group)
                    general_layout.addRow(shortcuts_group)
                    general_layout.addRow(font_group)
                    general_layout.addRow(icon_group)
                
                    general_tab.setLayout(general_layout)
            # AI Settings Tab
                    ai_tab = QWidget()
                    ai_layout = QVBoxLayout()
                    ai_layout.setContentsMargins(15, 15, 15, 15)
                    ai_layout.setSpacing(15)
                
                    # Model selection card
                    model_card = QFrame()
                    model_card.setFrameShape(QFrame.StyledPanel)
                    model_card.setStyleSheet("""
                        QFrame {
                            border: 1px solid #C0C0C0;
                            border-radius: 6px;
                            background-color: white;
                            padding: 10px;
                        }
                    """)
                
                    model_layout = QVBoxLayout(model_card)
                
                    model_title = QLabel("Groq LLM Model")
                    model_title.setStyleSheet("font-weight: bold; font-size: 14px;")
                    model_layout.addWidget(model_title)
                
                    model_description = QLabel("Select which Groq model to use for AI analysis:")
                    model_layout.addWidget(model_description)
                
                    self.model_combo = QComboBox()
                    self.model_combo.addItems(AVAILABLE_MODELS)
                    self.model_combo.setStyleSheet("""
                        QComboBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 8px 12px;
                            background: #F7F7F7;
                            min-width: 250px;
                            font-weight: 500;
                        }
                        QComboBox::drop-down {
                            subcontrol-origin: padding;
                            subcontrol-position: top right;
                            width: 20px;
                            border-left: 1px solid #C0C0C0;
                            border-top-right-radius: 4px;
                            border-bottom-right-radius: 4px;
                        }
                    """)
                    model_layout.addWidget(self.model_combo)
                
                    # Model descriptions
                    model_info = QLabel("""
                        <b>llama3-70b-8192:</b> Highest quality with deep reasoning<br>
                        <b>llama3-8b-8192:</b> Faster responses for simpler tasks<br>
                        <b>mixtral-8x7b-32768:</b> Balanced between speed and quality<br>
                        <b>gemma-7b-it:</b> Efficient for straightforward analyses
                    """)
                    model_info.setStyleSheet("""
                        QLabel {
                            background-color: #F5F5F5;
                            border-radius: 4px;
                            padding: 8px;
                            color: #555;
                            font-size: 12px;
                        }
                    """)
                    model_layout.addWidget(model_info)
                
                    ai_layout.addWidget(model_card)
            # Summary intervals card
                    interval_card = QFrame()
                    interval_card.setFrameShape(QFrame.StyledPanel)
                    interval_card.setStyleSheet("""
                        QFrame {
                            border: 1px solid #C0C0C0;
                            border-radius: 6px;
                            background-color: white;
                            padding: 10px;
                        }
                    """)
                
                    interval_layout = QVBoxLayout(interval_card)
                
                    interval_title = QLabel("Analysis Frequency")
                    interval_title.setStyleSheet("font-weight: bold; font-size: 14px;")
                    interval_layout.addWidget(interval_title)
                
                    interval_form = QFormLayout()
                    interval_form.setVerticalSpacing(10)
                
                    self.line_interval_spin = QSpinBox()
                    self.line_interval_spin.setRange(3, 50)
                    self.line_interval_spin.setStyleSheet("""
                        QSpinBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 6px 10px;
                            background: #F7F7F7;
                            min-width: 80px;
                        }
                    """)
                    interval_form.addRow("Summarize every N lines:", self.line_interval_spin)
                
                    self.time_interval_spin = QSpinBox()
                    self.time_interval_spin.setRange(30, 600)
                    self.time_interval_spin.setSingleStep(30)
                    self.time_interval_spin.setStyleSheet("""
                        QSpinBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 6px 10px;
                            background: #F7F7F7;
                            min-width: 80px;
                        }
                    """)
                    interval_form.addRow("Summarize every N seconds:", self.time_interval_spin)
                
                    interval_layout.addLayout(interval_form)
                
                    ai_layout.addWidget(interval_card)
            # Prompt customization card
                    prompt_card = QFrame()
                    prompt_card.setFrameShape(QFrame.StyledPanel)
                    prompt_card.setStyleSheet("""
                        QFrame {
                            border: 1px solid #C0C0C0;
                            border-radius: 6px;
                            background-color: white;
                            padding: 10px;
                        }
                    """)
                
                    prompt_layout = QVBoxLayout(prompt_card)
                
                    prompt_title = QLabel("Prompt Templates")
                    prompt_title.setStyleSheet("font-weight: bold; font-size: 14px;")
                    prompt_layout.addWidget(prompt_title)
                
                    prompt_description = QLabel("Customize the instructions given to the AI:")
                    prompt_layout.addWidget(prompt_description)
                
                    prompt_form = QFormLayout()
                
                    self.prompt_combo = QComboBox()
                    self.prompt_combo.setStyleSheet("""
                        QComboBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 8px 12px;
                            background: #F7F7F7;
                            min-width: 250px;
                            font-weight: 500;
                        }
                        QComboBox::drop-down {
                            subcontrol-origin: padding;
                            subcontrol-position: top right;
                            width: 20px;
                            border-left: 1px solid #C0C0C0;
                            border-top-right-radius: 4px;
                            border-bottom-right-radius: 4px;
                        }
                    """)
                    prompt_form.addRow("Active Prompt:", self.prompt_combo)
                
                    prompt_layout.addLayout(prompt_form)
                
                    self.prompt_edit = QTextEdit()
                    self.prompt_edit.setMinimumHeight(150)
                    self.prompt_edit.setStyleSheet("""
                        QTextEdit {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 8px;
                            background: #F7F7F7;
                            font-family: "Menlo", monospace;
                            font-size: 13px;
                        }
                    """)
                    prompt_layout.addWidget(self.prompt_edit)
                
                    # Add/edit/delete prompt buttons
                    prompt_btn_layout = QHBoxLayout()
                
                    add_prompt_btn = QPushButton("Add New")
                    add_prompt_btn.setIcon(QIcon.fromTheme("document-new"))
                    add_prompt_btn.setStyleSheet("""
                        QPushButton {
                            background-color: #34C759;
                            color: white;
                            border: none;
                            border-radius: 4px;
                            padding: 8px 15px;
                            font-weight: 500;
                        }
                        QPushButton:hover {
                            background-color: #30B955;
                        }
                        QPushButton:pressed {
                            background-color: #2CAA50;
                        }
                    """)
                
                    save_prompt_btn = QPushButton("Save Changes")
                    save_prompt_btn.setIcon(QIcon.fromTheme("document-save"))
                    save_prompt_btn.setStyleSheet("""
                        QPushButton {
                            background-color: #007AFF;
                            color: white;
                            border: none;
                            border-radius: 4px;
                            padding: 8px 15px;
                            font-weight: 500;
                        }
                        QPushButton:hover {
                            background-color: #0071EA;
                        }
                        QPushButton:pressed {
                            background-color: #0068D7;
                        }
                    """)
                
                    delete_prompt_btn = QPushButton("Delete")
                    delete_prompt_btn.setIcon(QIcon.fromTheme("edit-delete"))
                    delete_prompt_btn.setStyleSheet("""
                        QPushButton {
                            background-color: #FF3B30;
                            color: white;
                            border: none;
                            border-radius: 4px;
                            padding: 8px 15px;
                            font-weight: 500;
                        }
                        QPushButton:hover {
                            background-color: #E0352B;
                        }
                        QPushButton:pressed {
                            background-color: #D03027;
                        }
                    """)
                
                    add_prompt_btn.clicked.connect(self.add_new_prompt)
                    save_prompt_btn.clicked.connect(self.save_prompt_changes)
                    delete_prompt_btn.clicked.connect(self.delete_prompt)
                
                    prompt_btn_layout.addWidget(add_prompt_btn)
                    prompt_btn_layout.addWidget(save_prompt_btn)
                    prompt_btn_layout.addWidget(delete_prompt_btn)
                    prompt_layout.addLayout(prompt_btn_layout)
                
                    # Connect prompt combo to update text edit
                    self.prompt_combo.currentTextChanged.connect(self.update_prompt_edit)
                
                    ai_layout.addWidget(prompt_card)
                
                    ai_tab.setLayout(ai_layout)
            # Advanced Settings Tab
                    adv_tab = QWidget()
                    adv_layout = QVBoxLayout()
                    adv_layout.setContentsMargins(15, 15, 15, 15)
                    adv_layout.setSpacing(15)
                
                    # Polling settings card
                    polling_card = QFrame()
                    polling_card.setFrameShape(QFrame.StyledPanel)
                    polling_card.setStyleSheet("""
                        QFrame {
                            border: 1px solid #C0C0C0;
                            border-radius: 6px;
                            background-color: white;
                            padding: 10px;
                        }
                    """)
                
                    polling_layout = QVBoxLayout(polling_card)
                
                    polling_title = QLabel("Transcript Polling")
                    polling_title.setStyleSheet("font-weight: bold; font-size: 14px;")
                    polling_layout.addWidget(polling_title)
                
                    polling_description = QLabel("How frequently to check for new transcript content:")
                    polling_layout.addWidget(polling_description)
                
                    self.poll_interval_spin = QSpinBox()
                    self.poll_interval_spin.setRange(1, 30)
                    self.poll_interval_spin.setStyleSheet("""
                        QSpinBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 6px 10px;
                            background: #F7F7F7;
                            min-width: 80px;
                        }
                    """)
                
                    polling_form = QFormLayout()
                    polling_form.addRow("Polling interval (seconds):", self.poll_interval_spin)
                
                    polling_note = QLabel("Note: Lower values increase responsiveness but may use more resources.")
                    polling_note.setStyleSheet("font-style: italic; color: #666;")
                
                    polling_layout.addLayout(polling_form)
                    polling_layout.addWidget(polling_note)
                
                    adv_layout.addWidget(polling_card)
                
                    # Transcript features card
                    transcript_card = QFrame()
                    transcript_card.setFrameShape(QFrame.StyledPanel)
                    transcript_card.setStyleSheet("""
                        QFrame {
                            border: 1px solid #C0C0C0;
                            border-radius: 6px;
                            background-color: white;
                            padding: 10px;
                        }
                    """)
                
                    transcript_layout = QVBoxLayout(transcript_card)
                
                    transcript_title = QLabel("Transcript Features")
                    transcript_title.setStyleSheet("font-weight: bold; font-size: 14px;")
                    transcript_layout.addWidget(transcript_title)
                
                    # Anonymize transcript
                    self.anonymize_check = QCheckBox("Anonymize names in transcript")
                    self.anonymize_check.setStyleSheet("""
                        QCheckBox {
                            spacing: 8px;
                        }
                        QCheckBox::indicator {
                            width: 18px;
                            height: 18px;
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                        }
                        QCheckBox::indicator:checked {
                            background-color: #007AFF;
                            border: 1px solid #007AFF;
                            image: url(data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxOCIgaGVpZ2h0PSIxOCIgdmlld0JveD0iMCAwIDE4IDE4Ij48cGF0aCBmaWxsPSIjRkZGRkZGIiBkPSJNNi45NTI4MTYxLDEyLjc2MzYzNjQgTDMuMTEzMTI5ODcsOC45MjM5NTgyNCBMMiw5Ljk5NzgwMjc0IEw2Ljk1MjgxNjEsMTQuOTUwNjI2OSBMMTYuMTc2Nzc2OCw1LjcyNjY2NjE4IEwxNS4wNDI5MzMxLDQuNTkyODIxNDQgTDYuOTUyODE2MSwxMi43NjM2MzY0IFoiLz48L3N2Zz4=);
                        }
                        QCheckBox::indicator:unchecked:hover {
                            border: 1px solid #007AFF;
                        }
                    """)
                    transcript_layout.addWidget(self.anonymize_check)
                
                    anonymize_description = QLabel("Replaces detected names with generic identifiers (Person-1, Person-2, etc.)")
                    anonymize_description.setStyleSheet("font-size: 12px; color: #666; margin-left: 26px;")
                    transcript_layout.addWidget(anonymize_description)
                
                    # Spell checking
                    self.spellcheck_check = QCheckBox("Enable spell checking")
                    self.spellcheck_check.setStyleSheet("""
                        QCheckBox {
                            spacing: 8px;
                            margin-top: 10px;
                        }
                        QCheckBox::indicator {
                            width: 18px;
                            height: 18px;
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                        }
                        QCheckBox::indicator:checked {
                            background-color: #007AFF;
                            border: 1px solid #007AFF;
                            image: url(data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxOCIgaGVpZ2h0PSIxOCIgdmlld0JveD0iMCAwIDE4IDE4Ij48cGF0aCBmaWxsPSIjRkZGRkZGIiBkPSJNNi45NTI4MTYxLDEyLjc2MzYzNjQgTDMuMTEzMTI5ODcsOC45MjM5NTgyNCBMMiw5Ljk5NzgwMjc0IEw2Ljk1MjgxNjEsMTQuOTUwNjI2OSBMMTYuMTc2Nzc2OCw1LjcyNjY2NjE4IEwxNS4wNDI5MzMxLDQuNTkyODIxNDQgTDYuOTUyODE2MSwxMi43NjM2MzY0IFoiLz48L3N2Zz4=);
                        }
                        QCheckBox::indicator:unchecked:hover {
                            border: 1px solid #007AFF;
                        }
                    """)
                    transcript_layout.addWidget(self.spellcheck_check)
                
                    spellcheck_description = QLabel("Automatically corrects common spelling errors in the transcript")
                    spellcheck_description.setStyleSheet("font-size: 12px; color: #666; margin-left: 26px;")
                    transcript_layout.addWidget(spellcheck_description)
                
                    adv_layout.addWidget(transcript_card)
                
                    # Meeting timer card
                    timer_card = QFrame()
                    timer_card.setFrameShape(QFrame.StyledPanel)
                    timer_card.setStyleSheet("""
                        QFrame {
                            border: 1px solid #C0C0C0;
                            border-radius: 6px;
                            background-color: white;
                            padding: 10px;
                        }
                    """)
                
                    timer_layout = QVBoxLayout(timer_card)
                
                    timer_title = QLabel("Meeting Timer")
                    timer_title.setStyleSheet("font-weight: bold; font-size: 14px;")
                    timer_layout.addWidget(timer_title)
                
                    # Timer enabled
                    self.timer_check = QCheckBox("Enable meeting timer")
                    self.timer_check.setStyleSheet("""
                        QCheckBox {
                            spacing: 8px;
                        }
                        QCheckBox::indicator {
                            width: 18px;
                            height: 18px;
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                        }
                        QCheckBox::indicator:checked {
                            background-color: #007AFF;
                            border: 1px solid #007AFF;
                            image: url(data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxOCIgaGVpZ2h0PSIxOCIgdmlld0JveD0iMCAwIDE4IDE4Ij48cGF0aCBmaWxsPSIjRkZGRkZGIiBkPSJNNi45NTI4MTYxLDEyLjc2MzYzNjQgTDMuMTEzMTI5ODcsOC45MjM5NTgyNCBMMiw5Ljk5NzgwMjc0IEw2Ljk1MjgxNjEsMTQuOTUwNjI2OSBMMTYuMTc2Nzc2OCw1LjcyNjY2NjE4IEwxNS4wNDI5MzMxLDQuNTkyODIxNDQgTDYuOTUyODE2MSwxMi43NjM2MzY0IFoiLz48L3N2Zz4=);
                        }
                        QCheckBox::indicator:unchecked:hover {
                            border: 1px solid #007AFF;
                        }
                    """)
                    timer_layout.addWidget(self.timer_check)
                
                    timer_form = QFormLayout()
                    timer_form.setVerticalSpacing(10)
                
                    self.reminder_spin = QSpinBox()
                    self.reminder_spin.setRange(5, 60)
                    self.reminder_spin.setSingleStep(5)
                    self.reminder_spin.setStyleSheet("""
                        QSpinBox {
                            border: 1px solid #C0C0C0;
                            border-radius: 4px;
                            padding: 6px 10px;
                            background: #F7F7F7;
                            min-width: 80px;
                        }
                    """)
                    timer_form.addRow("Meeting reminder interval (minutes):", self.reminder_spin)
                
                    timer_layout.addLayout(timer_form)
                
                    timer_note = QLabel("Notifications will remind you how long the meeting has been running.")
                    timer_note.setStyleSheet("font-style: italic; color: #666;")
                    timer_layout.addWidget(timer_note)
                
                    adv_layout.addWidget(timer_card)
                
                    # Add stretch to push everything to the top
                    adv_layout.addStretch()
                
                    adv_tab.setLayout(adv_layout)
                
                    # Add tabs to widget
                    tabs.addTab(general_tab, "General")
                    tabs.addTab(ai_tab, "AI Settings")
                    tabs.addTab(adv_tab, "Advanced")
                
                    layout.addWidget(tabs)
                
                    # Buttons
                    button_layout = QHBoxLayout()
                
                    reset_btn = QPushButton("Reset to Defaults")
                    reset_btn.setStyleSheet("""
                        QPushButton {
                            background-color: #E0E0E0;
                            border: 1px solid #C0C0C0;
                            border-radius: 6px;
                            padding: 10px 18px;
                            font-weight: 500;
                            color: #404040;
                        }
                        QPushButton:hover {
                            background-color: #D0D0D0;
                        }
                        QPushButton:pressed {
                            background-color: #C0C0C0;
                        }
                    """)
                
                    save_btn = QPushButton("Save")
                    save_btn.setStyleSheet("""
                        QPushButton {
                            background-color: #007AFF;
                            color: white;
                            border: none;
                            border-radius: 6px;
                            padding: 10px 25px;
                            font-weight: 600;
                        }
                        QPushButton:hover {
                            background-color: #0071EA;
                        }
                        QPushButton:pressed {
                            background-color: #0068D7;
                        }
                    """)
                
                    cancel_btn = QPushButton("Cancel")
                    cancel_btn.setStyleSheet("""
                        QPushButton {
                            background-color: #F2F2F2;
                            border: 1px solid #D0D0D0;
                            border-radius: 6px;
                            padding: 10px 18px;
                            font-weight: 500;
                        }
                        QPushButton:hover {
                            background-color: #E5E5E5;
                        }
                        QPushButton:pressed {
                            background-color: #D8D8D8;
                        }
                    """)
                
                    reset_btn.clicked.connect(self.reset_to_defaults)
                    save_btn.clicked.connect(self.save_settings)
                    cancel_btn.clicked.connect(self.reject)
                
                    button_layout.addWidget(reset_btn)
                    button_layout.addStretch()
                    button_layout.addWidget(save_btn)
                    button_layout.addWidget(cancel_btn)
                
                    layout.addLayout(button_layout)
                
                    self.setLayout(layout)
            def choose_custom_font(self):
                    """Open font dialog to choose a custom font."""
                    current_font = QFont(self.font_family_combo.currentText(), self.font_size_spin.value())
                    font, ok = QFontDialog.getFont(current_font, self)
                    if ok:
                        # Update font family and size if a font was selected
                        self.font_family_combo.setCurrentText(font.family())
                        self.font_size_spin.setValue(font.pointSize())
                
                        # Update the preview
                        font_preview = self.findChild(QLabel, "font_preview")
                        if font_preview:
                            font_preview.setFont(font)
                
                def browse_icon(self, line_edit):
                    """Open file dialog to choose an icon file."""
                    file_path, _ = QFileDialog.getOpenFileName(
                        self, "Select Icon", "", "Image Files (*.png *.jpg *.jpeg *.ico *.svg)"
                    )
                    if file_path:
                        line_edit.setText(file_path)
                
                        # Show preview of the selected icon
                        try:
                            pixmap = QPixmap(file_path)
                            if not pixmap.isNull():
                                icon_preview = QLabel()
                                icon_preview.setPixmap(pixmap.scaled(32, 32, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                
                                # Add preview next to the text field
                                parent_layout = line_edit.parent().layout()
                
                                # Check if we already have a preview
                                for i in range(parent_layout.count()):
                                    item = parent_layout.itemAt(i)
                                    if item.widget() and isinstance(item.widget(), QLabel) and item.widget() != line_edit:
                                        # Replace existing preview
                                        item.widget().deleteLater()
                                        parent_layout.replaceWidget(item.widget(), icon_preview)
                                        return
                
                                # No preview yet, add a new one
                                parent_layout.addWidget(icon_preview)
                        except Exception as e:
                            print(f"Error previewing icon: {e}")
            def add_new_prompt(self):
                    """Add a new prompt template."""
                    prompt_name, ok = QInputDialog.getText(self, "New Prompt", "Enter prompt name:")
                    if ok and prompt_name:
                        # Check if name already exists
                        if prompt_name in [self.prompt_combo.itemText(i) for i in range(self.prompt_combo.count())]:
                            QMessageBox.warning(self, "Duplicate Name", "A prompt with this name already exists.")
                            return
                
                        # Add to combo and set as current
                        self.prompt_combo.addItem(prompt_name)
                        self.prompt_combo.setCurrentText(prompt_name)
                        self.prompt_edit.setText("Enter your prompt template here...")
                
                        # Highlight the text to make it easy to replace
                        self.prompt_edit.selectAll()
                        self.prompt_edit.setFocus()
                
                def save_prompt_changes(self):
                    """Save changes to the current prompt."""
                    current_name = self.prompt_combo.currentText()
                    if not current_name:
                        return
                
                    # Update the prompt in settings
                    prompts = self.settings.get("prompts", [])
                    prompt_text = self.prompt_edit.toPlainText()
                
                    # Find if this prompt already exists
                    found = False
                    for i, prompt in enumerate(prompts):
                        if prompt["name"] == current_name:
                            prompts[i]["prompt"] = prompt_text
                            found = True
                            break
                
                    # If it's a new prompt, add it
                    if not found:
                        prompts.append({"name": current_name, "prompt": prompt_text})
                
                    self.settings["prompts"] = prompts
                
                    # Show a save animation
                    save_animation = QLabel("✓ Saved")
                    save_animation.setStyleSheet("""
                        QLabel {
                            background-color: #34C759;
                            color: white;
                            border-radius: 4px;
                            padding: 5px 10px;
                            font-weight: bold;
                        }
                    """)
                
                    # Add the animation to prompt layout
                    prompt_card = self.prompt_edit.parent()
                    if prompt_card:
                        prompt_layout = prompt_card.layout()
                        prompt_layout.addWidget(save_animation)
                
                        # Remove after a delay
                        QTimer.singleShot(2000, lambda: save_animation.deleteLater())
            def delete_prompt(self):
                    """Delete the current prompt."""
                    current_name = self.prompt_combo.currentText()
                    if not current_name:
                        return
                
                    # Confirm deletion
                    confirm = QMessageBox.question(
                        self, "Confirm Deletion", 
                        f"Are you sure you want to delete the prompt '{current_name}'?",
                        QMessageBox.Yes | QMessageBox.No,
                        QMessageBox.No  # Default to No for safety
                    )
                
                    if confirm != QMessageBox.Yes:
                        return
                
                    # Remove from settings
                    prompts = self.settings.get("prompts", [])
                    self.settings["prompts"] = [p for p in prompts if p["name"] != current_name]
                
                    # Remove from combo
                    index = self.prompt_combo.findText(current_name)
                    if index >= 0:
                        self.prompt_combo.removeItem(index)
                
                    # Clear text edit if no prompts left
                    if self.prompt_combo.count() == 0:
                        self.prompt_edit.clear()
                
                def update_prompt_edit(self, prompt_name):
                    """Update the prompt edit box with the selected prompt."""
                    if not prompt_name:
                        self.prompt_edit.clear()
                        return
                
                    # Find the prompt in settings
                    prompts = self.settings.get("prompts", [])
                    for prompt in prompts:
                        if prompt["name"] == prompt_name:
                            self.prompt_edit.setText(prompt["prompt"])
                            return
                
                    # If not found, clear the edit box
                    self.prompt_edit.clear()
                
                def load_settings(self):
                    """Load current settings into the dialog."""
                    # General settings
                    self.theme_combo.setCurrentText(self.settings.get("theme", "macOS"))
                    self.start_bg_check.setChecked(self.settings.get("start_in_background", False))
                    self.shortcut_edit.setText(self.settings.get("keyboard_shortcut", "ctrl+i"))
                    self.code_word_edit.setText(self.settings.get("code_word", DEFAULT_CODE_WORD))
                
                    # Visual effects
                    self.rounded_corners_check.setChecked(self.settings.get("use_rounded_corners", True))
                    self.animations_check.setChecked(self.settings.get("use_animations", True))
                    self.transparency_check.setChecked(self.settings.get("use_transparency", True))
                
                    # Font settings
                    self.font_family_combo.setCurrentText(self.settings.get("font_family", DEFAULT_FONT_FAMILY))
                    self.font_size_spin.setValue(self.settings.get("font_size", DEFAULT_FONT_SIZE))
                    self.line_spacing_spin.setValue(self.settings.get("line_spacing", DEFAULT_LINE_SPACING))
                    self.word_spacing_spin.setValue(int(self.settings.get("word_spacing", DEFAULT_WORD_SPACING) * 100))
                
                    # Update font preview
                    font_preview = self.findChild(QLabel, "font_preview")
                    if font_preview:
                        font = QFont(self.font_family_combo.currentText(), self.font_size_spin.value())
                        font_preview.setFont(font)
                
                    # Icon paths
                    self.app_icon_edit.setText(self.settings.get("app_icon_path", ""))
                    self.generate_icon_edit.setText(self.settings.get("generate_icon_path", ""))
                    self.settings_icon_edit.setText(self.settings.get("settings_icon_path", ""))
                
                    # Add previews for icons if paths exist
                    for field_name in ["app_icon_edit", "generate_icon_edit", "settings_icon_edit"]:
                        line_edit = getattr(self, field_name)
                        path = line_edit.text()
                        if path and os.path.exists(path):
                            self.browse_icon(line_edit)  # This will add the preview
                
                    # AI settings
                    self.model_combo.setCurrentText(self.settings.get("model", DEFAULT_MODEL))
                    self.line_interval_spin.setValue(self.settings.get("summary_line_interval", 10))
                    self.time_interval_spin.setValue(self.settings.get("summary_time_interval", 120))
                
                    # Load prompts
                    self.prompt_combo.clear()
                    prompts = self.settings.get("prompts", [])
                    for prompt in prompts:
                        self.prompt_combo.addItem(prompt["name"])
                
                    # Set active prompt
                    active_prompt = self.settings.get("active_prompt", "")
                    if active_prompt:
                        index = self.prompt_combo.findText(active_prompt)
                        if index >= 0:
                            self.prompt_combo.setCurrentIndex(index)
                            self.update_prompt_edit(active_prompt)
                
                    # Advanced settings
                    self.poll_interval_spin.setValue(self.settings.get("poll_interval", DEFAULT_POLL_INTERVAL))
                    self.anonymize_check.setChecked(self.settings.get("anonymize_transcript", False))
                    self.spellcheck_check.setChecked(self.settings.get("spellcheck_enabled", False))
                    self.timer_check.setChecked(self.settings.get("meeting_timer_enabled", True))
                    self.reminder_spin.setValue(self.settings.get("meeting_reminder_interval", 15))
            def reset_to_defaults(self):
                    """Reset all settings to defaults."""
                    confirm = QMessageBox.question(
                        self, "Confirm Reset", 
                        "Are you sure you want to reset all settings to default values?",
                        QMessageBox.Yes | QMessageBox.No
                    )
                
                    if confirm != QMessageBox.Yes:
                        return
                
                    # Reset to defaults but keep existing prompts
                    prompts = self.settings.get("prompts", [DEFAULT_SETTINGS["prompts"][0]])
                    self.settings = DEFAULT_SETTINGS.copy()
                    self.settings["prompts"] = prompts
                
                    # Reload settings into dialog
                    self.load_settings()
                
                    QMessageBox.information(self, "Reset Complete", "Settings have been reset to defaults.")
                
                def save_settings(self):
                    """Save settings from dialog to settings dict."""
                    # General settings
                    self.settings["theme"] = self.theme_combo.currentText()
                    self.settings["start_in_background"] = self.start_bg_check.isChecked()
                    self.settings["keyboard_shortcut"] = self.shortcut_edit.text()
                    self.settings["code_word"] = self.code_word_edit.text()
                
                    # Visual effects
                    self.settings["use_rounded_corners"] = self.rounded_corners_check.isChecked()
                    self.settings["use_animations"] = self.animations_check.isChecked()
                    self.settings["use_transparency"] = self.transparency_check.isChecked()
                
                    # Font settings
                    self.settings["font_family"] = self.font_family_combo.currentText()
                    self.settings["font_size"] = self.font_size_spin.value()
                    self.settings["line_spacing"] = self.line_spacing_spin.value()
                    self.settings["word_spacing"] = self.word_spacing_spin.value() / 100.0
                
                    # Icon paths
                    self.settings["app_icon_path"] = self.app_icon_edit.text()
                    self.settings["generate_icon_path"] = self.generate_icon_edit.text()
                    self.settings["settings_icon_path"] = self.settings_icon_edit.text()
                
                    # AI settings
                    self.settings["model"] = self.model_combo.currentText()
                    self.settings["summary_line_interval"] = self.line_interval_spin.value()
                    self.settings["summary_time_interval"] = self.time_interval_spin.value()
                
                    # Active prompt
                    current_prompt = self.prompt_combo.currentText()
                    if current_prompt:
                        self.settings["active_prompt"] = current_prompt
                
                        # Make sure current prompt changes are saved
                        self.save_prompt_changes()
                
                    # Advanced settings
                    self.settings["poll_interval"] = self.poll_interval_spin.value()
                    self.settings["anonymize_transcript"] = self.anonymize_check.isChecked()
                    self.settings["spellcheck_enabled"] = self.spellcheck_check.isChecked()
                    self.settings["meeting_timer_enabled"] = self.timer_check.isChecked()
                    self.settings["meeting_reminder_interval"] = self.reminder_spin.value()
                
                    # Accept dialog
                    self.accept()
            # --------------------------
            # Main Application Class
            # --------------------------
            class MeetingAssistant(QWidget):
                def __init__(self):
                    super().__init__()
                    self.setWindowTitle(f"{APP_NAME} v{VERSION}")
                    self.setGeometry(100, 100, 900, 700)
                
                    # Load icon from settings or use default
                    self.app_icon = load_icon_from_path_or_default(
                        settings.get("app_icon_path", ""), 
                        DEFAULT_APP_ICON
                    )
                    self.setWindowIcon(self.app_icon)
                
                    # Initialize state variables
                    self.running = False
                    self.browser = None
                    self.page = None
                    self.buffer = []
                    self.seen_lines = set()
                    self.summaries = []
                    self.suggested_responses = []
                    self.insights = []
                    self.known_names = []
                    self.meeting_start_time = None
                    self.last_reminder_time = None
                    self.code_word_detected = False
                    self.meeting_title = "Unknown Meeting"
                    self.key_phrases = []
                
                    # Create child windows
                    self.wordcloud_window = WordCloudWindow(self)
                    self.settings_dialog = None  # Create on demand
                
                    # Create timers
                    self.timer = QTimer()  # Main timer
                    self.elapsed_seconds = 0
                    self.progress_timer = QTimer()  # For AI progress indication
                    self.meeting_timer = QTimer()  # For meeting duration tracking
                
                    # Create signal bridge for thread-safe UI updates
                    self.signals = SignalBridge()
                
                    # Build UI first
                    self.build_ui()
                
                    # Connect signals for UI updates
                    self.signals.append_transcript.connect(self.append_transcript)
                    self.signals.append_suggested_response.connect(self.append_suggested_response)
                    self.signals.append_insights.connect(self.append_insights)
                    self.signals.update_status.connect(self.update_status)
                    self.signals.update_progress.connect(self.update_progress)
                    self.signals.highlight_transcript.connect(self.highlight_transcript)
                    self.signals.update_wordcloud.connect(self.update_wordcloud)
                    self.signals.update_meeting_timer.connect(self.update_meeting_timer)
                    self.signals.detected_key_phrases.connect(self.update_key_phrases)
                    self.signals.detected_meeting_title.connect(self.update_meeting_title)
                
                    # Apply theme and setup UI components
                    self.apply_theme()
                    self.setup_tray()
                    self.setup_hotkey()
                
                    # Start background if configured
                    if settings.get("start_in_background", False):
                        self.hide()
            def build_ui(self):
                    # Top bar with URL input, generate button and settings
                    self.url_input = QLineEdit()
                    self.url_input.setPlaceholderText("Enter Otter.ai meeting URL")
                    self.url_input.setMinimumHeight(32)
                
                    # Load icons
                    self.generate_icon = load_icon_from_path_or_default(
                        settings.get("generate_icon_path", ""), 
                        DEFAULT_GENERATE_ICON
                    )
                    self.settings_icon = load_icon_from_path_or_default(
                        settings.get("settings_icon_path", ""), 
                        DEFAULT_SETTINGS_ICON
                    )
                
                    # Create icon buttons
                    self.generate_btn = create_icon_button(
                        self.generate_icon, 
                        "Generate AI Analysis Now (Shortcut: " + settings.get("keyboard_shortcut", "Ctrl+I") + ")"
                    )
                    self.settings_btn = create_icon_button(
                        self.settings_icon, 
                        "Settings"
                    )
                
                    # Connect buttons
                    self.generate_btn.clicked.connect(self.on_demand_analysis)
                    self.settings_btn.clicked.connect(self.open_settings)
                
                    # Top bar layout with macOS-style search field
                    top_bar = QHBoxLayout()
                    top_bar.addWidget(QLabel("Meeting URL:"))
                
                    # URL input takes 2/3 of the space and has a search-like appearance
                    url_container = QWidget()
                    url_container.setObjectName("url_container")
                    url_container.setStyleSheet("""
                        #url_container {
                            background-color: rgba(255, 255, 255, 0.1);
                            border-radius: 8px;
                            padding: 0px;
                        }
                    """)
                    url_layout = QHBoxLayout(url_container)
                    url_layout.setContentsMargins(0, 0, 0, 0)
                    url_layout.addWidget(self.url_input)
                
                    top_bar.addWidget(url_container, 2)
                
                    # Generate button and settings button
                    button_spacer = QSpacerItem(10, 10, QSizePolicy.Fixed, QSizePolicy.Minimum)
                    top_bar.addItem(button_spacer)
                    top_bar.addWidget(self.generate_btn)
                    top_bar.addWidget(self.settings_btn)
                
                    # Status and progress bar in a card-like container
                    status_container = QFrame()
                    status_container.setFrameShape(QFrame.StyledPanel)
                    status_container.setStyleSheet("""
                        QFrame {
                            border: 1px solid rgba(0, 0, 0, 0.1);
                            border-radius: 8px;
                            background-color: rgba(255, 255, 255, 0.05);
                            padding: 6px;
                        }
                    """)
                
                    status_layout = QHBoxLayout(status_container)
                    status_layout.setContentsMargins(12, 8, 12, 8)
                
                    self.status = QLabel("Status: Ready")
                    self.status.setStyleSheet("font-weight: 500;")
                
                    self.progress_bar = QProgressBar()
                    self.progress_bar.setRange(0, 100)
                    self.progress_bar.setValue(0)
                    self.progress_bar.setVisible(False)
                    self.progress_bar.setMaximumWidth(150)
                    self.progress_bar.setMinimumWidth(150)
                
                    self.timer_label = QLabel("Meeting time: 00:00:00")
                    self.timer_label.setStyleSheet("font-weight: 500;")
                
                    status_layout.addWidget(self.status)
                    status_layout.addStretch()
                    status_layout.addWidget(self.progress_bar)
                    status_layout.addWidget(self.timer_label)
            # Action buttons in a modern toolbar style
                    btn_container = QFrame()
                    btn_container.setStyleSheet("""
                        QFrame {
                            background-color: rgba(0, 0, 0, 0.05);
                            border-radius: 8px;
                            padding: 4px;
                        }
                    """)
                
                    btn_layout = QHBoxLayout(btn_container)
                    btn_layout.setSpacing(12)
                    btn_layout.setContentsMargins(12, 8, 12, 8)
                
                    self.start_btn = create_styled_button("Start")
                    self.start_btn.setIcon(QIcon.fromTheme("media-playback-start"))
                    self.start_btn.setStyleSheet("""
                        QPushButton {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #4CAF50, stop:1 #43A047);
                            color: white;
                            border: none;
                            border-radius: 6px;
                            padding: 8px 20px;
                            font-weight: 600;
                        }
                        QPushButton:hover {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #66BB6A, stop:1 #4CAF50);
                        }
                        QPushButton:pressed {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #43A047, stop:1 #388E3C);
                            padding-top: 9px;
                            padding-bottom: 7px;
                        }
                    """)
                
                    self.stop_btn = create_styled_button("Stop")
                    self.stop_btn.setIcon(QIcon.fromTheme("media-playback-stop"))
                    self.stop_btn.setStyleSheet("""
                        QPushButton {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #F44336, stop:1 #E53935);
                            color: white;
                            border: none;
                            border-radius: 6px;
                            padding: 8px 20px;
                            font-weight: 600;
                        }
                        QPushButton:hover {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #EF5350, stop:1 #F44336);
                        }
                        QPushButton:pressed {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #E53935, stop:1 #D32F2F);
                            padding-top: 9px;
                            padding-bottom: 7px;
                        }
                    """)
                
                    self.save_btn = create_styled_button("Save")
                    self.save_btn.setIcon(QIcon.fromTheme("document-save"))
                    self.save_btn.setStyleSheet("""
                        QPushButton {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2196F3, stop:1 #1E88E5);
                            color: white;
                            border: none;
                            border-radius: 6px;
                            padding: 8px 20px;
                            font-weight: 600;
                        }
                        QPushButton:hover {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #42A5F5, stop:1 #2196F3);
                        }
                        QPushButton:pressed {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #1E88E5, stop:1 #1976D2);
                            padding-top: 9px;
                            padding-bottom: 7px;
                        }
                    """)
                
                    self.cloud_btn = create_styled_button("Word Cloud")
                    self.cloud_btn.setIcon(QIcon.fromTheme("view-grid"))
                    self.cloud_btn.setStyleSheet("""
                        QPushButton {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #9C27B0, stop:1 #8E24AA);
                            color: white;
                            border: none;
                            border-radius: 6px;
                            padding: 8px 20px;
                            font-weight: 600;
                        }
                        QPushButton:hover {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #AB47BC, stop:1 #9C27B0);
                        }
                        QPushButton:pressed {
                            background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #8E24AA, stop:1 #7B1FA2);
                            padding-top: 9px;
                            padding-bottom: 7px;
                        }
                    """)
                
                    self.start_btn.clicked.connect(self.start)
                    self.stop_btn.clicked.connect(self.stop)
                    self.save_btn.clicked.connect(self.save)
                    self.cloud_btn.clicked.connect(self.show_wordcloud)
                
                    btn_layout.addWidget(self.start_btn)
                    btn_layout.addWidget(self.stop_btn)
                    btn_layout.addWidget(self.save_btn)
                    btn_layout.addWidget(self.cloud_btn)
                
                    # Model and prompt selector in a elegant card-like container
                    model_container = QFrame()
                    model_container.setFrameShape(QFrame.StyledPanel)
                    model_container.setStyleSheet("""
                        QFrame {
                            border: 1px solid rgba(0, 0, 0, 0.1);
                            border-radius: 8px;
                            background-color: rgba(255, 255, 255, 0.05);
                            padding: 4px;
                        }
                        QLabel {
                            font-weight: 500;
                        }
                    """)
                
                    model_layout = QHBoxLayout(model_container)
                    model_layout.setContentsMargins(12, 8, 12, 8)
                
                    model_label = QLabel("AI Model:")
                    model_label.setStyleSheet("font-weight: 500;")
                
                    self.model_selector = QComboBox()
                    self.model_selector.addItems(AVAILABLE_MODELS)
                    self.model_selector.setCurrentText(settings.get("model", DEFAULT_MODEL))
                    self.model_selector.setMinimumWidth(150)
                    self.model_selector.setStyleSheet("""
                        QComboBox {
                            padding: 6px 12px;
                            border-radius: 6px;
                            font-weight: 500;
                        }
                    """)
                
                    model_layout.addWidget(model_label)
                    model_layout.addWidget(self.model_selector)
                    model_layout.addSpacing(20)
                
                    prompt_label = QLabel("Prompt:")
                    prompt_label.setStyleSheet("font-weight: 500;")
                
                    self.prompt_selector = QComboBox()
                    self.prompt_selector.addItems([p["name"] for p in settings.get("prompts", [])])
                    self.prompt_selector.setCurrentText(settings.get("active_prompt", ""))
                    self.prompt_selector.currentTextChanged.connect(self.change_prompt)
                    self.prompt_selector.setMinimumWidth(200)
                    self.prompt_selector.setStyleSheet("""
                        QComboBox {
                            padding: 6px 12px;
                            border-radius: 6px;
                            font-weight: 500;
                        }
                    """)
                
                    model_layout.addWidget(prompt_label)
                    model_layout.addWidget(self.prompt_selector)
                
                    # Main content area
                    content_layout = QVBoxLayout()
                
                    # Meeting topic with a stylish header
                    meeting_header = QFrame()
                    meeting_header.setFrameShape(QFrame.StyledPanel)
                    meeting_header.setStyleSheet("""
                        QFrame {
                            border: 1px solid rgba(0, 0, 0, 0.1);
                            border-radius: 8px;
                            background-color: rgba(255, 255, 255, 0.05);
                            padding: 4px;
                        }
                    """)
                
                    meeting_layout = QHBoxLayout(meeting_header)
                    meeting_layout.setContentsMargins(12, 8, 12, 8)
                
                    topic_label = QLabel("<b>Meeting Topic:</b>")
                    topic_label.setStyleSheet("font-weight: 600;")
                
                    self.topic_value = QLabel("<span id='meeting-title'>Not detected yet</span>")
                    self.topic_value.setStyleSheet("font-style: italic;")
                
                    meeting_layout.addWidget(topic_label)
                    meeting_layout.addWidget(self.topic_value)
                    meeting_layout.addStretch()
                
                    # Groq output panel (at the top) with card-like styling
                    self.groq_panel = QSplitter(Qt.Horizontal)
                    self.groq_panel.setStyleSheet("""
                        QSplitter::handle {
                            background-color: rgba(0, 0, 0, 0.2);
                            margin: 2px;
                        }
                        QSplitter::handle:hover {
                            background-color: rgba(0, 0, 0, 0.3);
                        }
                    """)
                
                    # Left side: Suggested responses with card styling
                    suggested_frame = QFrame()
                    suggested_frame.setFrameShape(QFrame.StyledPanel)
                    suggested_frame.setStyleSheet("""
                        QFrame {
                            border: 1px solid rgba(0, 0, 0, 0.1);
                            border-radius: 8px;
                            background-color: rgba(255, 255, 255, 0.03);
                        }
                    """)
                
                    suggested_layout = QVBoxLayout(suggested_frame)
                    suggested_layout.setContentsMargins(0, 0, 0, 0)
                
                    suggested_header = QFrame()
                    suggested_header.setStyleSheet("""
                        QFrame {
                            background-color: rgba(0, 0, 0, 0.1);
                            border-top-left-radius: 8px;
                            border-top-right-radius: 8px;
                            padding: 8px;
                            margin: 0px;
                        }
                    """)
                    suggested_header_layout = QHBoxLayout(suggested_header)
                    suggested_header_layout.setContentsMargins(10, 5, 10, 5)
                
                    suggested_title = QLabel("<b>Suggested Responses</b>")
                    suggested_title.setStyleSheet("font-weight: 600; font-size: 14px;")
                    suggested_header_layout.addWidget(suggested_title)
                
                    self.suggested_response_box = CustomTextEdit()
                    self.suggested_response_box.setReadOnly(True)
                    self.suggested_response_box.document().setDefaultFont(
                        QFont(settings.get("font_family", "SF Pro Text"), 
                            settings.get("font_size", DEFAULT_FONT_SIZE))
                    )
                    self.suggested_response_box.setStyleSheet("""
                        QTextEdit {
                            border: none;
                            border-bottom-left-radius: 8px;
                            border-bottom-right-radius: 8px;
                            background-color: transparent;
                            padding: 12px;
                        }
                    """)
                
                    suggested_layout.addWidget(suggested_header)
                    suggested_layout.addWidget(self.suggested_response_box)
            # Right side: Insights with card styling
            insights_frame = QFrame()
            insights_frame.setFrameShape(QFrame.StyledPanel)
            insights_frame.setStyleSheet("""
                QFrame {
                    border: 1px solid rgba(0, 0, 0, 0.1);
                    border-radius: 8px;
                    background-color: rgba(255, 255, 255, 0.03);
                }
            """)
                
            insights_layout = QVBoxLayout(insights_frame)
            insights_layout.setContentsMargins(0, 0, 0, 0)
                
            insights_header = QFrame()
            insights_header.setStyleSheet("""
                QFrame {
                    background-color: rgba(0, 0, 0, 0.1);
                    border-top-left-radius: 8px;
                    border-top-right-radius: 8px;
                    padding: 8px;
                    margin: 0px;
                }
            """)
            insights_header_layout = QHBoxLayout(insights_header)
            insights_header_layout.setContentsMargins(10, 5, 10, 5)
                
            insights_title = QLabel("<b>Key Insights & Analysis</b>")
            insights_title.setStyleSheet("font-weight: 600; font-size: 14px;")
            insights_header_layout.addWidget(insights_title)
                
            self.insights_box = CustomTextEdit()
            self.insights_box.setReadOnly(True)
            self.insights_box.document().setDefaultFont(
                QFont(settings.get("font_family", "SF Pro Text"), 
                    settings.get("font_size", DEFAULT_FONT_SIZE))
            )
            self.insights_box.setStyleSheet("""
                QTextEdit {
                    border: none;
                    border-bottom-left-radius: 8px;
                    border-bottom-right-radius: 8px;
                    background-color: transparent;
                    padding: 12px;
                }
            """)
                
            insights_layout.addWidget(insights_header)
            insights_layout.addWidget(self.insights_box)
                
            # Add to splitter
            self.groq_panel.addWidget(suggested_frame)
            self.groq_panel.addWidget(insights_frame)
            self.groq_panel.setSizes([int(self.width() * 0.5), int(self.width() * 0.5)])
                
            # Transcript panel (at the bottom) with card styling
            transcript_frame = QFrame()
            transcript_frame.setFrameShape(QFrame.StyledPanel)
            transcript_frame.setStyleSheet("""
                QFrame {
                    border: 1px solid rgba(0, 0, 0, 0.1);
                    border-radius: 8px;
                    background-color: rgba(255, 255, 255, 0.03);
                }
            """)
                
            transcript_layout = QVBoxLayout(transcript_frame)
            transcript_layout.setContentsMargins(0, 0, 0, 0)
                
            transcript_header = QFrame()
            transcript_header.setStyleSheet("""
                QFrame {
                    background-color: rgba(0, 0, 0, 0.1);
                    border-top-left-radius: 8px;
                    border-top-right-radius: 8px;
                    padding: 8px;
                    margin: 0px;
                }
            """)
                
            transcript_header_layout = QHBoxLayout(transcript_header)
            transcript_header_layout.setContentsMargins(10, 5, 10, 5)
                
            transcript_title = QLabel("<b>Live Transcript</b>")
            transcript_title.setStyleSheet("font-weight: 600; font-size: 14px;")
            transcript_header_layout.addWidget(transcript_title)
                
            # Search box with modern styling
            search_container = QFrame()
            search_container.setStyleSheet("""
                QFrame {
                    background-color: rgba(255, 255, 255, 0.1);
                    border-radius: 6px;
                    padding: 0px;
                }
            """)
                
            search_layout = QHBoxLayout(search_container)
            search_layout.setContentsMargins(8, 0, 8, 0)
            search_layout.setSpacing(5)
                
            self.search_box = QLineEdit()
            self.search_box.setPlaceholderText("Search transcript...")
            self.search_box.setMaximumWidth(250)
            self.search_box.setMinimumHeight(28)
            self.search_box.textChanged.connect(self.search_transcript)
            self.search_box.setStyleSheet("""
                QLineEdit {
                    border: none;
                    background: transparent;
                    padding: 4px;
                }
            """)
                
            search_icon = QLabel()
            search_icon.setPixmap(QIcon.fromTheme("edit-find").pixmap(16, 16))
                
            search_layout.addWidget(search_icon)
            search_layout.addWidget(self.search_box)
                
            self.search_btn = QPushButton("Find")
            self.search_btn.setMaximumWidth(60)
            self.search_btn.clicked.connect(lambda: self.search_transcript(self.search_box.text()))
            self.search_btn.setStyleSheet("""
                QPushButton {
                    background-color: rgba(0, 0, 0, 0.2);
                    border: none;
                    border-radius: 4px;
                    padding: 4px 8px;
                    color: white;
                    font-weight: 500;
                }
                QPushButton:hover {
                    background-color: rgba(0, 0, 0, 0.3);
                }
                QPushButton:pressed {
                    background-color: rgba(0, 0, 0, 0.4);
                }
            """)
                
            search_layout.addWidget(self.search_btn)
                
            transcript_header_layout.addWidget(search_container)
            # Transcript box with modern styling
            self.transcript_box = CustomTextEdit()
            self.transcript_box.setReadOnly(True)
            self.transcript_box.document().setDefaultFont(
                QFont(settings.get("font_family", "SF Pro Text"), 
                    settings.get("font_size", DEFAULT_FONT_SIZE))
            )
            self.transcript_box.setStyleSheet("""
                QTextEdit {
                    border: none;
                    border-bottom-left-radius: 8px;
                    border-bottom-right-radius: 8px;
                    background-color: transparent;
                    padding: 12px;
                }
            """)
                
            transcript_layout.addWidget(transcript_header)
            transcript_layout.addWidget(self.transcript_box)
                
            # Add panels to content layout
            content_layout.addWidget(meeting_header)
            content_layout.addWidget(self.groq_panel, 1)  # 1 stretch factor
            content_layout.addWidget(transcript_frame, 1)  # 1 stretch factor
                
            # Main layout with proper margins and spacing
            main_layout = QVBoxLayout()
            main_layout.setSpacing(12)
            main_layout.setContentsMargins(16, 16, 16, 16)
                
            main_layout.addLayout(top_bar)
            main_layout.addWidget(status_container)
            main_layout.addWidget(btn_container)
            main_layout.addWidget(model_container)
            main_layout.addLayout(content_layout)
                
            self.setLayout(main_layout)
                
            def apply_theme(self):
                """Apply the selected theme to the application."""
                theme_name = settings.get("theme", "macOS")
                if theme_name not in THEMES:
                    theme_name = "macOS"  # Default to macOS theme if not found
                
                theme = THEMES[theme_name]
                
                # Create application palette
                palette = QPalette()
                
                # Set colors
                palette.setColor(QPalette.Window, QColor(theme["window"]))
                palette.setColor(QPalette.WindowText, QColor(theme["windowText"]))
                palette.setColor(QPalette.Base, QColor(theme["base"]))
                palette.setColor(QPalette.AlternateBase, QColor(theme["alternateBase"]))
                palette.setColor(QPalette.Text, QColor(theme["text"]))
                palette.setColor(QPalette.Button, QColor(theme["button"]))
                palette.setColor(QPalette.ButtonText, QColor(theme["buttonText"]))
                palette.setColor(QPalette.Highlight, QColor(theme["highlight"]))
                palette.setColor(QPalette.HighlightedText, QColor(theme["highlightedText"]))
                palette.setColor(QPalette.Link, QColor(theme["link"]))
                palette.setColor(QPalette.BrightText, QColor(theme["brightText"]))
                
                # Apply palette to application
                self.setPalette(palette)
                
                # Theme-specific UI adjustments
                is_dark = theme_name in ["dark", "navy", "dracula", "macOSDark"]
                
                # Apply theme to specific widgets
                search_color = "rgba(255, 255, 255, 0.1)" if is_dark else "rgba(0, 0, 0, 0.05)"
                self.search_box.parent().setStyleSheet(f"""
                    QFrame {{
                        background-color: {search_color};
                        border-radius: 6px;
                        padding: 0px;
                    }}
                """)
                
                # Apply font settings
                font_family = settings.get("font_family", DEFAULT_FONT_FAMILY)
                font_size = settings.get("font_size", DEFAULT_FONT_SIZE)
                line_spacing = settings.get("line_spacing", DEFAULT_LINE_SPACING) / 100.0
                word_spacing = settings.get("word_spacing", DEFAULT_WORD_SPACING)
                
                # Apply font to text boxes
                font = QFont(font_family, font_size)
                for text_edit in [self.transcript_box, self.suggested_response_box, self.insights_box]:
                    text_edit.document().setDefaultFont(font)
                
                    # Set line and word spacing
                    doc_format = text_edit.document().defaultTextOption()
                    doc_format.setTextDirection(Qt.LeftToRight)
                    doc_format.setWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)
                    text_edit.document().setDefaultTextOption(doc_format)
                
                    block_format = QTextBlockFormat()
                    block_format.setLineHeight(int(100 * line_spacing), QTextBlockFormat.ProportionalHeight)
                    cursor = QTextCursor(text_edit.document())
                    cursor.movePosition(QTextCursor.Start)
                    cursor.movePosition(QTextCursor.End, QTextCursor.KeepAnchor)
                    cursor.setBlockFormat(block_format)
                
                # Apply rounded corners if enabled
                if settings.get("use_rounded_corners", True):
                    self.setStyleSheet(f"""
                        QWidget {{
                            border-radius: 8px;
                            font-family: '{font_family}';
                        }}
                        QFrame {{
                            border-radius: 6px;
                        }}
                        QPushButton {{
                            border-radius: 6px;
                        }}
                        QLineEdit {{
                            border-radius: 4px;
                        }}
                        QComboBox {{
                            border-radius: 4px;
                        }}
                    """)
                
            def setup_tray(self):
                """Set up system tray icon and menu."""
                self.tray_icon = QSystemTrayIcon(self)
                self.tray_icon.setIcon(self.app_icon)
                
                # Create tray menu
                tray_menu = QMenu()
                
                show_action = QAction("Show", self)
                show_action.triggered.connect(self.show)
                
                hide_action = QAction("Hide", self)
                hide_action.triggered.connect(self.hide)
                
                quit_action = QAction("Quit", self)
                quit_action.triggered.connect(self.close)
                
                tray_menu.addAction(show_action)
                tray_menu.addAction(hide_action)
                tray_menu.addSeparator()
                tray_menu.addAction(quit_action)
                
                self.tray_icon.setContextMenu(tray_menu)
                self.tray_icon.activated.connect(self.tray_activated)
                
                self.tray_icon.show()
            def setup_hotkey(self):
                """Set up global hotkey for triggering analysis."""
                try:
                    keyboard.add_hotkey(
                        settings.get("keyboard_shortcut", "ctrl+i"), 
                        self.on_demand_analysis, 
                        suppress=False
                    )
                except Exception as e:
                    print(f"Error setting up hotkey: {e}")
                
            def tray_activated(self, reason):
                """Handle tray icon activation."""
                if reason == QSystemTrayIcon.DoubleClick:
                    if self.isVisible():
                        self.hide()
                    else:
                        self.show()
                        self.activateWindow()
                
            def open_settings(self):
                """Open settings dialog."""
                if not self.settings_dialog:
                    self.settings_dialog = SettingsDialog(self, settings)
                
                if self.settings_dialog.exec_():
                    # Save settings
                    global settings
                    settings = self.settings_dialog.settings
                
                    # Apply new settings
                    self.apply_theme()
                
                    # Update prompts
                    self.prompt_selector.clear()
                    self.prompt_selector.addItems([p["name"] for p in settings.get("prompts", [])])
                    self.prompt_selector.setCurrentText(settings.get("active_prompt", ""))
                
                    # Update model
                    self.model_selector.setCurrentText(settings.get("model", DEFAULT_MODEL))
                
                    # Update keyboard shortcut
                    try:
                        keyboard.remove_all_hotkeys()
                        keyboard.add_hotkey(
                            settings.get("keyboard_shortcut", "ctrl+i"), 
                            self.on_demand_analysis, 
                            suppress=False
                        )
                    except Exception as e:
                        print(f"Error updating hotkey: {e}")
                
                    # Save settings to file
                    self.save_settings_to_file()
                
                    # Show success message
                    self.log("Settings updated successfully")
                
            def save_settings_to_file(self):
                """Save settings to JSON file."""
                try:
                    with open("settings.json", "w") as f:
                        json.dump(settings, f, indent=4)
                except Exception as e:
                    print(f"Error saving settings: {e}")
                
            def change_prompt(self, prompt_name):
                """Change the active prompt."""
                if not prompt_name:
                    return
                
                settings["active_prompt"] = prompt_name
                self.save_settings_to_file()
                self.log(f"Prompt changed to: {prompt_name}")
                
            @asyncSlot()
            async def start(self):
                """Start the transcript monitoring."""
                if self.running:
                    self.log("Already running")
                    return
                
                url = self.url_input.text().strip()
                if not url:
                    QMessageBox.warning(self, "Error", "Please enter a valid Otter.ai URL")
                    return
                
                if not url.startswith("https://otter.ai/"):
                    confirm = QMessageBox.question(
                        self, "URL Verification", 
                        "The URL doesn't appear to be from Otter.ai. Continue anyway?",
                        QMessageBox.Yes | QMessageBox.No, 
                        QMessageBox.No
                    )
                    if confirm == QMessageBox.No:
                        return
                
                # Reset state
                self.running = True
                self.buffer = []
                self.seen_lines = set()
                self.summaries = []
                self.suggested_responses = []
                self.insights = []
                self.key_phrases = []
                
                # Clear text boxes
                self.transcript_box.clear()
                self.suggested_response_box.clear() 
                self.insights_box.clear()
                
                # Update UI
                self.start_btn.setEnabled(False)
                self.stop_btn.setEnabled(True)
                self.log("Starting transcript monitoring...")
                self.progress_bar.setVisible(True)
                self.update_progress(10)
                
                # Start meeting timer if enabled
                if settings.get("meeting_timer_enabled", True):
                    self.meeting_start_time = datetime.datetime.now()
                    self.last_reminder_time = self.meeting_start_time
                    self.meeting_timer.timeout.connect(self.update_meeting_time)
                    self.meeting_timer.start(1000)  # Update every second
                
                # Start in a separate thread
                try:
                    await self.scrape_otter(url)
                except Exception as e:
                    self.log(f"Error: {e}")
                    self.stop()
            @asyncSlot()
            async def stop(self):
                """Stop transcript monitoring."""
                if not self.running:
                    return
                
                self.running = False
                self.log("Stopping transcript monitoring...")
                
                # Close browser
                if self.browser:
                    try:
                        await self.browser.close()
                        self.browser = None
                        self.page = None
                    except Exception as e:
                        print(f"Error closing browser: {e}")
                
                # Update UI
                self.start_btn.setEnabled(True)
                self.stop_btn.setEnabled(False)
                self.progress_bar.setVisible(False)
                
                # Stop timers
                self.timer.stop()
                self.meeting_timer.stop()
                
                self.log("Transcript monitoring stopped")
                
            def save(self):
                """Save transcript to file."""
                if not self.transcript_box.toPlainText():
                    QMessageBox.warning(self, "Error", "No transcript to save")
                    return
                
                # Choose file format
                formats = {
                    "Text File (*.txt)": self.save_as_txt,
                    "Word Document (*.docx)": self.save_as_docx,
                    "PDF Document (*.pdf)": self.save_as_pdf,
                    "SRT Subtitle (*.srt)": self.save_as_srt,
                    "HTML Document (*.html)": self.save_as_html,
                    "Markdown Document (*.md)": self.save_as_markdown
                }
                
                format_str, _ = QInputDialog.getItem(
                    self, "Select Format", "Choose file format:", 
                    list(formats.keys()), 0, False
                )
                
                if not format_str:
                    return
                
                # Get file path
                extension = format_str.split("*")[1].rstrip(")")
                file_path, _ = QFileDialog.getSaveFileName(
                    self, "Save Transcript", 
                    f"transcript_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}{extension}",
                    format_str
                )
                
                if not file_path:
                    return
                
                # Save in selected format
                try:
                    save_function = formats[format_str]
                    save_function(file_path)
                
                    self.log(f"Transcript saved to {file_path}")
                
                    # Ask if they want to open the file
                    confirm = QMessageBox.question(
                        self, "Open File", 
                        "Transcript saved. Would you like to open it now?",
                        QMessageBox.Yes | QMessageBox.No
                    )
                
                    if confirm == QMessageBox.Yes:
                        try:
                            if sys.platform == 'win32':
                                os.startfile(file_path)
                            elif sys.platform == 'darwin':  # macOS
                                os.system(f"open {file_path}")
                            else:  # Linux
                                os.system(f"xdg-open {file_path}")
                        except Exception as e:
                            print(f"Error opening file: {e}")
                
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to save transcript: {e}")
                
            def save_as_txt(self, file_path):
                """Save transcript as plain text."""
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(self.transcript_box.toPlainText())
                
            def save_as_docx(self, file_path):
                """Save transcript as Word document."""
                doc = Document()
                
                # Add title
                doc.add_heading("Meeting Transcript", 0)
                
                # Add metadata
                doc.add_paragraph(f"Generated by {APP_NAME} v{VERSION}")
                doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                if self.meeting_start_time:
                    duration = datetime.datetime.now() - self.meeting_start_time
                    hours, remainder = divmod(duration.seconds, 3600)
                    minutes, seconds = divmod(remainder, 60)
                    doc.add_paragraph(f"Meeting Duration: {hours:02}:{minutes:02}:{seconds:02}")
                
                # Add meeting title if detected
                if self.meeting_title and self.meeting_title != "Unknown Meeting":
                    doc.add_paragraph(f"Meeting Title: {self.meeting_title}")
                
                # Add transcript
                doc.add_heading("Transcript", 1)
                for line in self.transcript_box.toPlainText().split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                
                # Add AI insights if available
                if self.insights_box.toPlainText():
                    doc.add_heading("AI Analysis & Insights", 1)
                    for line in self.insights_box.toPlainText().split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)
                
                # Save document
                doc.save(file_path)
                
            def save_as_pdf(self, file_path):
                """Save transcript as PDF document."""
                pdf = FPDF()
                pdf.add_page()
                
                # Add title
                pdf.set_font("Arial", 'B', 16)
                pdf.cell(0, 10, "Meeting Transcript", 0, 1, 'C')
                pdf.ln(5)
                
                # Add metadata
                pdf.set_font("Arial", '', 10)
                pdf.cell(0, 5, f"Generated by {APP_NAME} v{VERSION}", 0, 1)
                pdf.cell(0, 5, f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1)
                if self.meeting_start_time:
                    duration = datetime.datetime.now() - self.meeting_start_time
                    hours, remainder = divmod(duration.seconds, 3600)
                    minutes, seconds = divmod(remainder, 60)
                    pdf.cell(0, 5, f"Meeting Duration: {hours:02}:{minutes:02}:{seconds:02}", 0, 1)
                
                # Add meeting title if detected
                if self.meeting_title and self.meeting_title != "Unknown Meeting":
                    pdf.cell(0, 5, f"Meeting Title: {self.meeting_title}", 0, 1)
                
                pdf.ln(5)
            # Add transcript
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, "Transcript", 0, 1)
            pdf.set_font("Arial", '', 10)
                
            # Add transcript text
            transcript_text = self.transcript_box.toPlainText()
            pdf.multi_cell(0, 5, transcript_text)
                
            # Add AI insights if available
            if self.insights_box.toPlainText():
                pdf.add_page()
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, "AI Analysis & Insights", 0, 1)
                pdf.set_font("Arial", '', 10)
                
                insights_text = self.insights_box.toPlainText()
                pdf.multi_cell(0, 5, insights_text)
                
            # Save PDF
            pdf.output(file_path)
                
            def save_as_srt(self, file_path):
                """Save transcript as SRT subtitle format."""
                lines = self.transcript_box.toPlainText().strip().split('\n')
                srt = ""
                
                subtitle_index = 1
                last_timestamp = None
                current_text = ""
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                
                    # Check if line starts with timestamp
                    timestamp_match = re.match(r'\[(\d+:\d+:\d+)\]', line)
                    if timestamp_match:
                        # If we had previous text, add it as a subtitle
                        if last_timestamp and current_text:
                            start_time = last_timestamp
                            # Calculate end time (5 seconds after start)
                            h, m, s = map(int, start_time.split(':'))
                            start_seconds = h * 3600 + m * 60 + s
                            end_seconds = start_seconds + 5
                
                            end_h, end_remainder = divmod(end_seconds, 3600)
                            end_m, end_s = divmod(end_remainder, 60)
                            end_time = f"{end_h:02d}:{end_m:02d}:{end_s:02d}"
                
                            # Format for SRT
                            srt_start = f"{start_time},000"
                            srt_end = f"{end_time},000"
                
                            srt += f"{subtitle_index}\n{srt_start} --> {srt_end}\n{current_text}\n\n"
                            subtitle_index += 1
                            current_text = ""
                
                        # Get timestamp from current line
                        last_timestamp = timestamp_match.group(1)
                        # Remove timestamp from line
                        line = line[line.find(']')+1:].strip()
                
                    # Add line to current text
                    if current_text:
                        current_text += "\n"
                    current_text += line
                
                # Add final subtitle
                if last_timestamp and current_text:
                    start_time = last_timestamp
                    # Calculate end time (5 seconds after start)
                    h, m, s = map(int, start_time.split(':'))
                    start_seconds = h * 3600 + m * 60 + s
                    end_seconds = start_seconds + 5
                
                    end_h, end_remainder = divmod(end_seconds, 3600)
                    end_m, end_s = divmod(end_remainder, 60)
                    end_time = f"{end_h:02d}:{end_m:02d}:{end_s:02d}"
                
                    # Format for SRT
                    srt_start = f"{start_time},000"
                    srt_end = f"{end_time},000"
                
                    srt += f"{subtitle_index}\n{srt_start} --> {srt_end}\n{current_text}\n\n"
                
                # Write to file
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(srt)
                
            def save_as_html(self, file_path):
                """Save transcript as HTML document."""
                html = f"""<!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Meeting Transcript</title>
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        max-width: 800px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    h1, h2 {{
                        color: #333;
                    }}
                    .metadata {{
                        background-color: #f5f5f5;
                        padding: 10px;
                        border-radius: 5px;
                        margin-bottom: 20px;
                    }}
                    .transcript {{
                        margin-bottom: 30px;
                    }}
                    .insights {{
                        background-color: #e6f7ff;
                        padding: 15px;
                        border-radius: 5px;
                    }}
                    .footer {{
                        margin-top: 40px;
                        font-size: 0.8em;
                        color: #777;
                        text-align: center;
                    }}
                </style>
            </head>
            <body>
                <h1>Meeting Transcript</h1>
                
                <div class="metadata">
                    <p>Generated by {APP_NAME} v{VERSION}</p>
                    <p>Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            """
            # Add duration if available
            if self.meeting_start_time:
                duration = datetime.datetime.now() - self.meeting_start_time
                hours, remainder = divmod(duration.seconds, 3600)
                minutes, seconds = divmod(remainder, 60)
                html += f"<p>Meeting Duration: {hours:02}:{minutes:02}:{seconds:02}</p>\n"
                
            # Add meeting title if detected
            if self.meeting_title and self.meeting_title != "Unknown Meeting":
                html += f"<p>Meeting Title: {self.meeting_title}</p>\n"
                
            html += """
                </div>
                
                <h2>Transcript</h2>
                <div class="transcript">
            """
                
            # Add transcript with timestamps highlighted
            transcript_text = self.transcript_box.toPlainText()
            transcript_html = ""
            for line in transcript_text.split('\n'):
                if line.strip():
                    # Highlight timestamps
                    line_html = re.sub(
                        r'\[(\d+:\d+:\d+)\]', 
                        r'<span style="color: #777;">[<strong>\1</strong>]</span>', 
                        line
                    )
                    transcript_html += f"<p>{line_html}</p>\n"
                
            html += transcript_html
                
            html += """
                </div>
            """
                
            # Add AI insights if available
            if self.insights_box.toPlainText():
                html += """
                <h2>AI Analysis & Insights</h2>
                <div class="insights">
            """
                
                insights_text = self.insights_box.toPlainText()
                insights_html = ""
                for line in insights_text.split('\n'):
                    if line.strip():
                        # Highlight headers
                        if line.startswith("SUGGESTED RESPONSE:") or line.startswith("KEY INSIGHTS:") or line.startswith("TOPIC SUMMARY:"):
                            parts = line.split(":", 1)
                            if len(parts) == 2:
                                header, content = parts
                                insights_html += f"<p><strong style='color: #0066cc;'>{header}:</strong>{content}</p>\n"
                            else:
                                insights_html += f"<p>{line}</p>\n"
                        else:
                            insights_html += f"<p>{line}</p>\n"
                
                html += insights_html
                
                html += """
                </div>
            """
                
            html += f"""
                <div class="footer">
                    <p>Generated by {APP_NAME} v{VERSION} on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
                </div>
            </body>
            </html>
            """
                
            # Write to file
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(html)
                
            def save_as_markdown(self, file_path):
                """Save transcript as Markdown document."""
                md = f"# Meeting Transcript\n\n"
                
                # Add metadata
                md += f"Generated by {APP_NAME} v{VERSION}  \n"
                md += f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  \n"
                
                # Add duration if available
                if self.meeting_start_time:
                    duration = datetime.datetime.now() - self.meeting_start_time
                    hours, remainder = divmod(duration.seconds, 3600)
                    minutes, seconds = divmod(remainder, 60)
                    md += f"Meeting Duration: {hours:02}:{minutes:02}:{seconds:02}  \n"
                
                # Add meeting title if detected
                if self.meeting_title and self.meeting_title != "Unknown Meeting":
                    md += f"Meeting Title: {self.meeting_title}  \n"
                
                md += "\n## Transcript\n\n"
                
                # Add transcript
                transcript_text = self.transcript_box.toPlainText()
                for line in transcript_text.split('\n'):
                    if line.strip():
                        md += f"{line}  \n"
                
                # Add AI insights if available
                if self.insights_box.toPlainText():
                    md += "\n## AI Analysis & Insights\n\n"
                
                    insights_text = self.insights_box.toPlainText()
                    for line in insights_text.split('\n'):
                        if line.strip():
                            md += f"{line}  \n"
                
                # Write to file
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(md)
            def show_wordcloud(self):
                """Show word cloud window."""
                if not self.transcript_box.toPlainText():
                    QMessageBox.warning(self, "Error", "No transcript available for word cloud")
                    return
                
                # Generate word cloud from transcript
                self.refresh_wordcloud()
                
                # Connect signals if not already
                if not self.wordcloud_window.refresh_btn.receivers(self.wordcloud_window.refresh_btn.clicked):
                    self.wordcloud_window.refresh_btn.clicked.connect(self.refresh_wordcloud)
                
                if not self.wordcloud_window.export_btn.receivers(self.wordcloud_window.export_btn.clicked):
                    self.wordcloud_window.export_btn.clicked.connect(self.export_wordcloud)
                
                # Show window
                self.wordcloud_window.show()
                self.wordcloud_window.activateWindow()
                
            def refresh_wordcloud(self):
                """Refresh the word cloud with current transcript."""
                text = self.transcript_box.toPlainText()
                self.signals.update_wordcloud.emit(text)
                
            def export_wordcloud(self):
                """Export the word cloud as an image."""
                if not self.wordcloud_window.label.pixmap():
                    QMessageBox.warning(self, "Error", "No word cloud to export")
                    return
                
                file_path, _ = QFileDialog.getSaveFileName(
                    self, "Export Word Cloud", 
                    f"wordcloud_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
                    "PNG Image (*.png);;JPEG Image (*.jpg);;All Files (*)"
                )
                
                if not file_path:
                    return
                
                if self.wordcloud_window.export_wordcloud(file_path):
                    self.log(f"Word cloud exported to {file_path}")
                
                    # Ask if they want to open the file
                    confirm = QMessageBox.question(
                        self, "Open File", 
                        "Word cloud exported. Would you like to open it now?",
                        QMessageBox.Yes | QMessageBox.No
                    )
                
                    if confirm == QMessageBox.Yes:
                        try:
                            if sys.platform == 'win32':
                                os.startfile(file_path)
                            elif sys.platform == 'darwin':  # macOS
                                os.system(f"open '{file_path}'")
                            else:  # Linux
                                os.system(f"xdg-open '{file_path}'")
                        except Exception as e:
                            print(f"Error opening file: {e}")
                else:
                    QMessageBox.critical(self, "Error", "Failed to export word cloud")
                
            def update_wordcloud(self, text):
                """Update word cloud with new text."""
                self.wordcloud_window.update_wordcloud(text)
                
            def search_transcript(self, query):
                """Search transcript for query and highlight matches."""
                if not query:
                    self.transcript_box.clear_highlights()
                    return
                
                self.transcript_box.clear_highlights()
                self.transcript_box.highlight_text(re.escape(query), "#FFFF99")  # Yellow highlight
                
            def append_transcript(self, text):
                """Append text to transcript with formatting."""
                # Apply spell checking if enabled
                if settings.get("spellcheck_enabled", False):
                    text = correct_spelling(text)
                
                # Apply name anonymization if enabled
                if settings.get("anonymize_transcript", False):
                    text, name_mapping = anonymize_names(text, self.known_names)
                    # Update known names for future anonymization
                    self.known_names = list(name_mapping.keys())
                
                # Check for code word activation
                code_word = settings.get("code_word", DEFAULT_CODE_WORD).lower()
                if code_word and code_word in text.lower() and not self.code_word_detected:
                    self.code_word_detected = True
                    self.log(f"Code word '{code_word}' detected. Triggering analysis...")
                    QTimer.singleShot(500, self.on_demand_analysis)  # Slight delay for visual feedback
                
                # Format based on whether it contains a timestamp
                if text.startswith('[') and ']' in text:
                    # Split into timestamp and content
                    timestamp_end = text.find(']') + 1
                    timestamp = text[:timestamp_end]
                    content = text[timestamp_end:].strip()
                
                    # Add timestamp with gray color
                    self.transcript_box.append_formatted(
                        timestamp, {"color": "#888888", "font_size": 12}
                    )
                
                    # Add content with normal format
                    self.transcript_box.append_formatted(
                        content, {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                    )
                else:
                    # Insert regular text
                    self.transcript_box.append_formatted(
                        text, {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                    )
            def append_suggested_response(self, text):
                """Append text to suggested response box with formatting."""
                # Format "SUGGESTED RESPONSE:" header in bold blue
                if "SUGGESTED RESPONSE:" in text:
                    parts = text.split("SUGGESTED RESPONSE:", 1)
                    before = parts[0]
                    after = parts[1]
                
                    # Add any text before the header
                    if before.strip():
                        self.suggested_response_box.append_formatted(
                            before.strip(), 
                            {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                        )
                
                    # Add header in bold blue
                    self.suggested_response_box.append_formatted(
                        "SUGGESTED RESPONSE:", 
                        {
                            "font_size": 16, 
                            "bold": True, 
                            "color": "#007AFF"
                        }
                    )
                
                    # Add the content
                    self.suggested_response_box.append_formatted(
                        after.strip(), 
                        {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                    )
                else:
                    # Regular text
                    self.suggested_response_box.append_formatted(
                        text, 
                        {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                    )
                
            def append_insights(self, text):
                """Append text to insights box with formatting."""
                # Format headers in bold colored text
                headers = ["KEY INSIGHTS:", "TOPIC SUMMARY:"]
                
                for header in headers:
                    if header in text:
                        parts = text.split(header, 1)
                        before = parts[0]
                        after = parts[1]
                
                        # Add any text before the header
                        if before.strip():
                            self.insights_box.append_formatted(
                                before.strip(), 
                                {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                            )
                
                        # Add header in bold color
                        header_color = "#FF9500" if header == "KEY INSIGHTS:" else "#5AC8FA"
                        self.insights_box.append_formatted(
                            header, 
                            {
                                "font_size": 16, 
                                "bold": True, 
                                "color": header_color
                            }
                        )
                
                        # Add the content
                        self.insights_box.append_formatted(
                            after.strip(), 
                            {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                        )
                
                        return
                
                # If no headers found, just add as regular text
                self.insights_box.append_formatted(
                    text, 
                    {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                )
                
            def highlight_transcript(self, text, color="#FFFF00"):
                """Highlight text in transcript."""
                self.transcript_box.clear_highlights()
                if text:
                    self.transcript_box.highlight_text(re.escape(text), color)
                
            def update_key_phrases(self, phrases):
                """Update list of detected key phrases."""
                if not phrases:
                    return
                
                self.key_phrases = phrases
                
                # Add phrases to transcript highlighter
                for phrase in phrases:
                    self.transcript_box.highlight_text(
                        re.escape(phrase), 
                        "#E0F7FA"  # Light blue highlight
                    )
                
            def update_meeting_title(self, title):
                """Update the detected meeting title."""
                if not title or title == "Unknown Meeting":
                    return
                
                self.meeting_title = title
                self.topic_value.setText(title)
                
            def log(self, message):
                """Update status message."""
                timestamp = datetime.datetime.now().strftime("%H:%M:%S")
                self.signals.update_status.emit(f"[{timestamp}] {message}")
                
            def update_status(self, message):
                """Update status label."""
                self.status.setText(message)
                
            def update_progress(self, value):
                """Update progress bar."""
                self.progress_bar.setValue(value)
                
                # Hide progress bar if complete
                if value >= 100:
                    QTimer.singleShot(1000, lambda: self.progress_bar.setVisible(False))
                else:
                    self.progress_bar.setVisible(True)
                
            def update_meeting_time(self):
                """Update meeting timer display."""
                if not self.meeting_start_time:
                    return
                
                # Calculate elapsed time
                now = datetime.datetime.now()
                elapsed = now - self.meeting_start_time
                hours, remainder = divmod(elapsed.seconds, 3600)
                minutes, seconds = divmod(remainder, 60)
                
                # Update label
                time_str = f"Meeting time: {hours:02}:{minutes:02}:{seconds:02}"
                self.signals.update_meeting_timer.emit(time_str)
                
                # Check if reminder is due
                reminder_interval = settings.get("meeting_reminder_interval", 15) * 60  # Convert to seconds
                if reminder_interval > 0 and self.last_reminder_time:
                    time_since_reminder = (now - self.last_reminder_time).seconds
                    if time_since_reminder >= reminder_interval:
                        # Show reminder
                        self.tray_icon.showMessage(
                            "Meeting Time Reminder", 
                            f"Meeting has been running for {hours:02}:{minutes:02}:{seconds:02}",
                            QSystemTrayIcon.Information,
                            5000  # Show for 5 seconds
                        )
                        self.last_reminder_time = now
            def update_meeting_timer(self, time_str):
                """Update meeting timer label."""
                self.timer_label.setText(time_str)
                
            @asyncSlot()
            async def on_demand_analysis(self):
                """Trigger AI analysis on demand."""
                if not self.transcript_box.toPlainText():
                    self.log("No transcript available for analysis")
                    return
                
                self.log("Generating AI analysis...")
                self.update_progress(10)
                
                # Get current transcript
                transcript = self.transcript_box.toPlainText()
                
                # Get active prompt
                prompt_text = ""
                for p in settings.get("prompts", []):
                    if p["name"] == settings.get("active_prompt", ""):
                        prompt_text = p["prompt"]
                        break
                
                if not prompt_text:
                    prompt_text = DEFAULT_SETTINGS["prompts"][0]["prompt"]
                
                # Combine prompt and transcript
                full_prompt = f"{prompt_text}\n\nCurrent transcript:\n{transcript}"
                
                try:
                    # Show progress animation
                    self.update_progress(30)
                
                    # Call Groq API
                    response = groq_client.chat.completions.create(
                        model=settings.get("model", DEFAULT_MODEL),
                        messages=[
                            {"role": "system", "content": "You are an intelligent meeting assistant."},
                            {"role": "user", "content": full_prompt}
                        ],
                        temperature=0.3,
                        top_p=0.9,
                        max_tokens=1024
                    )
                
                    self.update_progress(80)
                
                    # Extract response
                    result = response.choices[0].message.content.strip()
                
                    # Process response
                    self.process_groq_response(result)
                
                    self.update_progress(100)
                    self.log("AI analysis complete")
                
                except Exception as e:
                    self.log(f"Error generating analysis: {e}")
                    self.update_progress(0)
                
            def process_groq_response(self, response):
                """Process the AI response and update UI."""
                lines = response.split('\n')
                current_section = None
                section_content = {
                    "SUGGESTED RESPONSE": [],
                    "KEY INSIGHTS": [],
                    "TOPIC SUMMARY": []
                }
                
                # Parse response into sections
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                
                    # Check for section headers
                    if "SUGGESTED RESPONSE:" in line:
                        current_section = "SUGGESTED RESPONSE"
                        content = line.split("SUGGESTED RESPONSE:", 1)[1].strip()
                        if content:
                            section_content[current_section].append(content)
                    elif "KEY INSIGHTS:" in line:
                        current_section = "KEY INSIGHTS"
                        content = line.split("KEY INSIGHTS:", 1)[1].strip()
                        if content:
                            section_content[current_section].append(content)
                    elif "TOPIC SUMMARY:" in line:
                        current_section = "TOPIC SUMMARY"
                        content = line.split("TOPIC SUMMARY:", 1)[1].strip()
                        if content:
                            section_content[current_section].append(content)
                    elif current_section:
                        section_content[current_section].append(line)
                
                # Update UI
                if section_content["SUGGESTED RESPONSE"]:
                    suggested = "SUGGESTED RESPONSE:\n" + "\n".join(section_content["SUGGESTED RESPONSE"])
                    self.signals.append_suggested_response.emit(suggested)
                
                if section_content["KEY INSIGHTS"]:
                    insights = "KEY INSIGHTS:\n" + "\n".join(section_content["KEY INSIGHTS"])
                    self.signals.append_insights.emit(insights)
                
                if section_content["TOPIC SUMMARY"]:
                    summary = "TOPIC SUMMARY:\n" + "\n".join(section_content["TOPIC SUMMARY"])
                    self.signals.append_insights.emit(summary)
                
                # If no structured content found, just add everything to insights
                if not any(section_content.values()):
                    self.signals.append_insights.emit(response)
                
            async def scrape_otter(self, url):
                """Scrape Otter.ai transcript."""
                print("[DEBUG] scrape_otter started")
                try:
                    # Initialize playwright
                    async with async_playwright() as p:
                        print("[DEBUG] launching browser")
                        self.browser = await p.chromium.launch(headless=False, timeout=30000)
                        self.page = await self.browser.new_page()
                
                        # Update status
                        self.log("Navigating to Otter.ai...")
                        self.update_progress(30)
                
                        # Navigate to URL with longer timeout
                        try:
                            await self.page.goto(url, timeout=90000)
                        except Exception as e:
                            self.log(f"Navigation timeout, but continuing: {e}")
                
                        # Wait for content to load
                        try:
                            # Try to wait for network idle, but don't fail if it times out
                            self.log("Waiting for page to load...")
                            self.update_progress(40)
                            await self.page.wait_for_load_state("domcontentloaded", timeout=30000)
                            print("[DEBUG] DOM content loaded, proceeding without waiting for network idle")
                        except Exception as e:
                            print(f"[DEBUG] Network idle wait timed out, but continuing: {e}")
                
                        # Longer wait for complete page render
                        print("[DEBUG] sleeping 10s for page render")
                        self.log("Waiting for page to render...")
                        self.update_progress(50)
                        await asyncio.sleep(10)
                
                        # Check if login is required
                        self.log("Checking if login is required...")
                        self.update_progress(60)
                        login_elements = await self.page.query_selector_all(
                            "input[type='password'], input[type='email'], button:has-text('Sign in'), button:has-text('Log in')"
                        )
                        if login_elements:
                            self.log("[WARNING] Login might be required. Please log in through the browser window.")
                            # Wait longer to give user time to log in manually if needed
                            self.update_progress(65)
                            await asyncio.sleep(20)
                
                        # Try to click 'Accept All Cookies' button if it exists
                        try:
                            print("[DEBUG] Looking for cookie consent dialog")
                            self.log("Handling cookie dialogs if present...")
                            self.update_progress(70)
                            accept_buttons = await self.page.query_selector_all(
                                "button:has-text('Accept'), button:has-text('Accept All'), button:has-text('Allow')"
                            )
                            if accept_buttons:
                                print("[DEBUG] Clicking cookie accept button")
                                await accept_buttons[0].click()
                                # Wait for dialog to close
                                await asyncio.sleep(2)
                            else:
                                print("[DEBUG] No cookie dialog found")
                        except Exception as e:
                            print(f"[DEBUG] Error handling cookie dialog: {e}")
                
                        # Start polling for transcript
                        self.log("Starting transcript polling...")
                        self.update_progress(100)
                
                        # Set up polling timer
                        poll_interval = settings.get("poll_interval", DEFAULT_POLL_INTERVAL) * 1000  # convert to ms
                        self.timer.timeout.connect(lambda: asyncio.create_task(self.poll_transcript()))
                        self.timer.start(poll_interval)
                
                except Exception as e:
                    self.log(f"Error: {e}")
                    self.stop()
                
            async def poll_transcript(self):
                """Poll transcript content from Otter.ai."""
                if not self.running or not self.page:
                    return
                
                try:
                    # Try to get meeting title if not already detected
                    if self.meeting_title == "Unknown Meeting":
                        try:
                            title_element = await self.page.query_selector("h1.title")
                            if title_element:
                                title = await title_element.inner_text()
                                if title and title.strip():
                                    self.signals.detected_meeting_title.emit(title.strip())
                        except Exception as e:
                            print(f"[DEBUG] Error getting meeting title: {e}")
                
                    # Get all transcript lines
                    transcript_elements = await self.page.query_selector_all(
                        "#__next .transcriptContainer .line, .transcript-container .transcript-element"
                    )
                
                    new_content = False
                
                    for element in transcript_elements:
                        try:
                            line_text = await element.inner_text()
                            line_text = line_text.strip()
                
                            if not line_text:
                                continue
                
                            # Only process lines we haven't seen before
                            line_hash = hash(line_text)
                            if line_hash not in self.seen_lines:
                                self.seen_lines.add(line_hash)
                                self.buffer.append(line_text)
                                new_content = True
                        except Exception as e:
                            print(f"[DEBUG] Error processing transcript line: {e}")
                
                    if new_content:
                        # Get formatted lines with timestamps
                        formatted_lines = []
                        line_count = 0
                
                        for line in self.buffer:
                            # Try to extract timestamp if not already present
                            if not line.startswith('['):
                                try:
                                    # Look for timestamp in span
                                    time_match = re.search(r'\b(\d{1,2}:\d{2}:\d{2})\b', line)
                                    if time_match:
                                        timestamp = time_match.group(1)
                                        # Remove timestamp from text
                                        text = re.sub(r'\b\d{1,2}:\d{2}:\d{2}\b', '', line).strip()
                                        formatted_lines.append(f"[{timestamp}] {text}")
                                    else:
                                        formatted_lines.append(line)
                                except Exception as e:
                                    formatted_lines.append(line)
                            else:
                                formatted_lines.append(line)
                
                            line_count += 1
                
                        # Add new lines to transcript
                        for line in formatted_lines:
                            self.signals.append_transcript.emit(line)
                
                        # Check if we should trigger analysis based on line interval
                        line_interval = settings.get("summary_line_interval", 10)
                        if line_count >= line_interval:
                            # Reset buffer and trigger analysis
                            self.buffer = []
                            await self.on_demand_analysis()
                
                except Exception as e:
                    print(f"[DEBUG] Error polling transcript: {e}")
                
            # --------------------------
            # Main Entry Point
            # --------------------------
            if __name__ == "__main__":
                # Set up async environment
                app = QApplication(sys.argv)
                loop = QEventLoop(app)
                asyncio.set_event_loop(loop)
                
                # Set application info and icon
                app.setApplicationName(APP_NAME)
                app.setApplicationVersion(VERSION)
                app.setWindowIcon(QIcon(QPixmap.fromImage(QImage.fromData(base64.b64decode(DEFAULT_APP_ICON)))))
                
                # Create and show main window
                window = MeetingAssistant()
                window.show()
                
                # Start event loop
                with loop:
                    sys.exit(loop.run_forever())