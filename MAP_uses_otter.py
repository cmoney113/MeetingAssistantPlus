# Merged Script: MeetingAssistantPlus v2.0 (Based on fixed3 + updated features)

import sys
import os
import asyncio
import base64
import json
import datetime
import re
import threading
import time
from io import BytesIO
from qasync import QEventLoop, asyncSlot

from PyQt5.QtWidgets import (
    QApplication, QWidget, QPushButton, QVBoxLayout, QLabel, QTextEdit, QLineEdit,
    QFileDialog, QMessageBox, QHBoxLayout, QSystemTrayIcon, QMenu, QAction, QComboBox,
    QSplitter, QTabWidget, QDialog, QFormLayout, QSpinBox, QCheckBox, QRadioButton,
    QButtonGroup, QToolButton, QSizePolicy, QShortcut, QProgressBar, QScrollArea,
    QFrame, QGroupBox, QGridLayout, QColorDialog, QFontDialog, QInputDialog, # Added QInputDialog
    QSpacerItem # Added QSpacerItem
)
from PyQt5.QtGui import (
    QPalette, QColor, QFont, QIcon, QPixmap, QKeySequence, QTextCharFormat,
    QSyntaxHighlighter, QTextCursor, QFontMetrics, QPainter, QTextFormat,
    QTextBlockFormat, QBrush, QTextDocument, QImage # Added QImage
)
from PyQt5.QtCore import (
    Qt, QTimer, pyqtSignal, QObject, QSize, QRect, QUrl, QThread,
    QRegExp, QEvent, QPoint, QMargins
)

from playwright.async_api import async_playwright
import groq
from docx import Document
from fpdf import FPDF
from wordcloud import WordCloud
from PIL import Image
import keyboard
from markdown import markdown
import webbrowser

# --------------------------
# Constants and Configuration
# --------------------------
VERSION = "2.0.0" # From updated
APP_NAME = "MeetingAssistant+" # From updated
SETTINGS_PATH = "settings.json" # From fixed3, kept for consistency
DEFAULT_FONT_SIZE = 14 # From updated
DEFAULT_FONT_FAMILY = "SF Pro Text" # From updated
DEFAULT_LINE_SPACING = 150  # From updated
DEFAULT_WORD_SPACING = 1.2  # From updated
DEFAULT_POLL_INTERVAL = 5   # From updated
DEFAULT_CODE_WORD = "blossom" # From updated

# --------------------------
#  API Setup (Using hardcoded key from fixed3)
# --------------------------
groq_client = groq.Client(api_key="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")

# --------------------------
# Default Model Configuration (from updated)
# --------------------------
DEFAULT_MODEL = "llama3-70b-8192"
AVAILABLE_MODELS = [
    "llama3-70b-8192",
    "llama3-8b-8192",
    "mixtral-8x7b-32768",
    "gemma-7b-it"
]

# --------------------------
# Theme Definitions (from updated)
# --------------------------
THEMES = {
    "dark": {
        "window": "#292A2D", "windowText": "#FFFFFF", "base": "#292A2D",
        "alternateBase": "#3C3E44", "text": "#FFFFFF", "button": "#3C3E44",
        "buttonText": "#FFFFFF", "highlight": "#4DB6AC", "highlightedText": "#000000",
        "link": "#88C0D0", "brightText": "#FF5555", "borderColor": "#4D4D4D",
        "gradientStart": "#292A2D", "gradientEnd": "#1E1F21"
    },
    "light": {
        "window": "#F5F6F7", "windowText": "#333333", "base": "#FFFFFF",
        "alternateBase": "#F0F0F0", "text": "#333333", "button": "#E4E4E4",
        "buttonText": "#333333", "highlight": "#4DB6AC", "highlightedText": "#FFFFFF",
        "link": "#0077CC", "brightText": "#FF5555", "borderColor": "#DDDDDD",
        "gradientStart": "#F5F6F7", "gradientEnd": "#E6E7E8"
    },
    "navy": {
        "window": "#0D101E", "windowText": "#E6F1FF", "base": "#0D101E",
        "alternateBase": "#172A45", "text": "#E6F1FF", "button": "#172A45",
        "buttonText": "#E6F1FF", "highlight": "#64FFDA", "highlightedText": "#0A192F",
        "link": "#64FFDA", "brightText": "#FF5555", "borderColor": "#253656",
        "gradientStart": "#0D101E", "gradientEnd": "#0A0D16"
    },
    "dracula": {
        "window": "#282A36", "windowText": "#F8F8F2", "base": "#282A36",
        "alternateBase": "#44475A", "text": "#F8F8F2", "button": "#44475A",
        "buttonText": "#F8F8F2", "highlight": "#FF79C6", "highlightedText": "#282A36",
        "link": "#8BE9FD", "brightText": "#FF5555", "borderColor": "#383A48",
        "gradientStart": "#282A36", "gradientEnd": "#21222C"
    },
    "macOS": {
        "window": "#EBEBEB", "windowText": "#333333", "base": "#FFFFFF",
        "alternateBase": "#F7F7F7", "text": "#333333", "button": "#F2F2F2",
        "buttonText": "#333333", "highlight": "#007AFF", "highlightedText": "#FFFFFF",
        "link": "#007AFF", "brightText": "#FF3B30", "borderColor": "#DDDDDD",
        "gradientStart": "#F7F7F7", "gradientEnd": "#EBEBEB"
    },
    "macOSDark": {
        "window": "#323232", "windowText": "#FFFFFF", "base": "#2A2A2A",
        "alternateBase": "#3A3A3A", "text": "#FFFFFF", "button": "#4A4A4A",
        "buttonText": "#FFFFFF", "highlight": "#0A84FF", "highlightedText": "#FFFFFF",
        "link": "#0A84FF", "brightText": "#FF453A", "borderColor": "#494949",
        "gradientStart": "#323232", "gradientEnd": "#2A2A2A"
    }
}

# --------------------------
# Default Settings (merged from updated and fixed3)
# --------------------------
DEFAULT_PROMPT_CONTENT = """You are an intelligent assistant attending a professional job interview. Analyze the conversation in real time and help the user prepare to answer if someone poses a question to them.

Instructions:
- Detect if another speaker asks a question directed at the user (e.g., using "you", "your experience", or direct names).
- When a question is asked, generate a contextual, helpful response or insight the user could consider sharing.
- Constantly monitor the discussion to summarize the current topic and key takeaways.
- Present this information clearly and concisely.
- If appropriate, include action points or things the user should pay attention to.

Output format:
SUGGESTED RESPONSE: [Your suggested response if a question was asked]

KEY INSIGHTS: [Important insights about the conversation]

TOPIC SUMMARY: [Brief summary of the current topic]
"""

DEFAULT_SETTINGS = {
    "active_prompt": "Job Interview Assistant",
    "prompts": [{
        "name": "Job Interview Assistant",
        "prompt": DEFAULT_PROMPT_CONTENT
    }],
    "summary_line_interval": 10,
    "summary_time_interval": 120, # seconds
    "poll_interval": DEFAULT_POLL_INTERVAL, # seconds
    "theme": "macOS",
    "start_in_background": False,
    "font_size": DEFAULT_FONT_SIZE,
    "font_family": DEFAULT_FONT_FAMILY,
    "line_spacing": DEFAULT_LINE_SPACING, # percentage
    "word_spacing": DEFAULT_WORD_SPACING, # multiplier
    "model": DEFAULT_MODEL,
    "keyboard_shortcut": "ctrl+i",
    "code_word": DEFAULT_CODE_WORD,
    "app_icon_path": "", # Paths for custom icons
    "generate_icon_path": "",
    "settings_icon_path": "",
    "anonymize_transcript": False,
    "spellcheck_enabled": False,
    "meeting_timer_enabled": True,
    "meeting_reminder_interval": 15,  # minutes
    "use_rounded_corners": True,
    "use_animations": True,
    "use_transparency": True,
}

# --------------------------
# Embedded Icons (Base64 - from updated)
# --------------------------
DEFAULT_APP_ICON = "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAACXBIWXMAAAsTAAALEwEAmpwYAAAF7ElEQVR4nO2bW2xURRjHf93dttBSKHcRqKWCBYJGQSAqIheuJhAUEzQqEBJvD5oYE33AB0x8UI2aGBOFqBiJGkWJd6IYRdAIhHspFGhpuRQo0ALttrB7fPhmOWfbc2b2zDlnd9v9JZPuzsx833z/uXxnZs+BEkooIXfiLuAssKzYisQ0eUCXeZxX0L7HAH3ALGDmIPQTKxwEHgTuNs9HB6HPVJBrBMwDngSmF0gXSYeRwMPAEuP0AqA3i+I6L2yqkRHgllFXcQcQ5sQ4RoCqAtCNDXxGgBoZSdIxAvpyJYjC7H8d2DpA5UKwB7jRJa8qPK+64pXc1AEMNpTbTbXAKeAPYHSK+DGIRGKwv0B8cBcQxEpgB/AzsNwlbw7wK/ArMNkl/1HEgLeBWYNhwO8h+VsTyNsd/rNIIKXCXqSdVAc5fAuwE7ijAPqNAK4DXgXWACnlRFUMFRqBy4BLgU7gR2BzwD1vApYCdcBxYD0wyuRNBC4CJgDHgF+Adqtsa8QE+wwVYiZsb91dttLdiSSCXcAmcz3NlB8FWgJ12AEsNG30ADPDGDCB4CJoRrA9zJ8EbqXzaYE8gDrgHiTiROmpAxYjgxi52NEdks9hynzgCZO3FVgBXBR4T52DdcAjwHeIDw6xCIpgW3kVw3jgNaAZ2Ae8iIwEhc1Ie00HXou0tREY53CfPUo2Bziq5MuMbKzJdL8iE2CHvRfQRrJttIQa1GZGxgzjfDuwzl4EVVlKPoKsx78Bs4EPzPU/yKTGxjZkONt4C6hR8vcAU03+bF8D3LAN+BE4gXS8FnkR37jkzUXmlOUW2WrEsAsRJ04j79fLkPoUwOeAlcA1yIgKrQPijGZgF/CJtg44iqxvNYH3dADrzPVnJv0Vn8rLwhhKZuIU7p6HA9WIfaFQw0dV9qnwfJtqrZvzBXLgNcAFyiQnaFBylRwxUkbCO3EYcKFyPxJJa4P0eD1izx7EvjlINHnXp/IBSAy02GXuDYFaXuOULLhz9ci8P0MHMnQPmvRF4G/kKA0y1aZZZQ4gS/RQZHRsBFab6y5gbcRoaVJlLlXypQm07HqoICPqZauOUWGjKWO0KfdpDqrrjkbmKltVGRs3K3mbkqm2txNQptx3tpKvV+6rkqXxdK4L+FDJg7EO+T5/ZJFrU7BXydrjf7vM47PIn9dAp/lgNDrhwGLvN8sYoLuuziFvV2U06O2wRnUdavQGj2sXNCjZYSWfoGRdNicxfpJVRof4euBdjyfhYk8DQuqALuC3WjnqAqkD7lP1eBswT8k+C2ijzWybSdW/y+i9JoDWGXURpFH0BuU+pA7Q+CJL3iLgTav8UPP3D0Lr6YpYOTVqVfucwA7rjWy3ufZZgnW7TQ5o40Ml6/CpfCiUGHkn/sRzFYiiNTJN6hFJBU5G7k6glUBHR+tHDjw97PxOcaH1VCCnVLuR1fJu4GXgc+T0OMtxGj1hehDYH5Kvc8E4TuGI+dM3X+tN5bOtkMnS0Hm6XsJG+pR9sZKpQA6+gnE56mwqoEbIcVgceNr0eQyZcJXlCDeQI/ORSu6dCzqATnnbPLcCwbOzdmBn6KGJjlw9/5Rl5dwE/B/wkWWmX7dXdvAQ+x1/eGAv0NfMXpXWxYAHgctDGk0gXlFldJi8zaSXIUO0FznYiaxGNLKpugVYinw7OI0cRn6G7HbrkGW7ATl9qkO+kB83tE5ZX1UgI2wRsmG5ErkvcnARDVvBwTLgIJJdjXxAvAmZS/o+UO/IOjtYjtQlv5jn95FtOhHaMk1g/GG7c3lD3Uw9ZBOzCNk5xlrNvJFnPbKiQBFRr8r2RXTKAKnz0/NpQJo0T3ehj9TqMnz7jjJCnx8D0iNDIXM2GeALnWs4R0bAfPIwZ58jo+chz4+LiWGpIzww3DxODs+nF+8YGaXI8hzhqELOyXVedLRGxidA6/9EvIQBxL/n+ibGRp+7nQAAAABJRU5ErkJggg=="
DEFAULT_SETTINGS_ICON = "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAACXBIWXMAAAsTAAALEwEAmpwYAAACnUlEQVR4nO2bS2sUQRDHf5OsJkGNEo0XUTxFiaCCeFE8KH4BRT2IB8WzFw9+A0G/gF/AgyjiK4ioRHzFq5hE8yDmgVHjBpNoMrO7HgZmDZr0TPfszux0wT9gDzM1Vf3vjZ6q7h0oKSkpyTe7gSHgFjAN/AG+A++Bh8B5oBfolPY7jee79DMs99j3TAMPgHNAt/RbCLqAu8An0ZQmPAeGgIPAtgb9twOH5Zlxea5RvwmgXyYrNxwB3ko0FWEZuAn0NTHGXuCWPKvJ/wg0aU0uRsElYMlDWf8DvAJGgUvAZWAM+OThPekzJWNsijMiC6eROCLCJjwEl4DRmISPgV6HMfeK0JrId8DF5GRuzHTF9Cv0D+AEsNHBz03ASeB7RPYrcCSR0C42AS8VuVlgh7YfEbmasGiw7Bb/d+Bn5d8+/dWciqBcRf5Jn4F9KoLP0wS+KPIDGvwBReYl2kxJL/apCT6QKWA4TeBbRX5Z6TvVBE9KuC+0t1qlvUuR+Qps1RBbh98x/o3OFN5rTOaFBr9XkaFegK+Q6XrLHyEHmBNOiW81yfkc+8OZ6prkUG0C9uy9VzszHxS5L9pcPUw+tZTp+1/v4bTlBVQzKgAAAABJRU5ErkJggg=="
DEFAULT_GENERATE_ICON = "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAACXBIWXMAAAsTAAALEwEAmpwYAAAD5klEQVR4nO2aXWjOYRjGr4+Z2Ij5KJGU+WimJhpSalIzopkPMR8H5IDIgTjAyYQDB1JKFGuKFCkHREopB8qBpChfOdqBKEVTTDN/13q73q3t/f6e93mfj3d7r+rX+3zc93U/7/s893PfzwvEiBFjDGAmgCYAhNkAGwCOARgXsm2jADQAOALAFbLzAI4DGBuBfLl4fABwCMDoUAwA8ES2nT5APYCxA2jbEAD1AJ6atr0DUDXQzlcYdATAkAFmwBAAhwF8MvJ9AbATgJ3rDbCdJQDqADzP0tl+AM8AbMl1wBHYRpnudFnwAVSGRn4vTgX0BMAyAG7I9jkA7gJ4F6ANnYYOb+dTQZDn6XxU+3o3gAJD9gxj+Swp9GQeMAR4GIA8pdz1nv93BYBbAOZ7rp+V4SuiREILtS3KApKnLPZuAMOMHWyXz3BFrXdE+f+2CWBNFPLFAG4DOKdsm+lhRyWAz552fQJQ5lPhZgVxH4AxWqUAjAZwV9kmrzfcxzSjjbX6OmvhWUYfHwMoA1Cr5LpGJm+RUu4Q45nfA0ARWUQBLPV0MgOgTcntM9ZYpZTnGoYEtM0qyoxlLAPghJbUcZzrAE4aJe/vGsf5Oy+OiYcAtuFl5c0TaHh6+VWYRQLPyPr8J54xq1DaSUDDm4PtMnZJW7lk+sKvAAUyZTCAxwKd7A7pA+qU62Sx/2Cwy+cw7QYsYLXQ5LMSVGbJbwlb0Hag3pE9hEXsEJr4AEm+S/aeM4JWJf1gGVOFJr0oQNJFSlkZtFOg2U+cHO0Ep0gBWiGFBHQA+CWEKlHihlDWVzFJSOGjEkzuCGWlEolENKEX04UUXipB3RgllZTwqsVRQsIrIYWpSrAEySW3lOOjAuWEWCkxS0jhsxIsFCStVsrfU/q5YoDQQgnmCCn0KMG8J0qbD0LCXKGEV4rQR0G+LEJ+XlGeGSZ0FHK1W5Iy5Z/KfrdHwRkGUQpgohJsUlaUvqQM6+QipGzTd4CuIMQGsaFDSFmvpLSRFraaYdCCzJNkIfZHwQkhXRHCdBLQqTBiI4ASIX2rAC/tWwbgp4BOhkmsE9L12UyIpJt3RSwAUCugl8pzx4bLNhH6C0AWCGol9LID6mfYL0S4UeDeXl4a3PGcPgXFX8uCGzY3GzjOv78D+J0JTsgdVSFRDuCcZwNTkWtQFBmZnZ904eRMbpwRdtjA7RFWNtnQWQmPL0pA9lVXKCezoSDbmGPz1X1hFWRRkMOO3aSHsRLtDZJlTHM7e4OkWWqSPtvQGYwRI0aMf/wGCaZOUO2BGa0AAAAASUVORK5CYII="

# --------------------------
# Settings Management (merged)
# --------------------------
def load_settings():
    """Load application settings from JSON file or use defaults."""
    try:
        if os.path.exists(SETTINGS_PATH):
            with open(SETTINGS_PATH, 'r') as f:
                saved_settings = json.load(f)
                # Merge saved settings with default settings to handle new keys
                merged_settings = DEFAULT_SETTINGS.copy()
                merged_settings.update(saved_settings)
                return merged_settings
    except Exception as e:
        print(f"Error loading settings: {e}")
    return DEFAULT_SETTINGS.copy()

settings = load_settings() # Load settings at startup

def save_settings_to_file(current_settings):
    """Save settings to JSON file."""
    try:
        with open(SETTINGS_PATH, "w") as f:
            json.dump(current_settings, f, indent=4)
    except Exception as e:
        print(f"Error saving settings: {e}")

# --------------------------
# Groq API Call Function (Using structure from fixed3)
# --------------------------
def ask_groq(prompt):
    """Sends a prompt to the configured Groq model with improved error handling."""
    if not prompt or not prompt.strip():
        return "[AI Error] Empty prompt provided to AI API"
        
    # Check API key validity
    if not groq_client.api_key or groq_client.api_key == "your_api_key_here":
        return "[AI Error] Invalid or missing Groq API key. Please check your settings."
    
    try:
        # Select model with validation
        model_name = settings.get("model", DEFAULT_MODEL)
        if model_name not in AVAILABLE_MODELS:
            print(f"Warning: Unknown model {model_name}, falling back to {DEFAULT_MODEL}")
            model_name = DEFAULT_MODEL
            
        # Prepare request with appropriate timeout
        print(f"Sending request to Groq API using model: {model_name}")
        response = groq_client.chat.completions.create(
            model=model_name,
            messages=[
                # System prompt could potentially be customized here later
                {"role": "system", "content": "You are an intelligent assistant helping with meeting transcripts."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1024,  # Increased max tokens
            temperature=0.3,  # Using lower temp for more focus
            top_p=0.9,        # Added top_p parameter
            timeout=60        # Set explicit timeout
        )
        
        # Validate response
        if not response or not hasattr(response, 'choices') or not response.choices:
            return "[AI Error] Received empty or invalid response from Groq API"
            
        return response.choices[0].message.content.strip()
        
    except groq.error.RateLimitError:
        return "[AI Error] Rate limit exceeded. Please try again in a few moments."
    except groq.error.AuthenticationError:
        return "[AI Error] API authentication failed. Please check your API key."
    except groq.error.InvalidRequestError as e:
        return f"[AI Error] Invalid request: {e}"
    except groq.error.APIError as e:
        return f"[AI Error] Groq API error: {e}"
    except groq.error.Timeout:
        return "[AI Error] Request timed out. Please try again."
    except Exception as e:
        print(f"Unexpected error contacting Groq API: {e}")
        import traceback
        traceback.print_exc()
        return f"[AI Error] Failed to get response from Groq: {str(e)[:100]}..."

def get_active_prompt_content():
    """Gets the content of the currently active prompt template."""
    active_prompt_name = settings.get("active_prompt", "")
    prompts = settings.get("prompts", [])
    for p in prompts:
        if p["name"] == active_prompt_name:
            return p["prompt"]
    # Fallback to the default prompt content if active one not found
    return DEFAULT_PROMPT_CONTENT

# --------------------------
# Helper Functions (merged from updated)
# --------------------------
def load_icon_from_path_or_default(path, default_base64):
    """Load an icon from a file path or fall back to the default base64 icon."""
    if path and os.path.exists(path):
        try:
            return QIcon(path)
        except Exception as e:
            print(f"Error loading icon from {path}: {e}")

    # Fall back to default
    pix = QPixmap()
    pix.loadFromData(base64.b64decode(default_base64))
    icon = QIcon()
    icon.addPixmap(pix)
    return icon

def create_styled_button(text, icon=None, tooltip=None, style_sheet=None):
    """Create a styled button with optional icon and tooltip."""
    btn = QPushButton(text)
    if icon:
        btn.setIcon(icon)
    if tooltip:
        btn.setToolTip(tooltip)
    btn.setMinimumHeight(36) # Default height from updated

    # Apply custom stylesheet if provided, else use default
    if style_sheet:
        btn.setStyleSheet(style_sheet)
    else: # Default style from updated
         btn.setStyleSheet("""
            QPushButton {
                background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #5E5E5E, stop:1 #444444);
                border: 1px solid #555;
                border-radius: 6px;
                padding: 5px 15px;
                color: white;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #656565, stop:1 #4F4F4F);
            }
            QPushButton:pressed {
                background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #444444, stop:1 #5E5E5E);
                padding-top: 6px;
                padding-bottom: 4px;
            }
            QPushButton:disabled {
                background-color: #3A3A3A;
                color: #7F7F7F;
            }
        """)
    return btn

def create_icon_button(icon, tooltip=None, size=32):
    """Create a button with just an icon (from updated)."""
    btn = QToolButton()
    btn.setIcon(icon)
    if tooltip:
        btn.setToolTip(tooltip)
    btn.setIconSize(QSize(size, size))
    btn.setMinimumSize(QSize(size + 12, size + 12))
    btn.setStyleSheet("""
        QToolButton {
            background-color: transparent;
            border: 1px solid #555;
            border-radius: 6px;
            padding: 3px;
        }
        QToolButton:hover {
            background-color: rgba(255, 255, 255, 0.1);
        }
        QToolButton:pressed {
            background-color: rgba(0, 0, 0, 0.1);
            padding-top: 4px;
            padding-bottom: 2px;
        }
    """)
    return btn

# --------------------------
# Export Functions (Enhanced versions from updated)
# --------------------------
def export_to_txt(text, filepath):
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(text)

def export_to_docx(text, filepath, metadata=None): # Added metadata option
    doc = Document()
    # Add title
    doc.add_heading("Meeting Transcript", 0)
    # Add metadata if provided
    if metadata:
        doc.add_paragraph(f"Generated by {metadata.get('app_name', APP_NAME)} v{metadata.get('version', VERSION)}")
        doc.add_paragraph(f"Date: {metadata.get('date', datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}")
        if metadata.get('duration'):
            doc.add_paragraph(f"Meeting Duration: {metadata.get('duration')}")
        if metadata.get('title') and metadata.get('title') != "Unknown Meeting":
             doc.add_paragraph(f"Meeting Title: {metadata.get('title')}")
        doc.add_paragraph() # Spacer

    doc.add_heading("Transcript", 1)
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    doc.save(filepath)

def export_to_pdf(text, filepath, metadata=None): # Added metadata option
    pdf = FPDF()
    pdf.add_page()
    # Add title
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "Meeting Transcript", 0, 1, 'C')
    pdf.ln(5)
    # Add metadata
    pdf.set_font("Arial", '', 10)
    if metadata:
        pdf.cell(0, 5, f"Generated by {metadata.get('app_name', APP_NAME)} v{metadata.get('version', VERSION)}", 0, 1)
        pdf.cell(0, 5, f"Date: {metadata.get('date', datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}", 0, 1)
        if metadata.get('duration'):
            pdf.cell(0, 5, f"Meeting Duration: {metadata.get('duration')}", 0, 1)
        if metadata.get('title') and metadata.get('title') != "Unknown Meeting":
            pdf.cell(0, 5, f"Meeting Title: {metadata.get('title')}", 0, 1)
    pdf.ln(5)
    # Add transcript
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, "Transcript", 0, 1)
    pdf.set_font("Arial", '', 10)
    pdf.multi_cell(0, 5, text.encode('latin-1', 'replace').decode('latin-1')) # Handle potential encoding issues
    pdf.output(filepath)

def export_to_srt(text, filepath): # Using smarter SRT logic from updated
    lines = text.strip().split('\n')
    srt = ""
    subtitle_index = 1
    last_timestamp = None
    current_text = ""

    for line in lines:
        line = line.strip()
        if not line: continue

        timestamp_match = re.match(r'\[(\d{1,2}:\d{1,2}:\d{1,2})\]', line)
        if timestamp_match:
            if last_timestamp and current_text:
                start_time = last_timestamp
                try:
                    h, m, s = map(int, start_time.split(':'))
                    start_seconds = h * 3600 + m * 60 + s
                    end_seconds = start_seconds + 5 # Simple duration estimate
                    end_h, rem = divmod(end_seconds, 3600)
                    end_m, end_s = divmod(rem, 60)
                    end_time = f"{end_h:02d}:{end_m:02d}:{end_s:02d}"
                    srt_start = f"{start_time},000"
                    srt_end = f"{end_time},000"
                    srt += f"{subtitle_index}\n{srt_start} --> {srt_end}\n{current_text}\n\n"
                    subtitle_index += 1
                except ValueError:
                    print(f"Warning: Could not parse timestamp {start_time} for SRT")
                current_text = ""

            last_timestamp = timestamp_match.group(1)
            line = line[line.find(']')+1:].strip() # Remove timestamp

        if line: # Only add if there's content after timestamp removal
            if current_text: current_text += "\n"
            current_text += line

    # Add final subtitle block
    if last_timestamp and current_text:
        start_time = last_timestamp
        try:
            h, m, s = map(int, start_time.split(':'))
            start_seconds = h * 3600 + m * 60 + s
            end_seconds = start_seconds + 5
            end_h, rem = divmod(end_seconds, 3600)
            end_m, end_s = divmod(rem, 60)
            end_time = f"{end_h:02d}:{end_m:02d}:{end_s:02d}"
            srt_start = f"{start_time},000"
            srt_end = f"{end_time},000"
            srt += f"{subtitle_index}\n{srt_start} --> {srt_end}\n{current_text}\n\n"
        except ValueError:
             print(f"Warning: Could not parse final timestamp {start_time} for SRT")

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(srt)

def export_to_html(text, filepath, metadata=None): # Added from updated
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Meeting Transcript</title>
    <style>
        body {{ font-family: sans-serif; line-height: 1.6; max-width: 800px; margin: 20px auto; padding: 15px; border: 1px solid #ddd; border-radius: 8px; }}
        h1, h2 {{ color: #333; border-bottom: 1px solid #eee; padding-bottom: 5px; }}
        .metadata {{ background-color: #f9f9f9; padding: 10px; border-radius: 5px; margin-bottom: 20px; font-size: 0.9em; color: #555; }}
        .transcript p {{ margin: 5px 0; }}
        .timestamp {{ color: #888; font-weight: bold; margin-right: 5px; }}
        .insights {{ background-color: #eef; padding: 15px; border-radius: 5px; margin-top: 20px; }}
        .insights strong {{ color: #0056b3; }}
    </style>
</head>
<body>
    <h1>Meeting Transcript</h1>
    <div class="metadata">"""
    if metadata:
        html += f"<p>Generated by {metadata.get('app_name', APP_NAME)} v{metadata.get('version', VERSION)}</p>"
        html += f"<p>Date: {metadata.get('date', datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}</p>"
        if metadata.get('duration'): html += f"<p>Meeting Duration: {metadata.get('duration')}</p>"
        if metadata.get('title') and metadata.get('title') != "Unknown Meeting": html += f"<p>Meeting Title: {metadata.get('title')}</p>"
    html += """</div>
    <h2>Transcript</h2>
    <div class="transcript">"""
    transcript_html = ""
    for line in text.split('\n'):
        if line.strip():
            line_html = re.sub(r'\[(\d{1,2}:\d{1,2}:\d{1,2})\]', r'<span class="timestamp">[\1]</span>', line)
            transcript_html += f"<p>{line_html}</p>\n"
    html += transcript_html
    html += """</div>"""
    # Add insights section if metadata contains it
    if metadata and metadata.get('insights'):
        html += """<h2>AI Analysis & Insights</h2><div class="insights">"""
        insights_html = ""
        for line in metadata['insights'].split('\n'):
             if line.strip():
                 # Highlight headers
                 if re.match(r'^(SUGGESTED RESPONSE|KEY INSIGHTS|TOPIC SUMMARY):', line):
                     parts = line.split(":", 1)
                     insights_html += f"<p><strong>{parts[0]}:</strong>{parts[1]}</p>\n"
                 else:
                     insights_html += f"<p>{line}</p>\n"
        html += insights_html
        html += """</div>"""

    html += """</body></html>"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)

def export_to_markdown(text, filepath, metadata=None): # Added from updated
    md = f"# Meeting Transcript\n\n"
    if metadata:
        md += f"Generated by {metadata.get('app_name', APP_NAME)} v{metadata.get('version', VERSION)}  \n"
        md += f"Date: {metadata.get('date', datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}  \n"
        #                                                                                      ^
        #                                                                                      Added missing parenthesis here
        if metadata.get('duration'): md += f"Meeting Duration: {metadata.get('duration')}  \n"
        if metadata.get('title') and metadata.get('title') != "Unknown Meeting": md += f"Meeting Title: {metadata.get('title')}  \n"
    md += "\n## Transcript\n\n"
    for line in text.split('\n'):
        if line.strip(): md += f"{line}  \n" # Add double space for line break
    if metadata and metadata.get('insights'):
        md += "\n## AI Analysis & Insights\n\n"
        for line in metadata['insights'].split('\n'):
            if line.strip():
                # Make headers bold
                line = re.sub(r'^(SUGGESTED RESPONSE|KEY INSIGHTS|TOPIC SUMMARY):', r'w', line)
                md += f"{line}  \n"
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(md)

# --------------------------
# Word Cloud Generator (from updated)
# --------------------------
def generate_wordcloud(text):
    if not text or text.isspace():
        print("Warning: Empty text provided for word cloud generation.")
        # Create a blank image or return None
        img = Image.new('RGB', (800, 400), color = 'white')
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        return buffer.getvalue()
    try:
        wc = WordCloud(width=800, height=400, background_color='black', colormap='Blues').generate(text)
        buffer = BytesIO()
        wc.to_image().save(buffer, format="PNG")
        return buffer.getvalue()
    except ValueError as e:
        print(f"Error generating word cloud (likely empty text after processing): {e}")
        img = Image.new('RGB', (800, 400), color = 'white')
        draw = ImageDraw.Draw(img)
        draw.text((10,10), "No words found for cloud", fill=(0,0,0))
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        return buffer.getvalue()
    except Exception as e:
        print(f"Unexpected error generating word cloud: {e}")
        return None # Indicate failure

# --------------------------
# Name Anonymizer (from updated)
# --------------------------
def anonymize_names(text, known_names=None):
    """Replace detected names with anonymous identifiers."""
    if not known_names: known_names = []
    name_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b' # More robust name pattern
    potential_names = set(re.findall(name_pattern, text))
    all_names = potential_names.union(set(known_names))
    # Filter out common words that might be capitalized at sentence start
    common_words = {'The', 'A', 'An', 'Is', 'Are', 'Was', 'Were', 'He', 'She', 'It', 'They', 'We', 'You', 'I'}
    filtered_names = {name for name in all_names if name not in common_words and len(name) > 2}

    name_mapping = {name: f"Person-{i+1}" for i, name in enumerate(filtered_names)}
    anonymized_text = text
    # Replace longer names first to avoid partial replacements
    for name in sorted(filtered_names, key=len, reverse=True):
        anonymized_text = re.sub(r'\b' + re.escape(name) + r'\b', name_mapping[name], anonymized_text)
    return anonymized_text, name_mapping

# --------------------------
# Simple Spell Checker (from updated)
# --------------------------
def correct_spelling(text):
    """Simple spell checking and correction."""
    corrections = {
        "teh": "the", "taht": "that", "thier": "their", "waht": "what", "dont": "don't",
        "cant": "can't", "wont": "won't", "im": "I'm", "youre": "you're", "theyre": "they're",
        "thats": "that's", "isnt": "isn't", "arent": "aren't", "hasnt": "hasn't", "havent": "haven't",
        "doesnt": "doesn't", "didnt": "didn't", "wasnt": "wasn't", "werent": "weren't",
        "wouldnt": "wouldn't", "shouldnt": "shouldn't", "couldnt": "couldn't", "alot": "a lot",
        "alright": "all right", "aswell": "as well", "becuase": "because", "definately": "definitely",
        "wierd": "weird", "recieve": "receive", "seperate": "separate", "untill": "until",
        "tommorow": "tomorrow", "tommorrow": "tomorrow", "alway": "always", "ofcourse": "of course",
        "ofcource": "of course", "allready": "already", "allright": "all right", "accross": "across",
        "beleive": "believe", "comunicate": "communicate", "greatful": "grateful", "dependance": "dependence",
        "excelent": "excellent", "existance": "existence", "foreward": "forward", "grammer": "grammar",
        "gaurd": "guard", "hight": "height", "humerous": "humorous", "independant": "independent",
        "liason": "liaison", "millenium": "millennium", "momento": "memento", "neccessary": "necessary",
        "occured": "occurred", "occurance": "occurrence", "peice": "piece", "preform": "perform",
        "prefered": "preferred", "probly": "probably", "reccommend": "recommend", "refered": "referred",
        "rember": "remember", "remeber": "remember", "sieze": "seize", "sucess": "success",
        "suprise": "surprise", "tomatos": "tomatoes", "vaccum": "vacuum", "writeing": "writing"
    }
    corrected_text = text
    # Use regex to replace whole words only, case-insensitively
    for wrong, right in corrections.items():
        pattern = re.compile(r'\b' + re.escape(wrong) + r'\b', re.IGNORECASE)
        # Function to handle case preservation during replacement
        def replace_case(match):
            original = match.group(0)
            if original.islower(): return right.lower()
            if original.isupper(): return right.upper()
            if original.istitle(): return right.title()
            return right # Default case
        corrected_text = pattern.sub(replace_case, corrected_text)
    return corrected_text

# --------------------------
# Signal Bridge (Enhanced from updated)
# --------------------------
class SignalBridge(QObject):
    append_transcript = pyqtSignal(str)
    append_suggested_response = pyqtSignal(str) # Specific signal
    append_insights = pyqtSignal(str) # Specific signal (includes summary)
    update_status = pyqtSignal(str)
    update_progress = pyqtSignal(int) # For progress bar
    highlight_transcript = pyqtSignal(str, str) # text, color hex
    update_wordcloud = pyqtSignal(str)
    update_meeting_timer = pyqtSignal(str) # For timer label
    detected_key_phrases = pyqtSignal(list) # from updated
    detected_meeting_title = pyqtSignal(str) # from updated

# --------------------------
# Custom Text Highlighter (from updated)
# --------------------------
class TranscriptHighlighter(QSyntaxHighlighter):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.highlight_patterns = {} # Use dict {pattern_str: (regexp, format)}

    def add_highlight_pattern(self, pattern_str, color_hex, weight=QFont.Normal):
        """Add or update a highlight pattern."""
        regexp = QRegExp(pattern_str, Qt.CaseInsensitive, QRegExp.RegExp2)
        text_format = QTextCharFormat()
        text_format.setBackground(QColor(color_hex))
        text_format.setFontWeight(weight)
        self.highlight_patterns[pattern_str] = (regexp, text_format)
        self.rehighlight()

    def remove_highlight_pattern(self, pattern_str):
        """Remove a highlight pattern."""
        if pattern_str in self.highlight_patterns:
            del self.highlight_patterns[pattern_str]
            self.rehighlight()

    def clear_highlight_patterns(self):
        """Clear all highlight patterns."""
        self.highlight_patterns = {}
        self.rehighlight()

    def highlightBlock(self, text):
        for pattern_str, (regexp, text_format) in self.highlight_patterns.items():
            index = regexp.indexIn(text)
            while index >= 0:
                length = regexp.matchedLength()
                if length > 0: # Ensure we matched something
                    self.setFormat(index, length, text_format)
                index = regexp.indexIn(text, index + length) # Continue search after match

# --------------------------
# Custom Text Edit (Enhanced from updated)
# --------------------------
class CustomTextEdit(QTextEdit):
    """Enhanced text edit with custom formatting capabilities."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.highlighter = TranscriptHighlighter(self.document())
        self.setReadOnly(True)
        self.setAcceptRichText(True) # Ensure rich text is accepted
        # Set default font and styling
        self.document().setDefaultFont(QFont(settings.get("font_family", DEFAULT_FONT_FAMILY),
                                           settings.get("font_size", DEFAULT_FONT_SIZE)))
        self.setStyleSheet("background-color: transparent; border: none;") # Make background transparent

    def highlight_text(self, pattern_str, color="#FFFF99"): # Yellow default
        """Highlight text matching the pattern string."""
        self.highlighter.add_highlight_pattern(pattern_str, color)

    def clear_highlights(self):
        """Clear all highlights."""
        self.highlighter.clear_highlight_patterns()

    def append_formatted(self, text, format_specs=None):
        """Add text with customized formatting."""
        print(f"DEBUG: append_formatted called with text: {text[:20]}...")
        """Append text with custom formatting."""
        cursor = self.textCursor()
        cursor.movePosition(QTextCursor.End)

        # Apply block format for line spacing
        block_fmt = QTextBlockFormat()
        block_fmt.setLineHeight(settings.get("line_spacing", DEFAULT_LINE_SPACING), QTextBlockFormat.ProportionalHeight)
        cursor.setBlockFormat(block_fmt)

        # Apply character format
        char_fmt = QTextCharFormat()
        default_font = QFont(settings.get("font_family", DEFAULT_FONT_FAMILY),
                             settings.get("font_size", DEFAULT_FONT_SIZE))
        char_fmt.setFont(default_font)

        if format_specs:
            if "font_size" in format_specs: char_fmt.setFontPointSize(format_specs["font_size"])
            if "bold" in format_specs and format_specs["bold"]: char_fmt.setFontWeight(QFont.Bold)
            else: char_fmt.setFontWeight(QFont.Normal) # Ensure not bold if not specified
            if "italic" in format_specs and format_specs["italic"]: char_fmt.setFontItalic(True)
            if "color" in format_specs: char_fmt.setForeground(QColor(format_specs["color"]))
            if "font_family" in format_specs: char_fmt.setFontFamily(format_specs["font_family"])

        # Set word spacing - applied document-wide, maybe not ideal here but can set on block
        # char_fmt.setWordSpacing(settings.get("word_spacing", DEFAULT_WORD_SPACING))

        cursor.insertText(text, char_fmt)
        cursor.insertBlock() # Ensure newline separation

        self.ensureCursorVisible()

# --------------------------
# WordCloud Window (Enhanced from updated)
# --------------------------
class WordCloudWindow(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_app = parent # Store reference to main app
        self.setWindowTitle("Live Word Cloud")
        self.setGeometry(300, 300, 800, 500)
        self.label = QLabel(self)
        self.label.setAlignment(Qt.AlignCenter)

        # Controls
        self.auto_refresh = QCheckBox("Auto-refresh (every 30 seconds)")
        self.auto_refresh.setChecked(True)
        self.refresh_btn = QPushButton("Refresh Now")
        self.export_btn = QPushButton("Export Word Cloud")

        # Layout
        control_layout = QHBoxLayout()
        control_layout.addWidget(self.auto_refresh)
        control_layout.addStretch()
        control_layout.addWidget(self.refresh_btn)
        control_layout.addWidget(self.export_btn)

        layout = QVBoxLayout()
        layout.addWidget(self.label, 1) # Label takes most space
        layout.addLayout(control_layout)
        self.setLayout(layout)

        # Timer for auto-refresh
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.auto_refresh_wordcloud)
        self.timer.start(30000) # 30 seconds

        # Connect button signals
        self.refresh_btn.clicked.connect(self.request_refresh)
        self.export_btn.clicked.connect(self.request_export)

    def update_wordcloud(self, text):
        """Update the word cloud with new text."""
        if not text:
            self.label.setText("No transcript data available.")
            return
        image_data = generate_wordcloud(text)
        if image_data:
            pixmap = QPixmap()
            pixmap.loadFromData(image_data)
            self.label.setPixmap(pixmap.scaled(self.label.size() * 0.95, Qt.KeepAspectRatio, Qt.SmoothTransformation)) # Scale slightly smaller than label
        else:
            self.label.setText("Error generating word cloud.")


    def auto_refresh_wordcloud(self):
        """Called by timer to refresh the word cloud."""
        if self.auto_refresh.isChecked() and self.isVisible() and self.parent_app:
            self.parent_app.refresh_wordcloud() # Call parent method

    def request_refresh(self):
         """Request parent app to refresh the word cloud."""
         if self.parent_app:
             self.parent_app.refresh_wordcloud()

    def request_export(self):
         """Request parent app to handle word cloud export."""
         if self.parent_app:
             self.parent_app.export_wordcloud()

    def get_pixmap(self):
        """Returns the current pixmap displayed."""
        return self.label.pixmap()

# --------------------------
# Settings Dialog (from updated)
# --------------------------
class SettingsDialog(QDialog):
    def __init__(self, parent=None, current_settings=None):
        super().__init__(parent)
        self.settings = current_settings.copy() if current_settings else DEFAULT_SETTINGS.copy() # Work on a copy
        self.setWindowTitle("Settings")
        self.setMinimumWidth(650)
        self.setMinimumHeight(550)

        # Apply styling based on current theme
        self.apply_dialog_theme()

        self.setup_ui()
        self.load_settings_into_dialog()

    def apply_dialog_theme(self):
        theme_name = self.settings.get("theme", "macOS")
        if theme_name not in THEMES: theme_name = "macOS"
        theme = THEMES[theme_name]

        is_dark = theme_name in ["dark", "navy", "dracula", "macOSDark"]
        base_bg = theme["window"]
        text_color = theme["windowText"]
        button_bg = theme["button"]
        button_text = theme["buttonText"]
        border_color = theme["borderColor"]
        highlight_bg = theme["highlight"]
        highlight_text = theme["highlightedText"]
        base_text_edit_bg = theme["base"] if is_dark else "#FFFFFF"
        control_bg = theme["alternateBase"] if is_dark else "#F7F7F7"

        # Remove the 'f' prefix here:
        self.setStyleSheet("""
          QDialog {{
            background-color: {base_bg};
            color: {text_color};
          }}
          QGroupBox {{
            font-weight: bold;
            border: 1px solid {border_color};
            border-radius: 6px;
            margin-top: 12px;
            padding-top: 15px;
            background-color: rgba(0,0,0, 0.02);
          }}
          QGroupBox::title {{
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 5px 0 5px;
            color: {text_color};
          }}
          QLabel {{ color: {text_color}; }}
          QLineEdit, QComboBox, QSpinBox {{
            border: 1px solid {border_color};
            border-radius: 4px;
            padding: 6px 10px;
            background-color: {control_bg};
            color: {text_color};
          }}
          QLineEdit:focus, QComboBox:focus, QSpinBox:focus {{
            border: 1px solid {highlight_bg};
          }}
          QComboBox::drop-down {{
            border-left: 1px solid {border_color};
            border-top-right-radius: 4px;
            border-bottom-right-radius: 4px;
          }}
          QTextEdit {{
            border: 1px solid {border_color};
            border-radius: 4px;
            padding: 8px;
            background-color: {base_text_edit_bg};
            color: {text_color};
            font-family: "Menlo", monospace;
          }}
          QCheckBox {{ spacing: 8px; color: {text_color}; }}
          QCheckBox::indicator {{
            width: 18px; height: 18px;
            border: 1px solid {border_color};
            border-radius: 4px;
            background-color: {control_bg};
          }}
          QCheckBox::indicator:checked {{
            background-color: {highlight_bg};
            border: 1px solid {highlight_bg};
            /* Use a placeholder {svg_fill_color} for the SVG fill */
            image: url(data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxOCIgaGVpZ2h0PSIxOCIgdmlld0JveD0iMCAwIDE4IDE4Ij48cGF0aCBmaWxsPSJ{svg_fill_color}\" d=\"M6.95 12.76L3.11 8.92L2 10L6.95 14.95L16.18 5.73L15.04 4.59L6.95 12.76z\"/></svg>);
          }}
          QTabWidget::pane {{
            border: 1px solid {border_color};
            border-radius: 6px;
            background-color: {base_bg};
            padding: 10px;
          }}
          QTabBar::tab {{
            background-color: {button_bg};
            border: 1px solid {border_color};
            border-bottom: none;
            border-top-left-radius: 6px;
            border-top-right-radius: 6px;
            padding: 8px 16px;
            margin-right: 2px;
            color: {button_text};
            font-weight: 500;
          }}
          QTabBar::tab:selected {{
            background: {base_bg};
            border-bottom-color: {base_bg};
            font-weight: bold;
          }}
          QPushButton {{
            background-color: {button_bg};
            border: 1px solid {border_color};
            border-radius: 6px;
            padding: 8px 15px;
            font-weight: 500;
            color: {button_text};
          }}
          /* Use .format() for parts needing variables like hover/pressed states */
          QPushButton:hover {{ background-color: {alternateBase}; }}
          QPushButton:pressed {{ background-color: {highlight_bg}; color: {highlight_text}; }}
          QPushButton#saveButton {{
            background-color: {highlight_bg};
            color: {highlight_text};
            font-weight: bold;
          }}
          /* These hover/pressed need careful formatting if using QColor methods */
          QPushButton#saveButton:hover {{ background-color: {save_button_hover_color}; }}
          QPushButton#saveButton:pressed {{ background-color: {save_button_pressed_color}; }}
          QPushButton#deleteButton {{ background-color: {delete_button_color}; color: white; }}
        """.format(
          # Pass all the theme variables needed in the template here
          base_bg=base_bg,
          text_color=text_color,
          border_color=border_color,
          control_bg=control_bg,
          highlight_bg=highlight_bg,
          base_text_edit_bg=base_text_edit_bg,
          svg_fill_color=theme['highlightedText'].replace('#', '%23'), # URL encode #
          button_bg=button_bg,
          button_text=button_text,
          alternateBase=theme["alternateBase"],
          highlight_text=theme["highlightedText"],
          # Calculate hover/pressed colors *before* formatting
          save_button_hover_color=QColor(highlight_bg).lighter(110).name(),
          save_button_pressed_color=QColor(highlight_bg).darker(110).name(),
          delete_button_color=theme["brightText"]
          # Add any other variables used in the CSS string above
        ))

        self.transcript_lock = asyncio.Lock()
        self.poll_task = asyncio.create_task(self.poll_transcript())

    async def poll_transcript(self):
        """Async polling mechanism for transcript updates"""
        while True:
            await asyncio.sleep(settings["poll_interval"])
            await self.check_transcript()

        layout = QVBoxLayout()
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        tabs = QTabWidget()
        # tabs.setStyleSheet is now handled by apply_dialog_theme

        # --- General Settings Tab ---
        general_tab = QWidget()
        general_layout = QVBoxLayout(general_tab) # Use QVBoxLayout for grouping
        general_layout.setSpacing(15)
        general_layout.setContentsMargins(0, 0, 0, 0) # No margins on inner layout

        # GroupBoxes for better organization
        theme_group = QGroupBox("Appearance")
        theme_layout = QFormLayout(theme_group)
        self.theme_combo = QComboBox()
        self.theme_combo.addItems(list(THEMES.keys()))
        theme_layout.addRow("Theme:", self.theme_combo)
        self.rounded_corners_check = QCheckBox("Use rounded corners")
        theme_layout.addRow(self.rounded_corners_check)
        self.animations_check = QCheckBox("Enable animations (requires restart)")
        theme_layout.addRow(self.animations_check)
        self.transparency_check = QCheckBox("Enable transparency effects (requires restart)")
        theme_layout.addRow(self.transparency_check)
        general_layout.addWidget(theme_group)

        startup_group = QGroupBox("Startup")
        startup_layout = QFormLayout(startup_group)
        self.start_bg_check = QCheckBox("Start application in background (system tray)")
        startup_layout.addRow(self.start_bg_check)
        general_layout.addWidget(startup_group)

        shortcuts_group = QGroupBox("Keyboard Shortcuts")
        shortcuts_layout = QFormLayout(shortcuts_group)
        self.shortcut_edit = QLineEdit()
        shortcuts_layout.addRow("Analysis shortcut:", self.shortcut_edit)
        self.code_word_edit = QLineEdit()
        shortcuts_layout.addRow("Trigger code word:", self.code_word_edit)
        general_layout.addWidget(shortcuts_group)

        font_group = QGroupBox("Font Settings")
        font_layout = QFormLayout(font_group)
        self.font_family_combo = QComboBox()
        self.font_family_combo.addItems(["SF Pro Text", "SF Pro Display", "Arial", "Verdana", "Roboto", "Times New Roman", "Helvetica Neue", "Courier New", "Consolas"])
        font_layout.addRow("Font family:", self.font_family_combo)
        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(8, 24)
        font_layout.addRow("Font size:", self.font_size_spin)
        self.line_spacing_spin = QSpinBox()
        self.line_spacing_spin.setRange(100, 300); self.line_spacing_spin.setSingleStep(10)
        font_layout.addRow("Line spacing (%):", self.line_spacing_spin)
        self.word_spacing_spin = QSpinBox() # Note: Word spacing isn't easily applied per-widget, maybe remove?
        self.word_spacing_spin.setRange(80, 200); self.word_spacing_spin.setSingleStep(5)
        # font_layout.addRow("Word spacing (%):", self.word_spacing_spin) # Commented out for now
        general_layout.addWidget(font_group)

        icon_group = QGroupBox("Application Icons (Requires Restart)")
        icon_layout = QFormLayout(icon_group)
        icon_fields = [("App Icon:", "app_icon_edit"), ("Generate Icon:", "generate_icon_edit"), ("Settings Icon:", "settings_icon_edit")]
        for label_text, field_name in icon_fields:
            setattr(self, field_name, QLineEdit())
            line_edit = getattr(self, field_name)
            browse_btn = QPushButton("Browse...")
            browse_btn.clicked.connect(lambda checked=False, le=line_edit: self.browse_icon(le))
            field_layout = QHBoxLayout(); field_layout.addWidget(line_edit); field_layout.addWidget(browse_btn)
            icon_layout.addRow(label_text, field_layout)
        general_layout.addWidget(icon_group)
        general_layout.addStretch() # Push content up

        # --- AI Settings Tab ---
        ai_tab = QWidget()
        ai_layout = QVBoxLayout(ai_tab)
        ai_layout.setSpacing(15)
        ai_layout.setContentsMargins(0,0,0,0)

        model_card = QGroupBox("Groq LLM Model")
        model_layout = QFormLayout(model_card)
        self.model_combo = QComboBox()
        self.model_combo.addItems(AVAILABLE_MODELS)
        model_layout.addRow("Select Model:", self.model_combo)
        ai_layout.addWidget(model_card)

        interval_card = QGroupBox("Analysis Frequency")
        interval_layout = QFormLayout(interval_card)
        self.line_interval_spin = QSpinBox()
        self.line_interval_spin.setRange(3, 50)
        interval_layout.addRow("Summarize every N lines:", self.line_interval_spin)
        self.time_interval_spin = QSpinBox()
        self.time_interval_spin.setRange(30, 600); self.time_interval_spin.setSingleStep(30)
        interval_layout.addRow("Summarize every N seconds:", self.time_interval_spin)
        ai_layout.addWidget(interval_card)

        prompt_card = QGroupBox("Prompt Templates")
        prompt_layout = QVBoxLayout(prompt_card)
        prompt_form = QFormLayout()
        self.prompt_combo = QComboBox()
        prompt_form.addRow("Active Prompt:", self.prompt_combo)
        prompt_layout.addLayout(prompt_form)
        self.prompt_edit = QTextEdit()
        self.prompt_edit.setMinimumHeight(150)
        prompt_layout.addWidget(self.prompt_edit)
        prompt_btn_layout = QHBoxLayout()
        add_prompt_btn = QPushButton("Add New"); save_prompt_btn = QPushButton("Save Changes"); delete_prompt_btn = QPushButton("Delete");
        delete_prompt_btn.setObjectName("deleteButton") # ID for styling
        prompt_btn_layout.addWidget(add_prompt_btn); prompt_btn_layout.addWidget(save_prompt_btn); prompt_btn_layout.addWidget(delete_prompt_btn);
        prompt_layout.addLayout(prompt_btn_layout)
        add_prompt_btn.clicked.connect(self.add_new_prompt)
        save_prompt_btn.clicked.connect(self.save_prompt_changes)
        delete_prompt_btn.clicked.connect(self.delete_prompt)
        self.prompt_combo.currentTextChanged.connect(self.update_prompt_edit)
        ai_layout.addWidget(prompt_card)
        ai_layout.addStretch()

        # --- Advanced Settings Tab ---
        adv_tab = QWidget()
        adv_layout = QVBoxLayout(adv_tab)
        adv_layout.setSpacing(15)
        adv_layout.setContentsMargins(0,0,0,0)

        polling_card = QGroupBox("Transcript Polling")
        polling_layout = QFormLayout(polling_card)
        self.poll_interval_spin = QSpinBox()
        self.poll_interval_spin.setRange(1, 30)
        polling_layout.addRow("Polling interval (seconds):", self.poll_interval_spin)
        polling_layout.addRow(QLabel("Note: Lower values increase responsiveness but may use more resources."))
        adv_layout.addWidget(polling_card)

        transcript_card = QGroupBox("Transcript Features")
        transcript_layout = QVBoxLayout(transcript_card) # Use QVBoxLayout for checkboxes
        self.anonymize_check = QCheckBox("Anonymize names in transcript")
        transcript_layout.addWidget(self.anonymize_check)
        transcript_layout.addWidget(QLabel("Replaces detected names with generic identifiers (Person-1, Person-2, etc.)", wordWrap=True))
        self.spellcheck_check = QCheckBox("Enable basic spell checking")
        transcript_layout.addWidget(self.spellcheck_check)
        transcript_layout.addWidget(QLabel("Automatically corrects common spelling errors", wordWrap=True))
        adv_layout.addWidget(transcript_card)

        timer_card = QGroupBox("Meeting Timer")
        timer_layout = QVBoxLayout(timer_card) # Use QVBoxLayout
        self.timer_check = QCheckBox("Enable meeting timer")
        timer_layout.addWidget(self.timer_check)
        timer_form = QFormLayout() # Use FormLayout inside for label/spinbox
        self.reminder_spin = QSpinBox()
        self.reminder_spin.setRange(5, 60); self.reminder_spin.setSingleStep(5)
        timer_form.addRow("Reminder interval (minutes):", self.reminder_spin)
        timer_layout.addLayout(timer_form)
        timer_layout.addWidget(QLabel("Notifications remind you how long the meeting has run.", wordWrap=True))
        adv_layout.addWidget(timer_card)
        adv_layout.addStretch()

        # Add tabs
        tabs.addTab(general_tab, "General")
        tabs.addTab(ai_tab, "AI Settings")
        tabs.addTab(adv_tab, "Advanced")
        layout.addWidget(tabs)

        # Buttons
        button_layout = QHBoxLayout()
        reset_btn = QPushButton("Reset to Defaults")
        save_btn = QPushButton("Save Settings")
        save_btn.setObjectName("saveButton") # ID for specific styling
        cancel_btn = QPushButton("Cancel")
        reset_btn.clicked.connect(self.reset_to_defaults)
        save_btn.clicked.connect(self.save_settings_and_accept) # Changed action
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(reset_btn); button_layout.addStretch(); button_layout.addWidget(save_btn); button_layout.addWidget(cancel_btn);
        layout.addLayout(button_layout)

        self.setLayout(layout)

    def browse_icon(self, line_edit):
        """Open file dialog to choose an icon file."""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Icon", "", "Image Files (*.png *.jpg *.jpeg *.ico *.svg)")
        if file_path: line_edit.setText(file_path)

    def add_new_prompt(self):
        """Add a new prompt template."""
        prompt_name, ok = QInputDialog.getText(self, "New Prompt", "Enter prompt name:")
        if ok and prompt_name:
            if prompt_name in [self.prompt_combo.itemText(i) for i in range(self.prompt_combo.count())]:
                QMessageBox.warning(self, "Duplicate Name", "A prompt with this name already exists.")
                return
            self.prompt_combo.addItem(prompt_name)
            self.prompt_combo.setCurrentText(prompt_name)
            self.prompt_edit.setText("# Enter your new prompt template here...\n")
            self.prompt_edit.selectAll(); self.prompt_edit.setFocus()

    def save_prompt_changes(self):
        """Save changes to the current prompt in the internal settings dict."""
        current_name = self.prompt_combo.currentText()
        if not current_name: return False # Indicate failure
        prompts = self.settings.get("prompts", [])
        prompt_text = self.prompt_edit.toPlainText()
        found = False
        for i, prompt in enumerate(prompts):
            if prompt["name"] == current_name:
                prompts[i]["prompt"] = prompt_text; found = True; break
        if not found: prompts.append({"name": current_name, "prompt": prompt_text})
        self.settings["prompts"] = prompts
        print(f"Prompt '{current_name}' saved internally.")
        return True # Indicate success

    def delete_prompt(self):
        """Delete the current prompt."""
        current_name = self.prompt_combo.currentText()
        if not current_name or current_name == DEFAULT_SETTINGS["prompts"][0]["name"]: # Prevent deleting default
            QMessageBox.warning(self, "Cannot Delete", "Cannot delete the default prompt.")
            return
        confirm = QMessageBox.question(self, "Confirm Deletion", f"Delete prompt '{current_name}'?", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if confirm != QMessageBox.Yes: return
        self.settings["prompts"] = [p for p in self.settings.get("prompts", []) if p["name"] != current_name]
        index = self.prompt_combo.findText(current_name)
        if index >= 0: self.prompt_combo.removeItem(index)
        # Select the default prompt if the deleted one was active
        if self.settings.get("active_prompt") == current_name:
            self.settings["active_prompt"] = DEFAULT_SETTINGS["prompts"][0]["name"]
            self.prompt_combo.setCurrentText(self.settings["active_prompt"])
        self.update_prompt_edit(self.prompt_combo.currentText()) # Refresh editor

    def update_prompt_edit(self, prompt_name):
        """Update the prompt edit box with the selected prompt."""
        if not prompt_name: self.prompt_edit.clear(); return
        prompts = self.settings.get("prompts", [])
        for prompt in prompts:
            if prompt["name"] == prompt_name: self.prompt_edit.setText(prompt["prompt"]); return
        self.prompt_edit.clear() # Clear if not found

    def load_settings_into_dialog(self):
        """Load current settings into the dialog widgets."""
        self.theme_combo.setCurrentText(self.settings.get("theme", "macOS"))
        self.start_bg_check.setChecked(self.settings.get("start_in_background", False))
        self.shortcut_edit.setText(self.settings.get("keyboard_shortcut", "ctrl+i"))
        self.code_word_edit.setText(self.settings.get("code_word", DEFAULT_CODE_WORD))
        self.rounded_corners_check.setChecked(self.settings.get("use_rounded_corners", True))
        self.animations_check.setChecked(self.settings.get("use_animations", True))
        self.transparency_check.setChecked(self.settings.get("use_transparency", True))
        self.font_family_combo.setCurrentText(self.settings.get("font_family", DEFAULT_FONT_FAMILY))
        self.font_size_spin.setValue(self.settings.get("font_size", DEFAULT_FONT_SIZE))
        self.line_spacing_spin.setValue(self.settings.get("line_spacing", DEFAULT_LINE_SPACING))
        # self.word_spacing_spin.setValue(int(self.settings.get("word_spacing", DEFAULT_WORD_SPACING) * 100))
        self.app_icon_edit.setText(self.settings.get("app_icon_path", ""))
        self.generate_icon_edit.setText(self.settings.get("generate_icon_path", ""))
        self.settings_icon_edit.setText(self.settings.get("settings_icon_path", ""))
        self.model_combo.setCurrentText(self.settings.get("model", DEFAULT_MODEL))
        self.line_interval_spin.setValue(self.settings.get("summary_line_interval", 10))
        self.time_interval_spin.setValue(self.settings.get("summary_time_interval", 120))
        self.prompt_combo.clear()
        prompts = self.settings.get("prompts", [])
        if not any(p['name'] == DEFAULT_SETTINGS['prompts'][0]['name'] for p in prompts):
             prompts.insert(0, DEFAULT_SETTINGS['prompts'][0]) # Ensure default exists
        self.prompt_combo.addItems([p["name"] for p in prompts])
        active_prompt = self.settings.get("active_prompt", DEFAULT_SETTINGS["prompts"][0]["name"])
        index = self.prompt_combo.findText(active_prompt)
        self.prompt_combo.setCurrentIndex(index if index >= 0 else 0)
        self.update_prompt_edit(self.prompt_combo.currentText())
        self.poll_interval_spin.setValue(self.settings.get("poll_interval", DEFAULT_POLL_INTERVAL))
        self.anonymize_check.setChecked(self.settings.get("anonymize_transcript", False))
        self.spellcheck_check.setChecked(self.settings.get("spellcheck_enabled", False))
        self.timer_check.setChecked(self.settings.get("meeting_timer_enabled", True))
        self.reminder_spin.setValue(self.settings.get("meeting_reminder_interval", 15))

    def reset_to_defaults(self):
        """Reset all settings to defaults and reload dialog."""
        confirm = QMessageBox.question(self, "Confirm Reset", "Reset all settings to defaults?", QMessageBox.Yes | QMessageBox.No)
        if confirm != QMessageBox.Yes: return
        # Keep custom prompts but reset everything else
        custom_prompts = [p for p in self.settings.get("prompts", []) if p['name'] != DEFAULT_SETTINGS['prompts'][0]['name']]
        self.settings = DEFAULT_SETTINGS.copy()
        self.settings["prompts"].extend(custom_prompts) # Add custom prompts back
        self.settings["active_prompt"] = DEFAULT_SETTINGS["prompts"][0]["name"] # Reset active prompt
        self.load_settings_into_dialog()
        QMessageBox.information(self, "Reset Complete", "Settings reset to defaults.")

    def save_settings_and_accept(self):
        """Save settings from dialog to internal dict and accept."""
        # Save current prompt changes first
        if not self.save_prompt_changes():
             QMessageBox.warning(self, "Save Error", "Could not save changes to the current prompt.")
             return # Don't proceed if prompt save failed

        # General
        self.settings["theme"] = self.theme_combo.currentText()
        self.settings["start_in_background"] = self.start_bg_check.isChecked()
        self.settings["keyboard_shortcut"] = self.shortcut_edit.text()
        self.settings["code_word"] = self.code_word_edit.text()
        # Visual
        self.settings["use_rounded_corners"] = self.rounded_corners_check.isChecked()
        self.settings["use_animations"] = self.animations_check.isChecked()
        self.settings["use_transparency"] = self.transparency_check.isChecked()
        # Font
        self.settings["font_family"] = self.font_family_combo.currentText()
        self.settings["font_size"] = self.font_size_spin.value()
        self.settings["line_spacing"] = self.line_spacing_spin.value()
        # self.settings["word_spacing"] = self.word_spacing_spin.value() / 100.0
        # Icons
        self.settings["app_icon_path"] = self.app_icon_edit.text()
        self.settings["generate_icon_path"] = self.generate_icon_edit.text()
        self.settings["settings_icon_path"] = self.settings_icon_edit.text()
        # AI
        self.settings["model"] = self.model_combo.currentText()
        self.settings["summary_line_interval"] = self.line_interval_spin.value()
        self.settings["summary_time_interval"] = self.time_interval_spin.value()
        self.settings["active_prompt"] = self.prompt_combo.currentText() # Already saved
        # Advanced
        self.settings["poll_interval"] = self.poll_interval_spin.value()
        self.settings["anonymize_transcript"] = self.anonymize_check.isChecked()
        self.settings["spellcheck_enabled"] = self.spellcheck_check.isChecked()
        self.settings["meeting_timer_enabled"] = self.timer_check.isChecked()
        self.settings["meeting_reminder_interval"] = self.reminder_spin.value()

        self.accept() # Close the dialog successfully

    def get_updated_settings(self):
        """Return the settings dictionary."""
        return self.settings

# --------------------------
# Main Application Class (Merged)
# --------------------------
class MeetingAssistant(QWidget):
    def __init__(self):
        super().__init__()
        self.setObjectName("mainWindow")
        self.setWindowTitle(f"{APP_NAME} v{VERSION}")
        self.setGeometry(100, 100, 900, 700) # Size from updated

        # Load settings initially
        global settings
        settings = load_settings()

        # Load icons based on settings
        self.app_icon = load_icon_from_path_or_default(settings.get("app_icon_path"), DEFAULT_APP_ICON)
        self.generate_icon = load_icon_from_path_or_default(settings.get("generate_icon_path"), DEFAULT_GENERATE_ICON)
        self.settings_icon = load_icon_from_path_or_default(settings.get("settings_icon_path"), DEFAULT_SETTINGS_ICON)
        self.setWindowIcon(self.app_icon)

        # State variables
        self.running = False
        self.browser = None
        self.page = None
        self.buffer = []
        self.seen_lines = set() # Using set for efficiency
        self.summaries = [] # Keep summaries list from fixed3
        self.suggested_responses = [] # From updated
        self.insights = [] # From updated
        self.known_names = [] # From updated
        self.meeting_start_time = None # From updated
        self.last_reminder_time = None # From updated
        self.code_word_detected = False # From updated
        self.meeting_title = "Unknown Meeting" # From updated
        self.key_phrases = [] # From updated

        # Child windows
        self.wordcloud_window = WordCloudWindow(self)
        self.settings_dialog = None # Create on demand

        # Timers
        self.main_timer = QTimer() # Use main_timer instead of timer from fixed3
        self.elapsed_seconds = 0
        self.progress_timer = QTimer() # from updated
        self.meeting_timer = QTimer() # from updated

        # Signal bridge
        self.signals = SignalBridge()

        # Build UI using the updated structure
        self.build_ui()

        # Connect signals
        self.signals.append_transcript.connect(self.append_transcript)
        print("Connected append_transcript signal")
        self.signals.append_suggested_response.connect(self.append_suggested_response) # New connection
        self.signals.append_insights.connect(self.append_insights) # New connection (replaces append_summary)
        self.signals.update_status.connect(self.update_status)
        self.signals.update_progress.connect(self.update_progress)
        self.signals.highlight_transcript.connect(self.highlight_transcript)
        self.signals.update_wordcloud.connect(self.wordcloud_window.update_wordcloud) # Connect directly
        self.signals.update_meeting_timer.connect(self.update_meeting_timer_label) # Renamed method
        self.signals.detected_key_phrases.connect(self.update_key_phrases)
        self.signals.detected_meeting_title.connect(self.update_meeting_title)

        # Apply theme, tray, hotkey
        self.apply_theme()
        self.setup_tray()
        self.setup_hotkey()

        # Start hidden if configured
        if settings.get("start_in_background"):
            QTimer.singleShot(100, self.hide) # Delay hiding slightly

    def build_ui(self): # Using build_ui from updated
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(12)
        main_layout.setContentsMargins(16, 16, 16, 16)

        # --- Top Bar ---
        top_bar = QHBoxLayout()
        top_bar.addWidget(QLabel("Meeting URL:"))
        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText("Enter Otter.ai meeting URL")
        self.url_input.setMinimumHeight(32)
        url_container = QFrame() # Container for styling
        url_container.setObjectName("url_container")
        url_layout = QHBoxLayout(url_container); url_layout.setContentsMargins(0,0,0,0); url_layout.addWidget(self.url_input)
        top_bar.addWidget(url_container, 2) # Takes 2/3 space
        top_bar.addSpacerItem(QSpacerItem(10,10, QSizePolicy.Fixed, QSizePolicy.Minimum))
        self.generate_btn = create_icon_button(self.generate_icon, f"Generate AI Analysis Now ({settings.get('keyboard_shortcut', 'Ctrl+I')})")
        self.settings_btn = create_icon_button(self.settings_icon, "Settings")
        self.generate_btn.clicked.connect(self.on_demand_analysis)
        self.settings_btn.clicked.connect(self.open_settings)
        top_bar.addWidget(self.generate_btn); top_bar.addWidget(self.settings_btn)
        main_layout.addLayout(top_bar)

        # --- Status Bar ---
        status_container = QFrame(); status_container.setFrameShape(QFrame.StyledPanel)
        status_layout = QHBoxLayout(status_container); status_layout.setContentsMargins(12, 8, 12, 8)
        self.status = QLabel("Status: Ready"); self.status.setStyleSheet("font-weight: 500;")
        self.progress_bar = QProgressBar(); self.progress_bar.setRange(0, 100); self.progress_bar.setValue(0); self.progress_bar.setVisible(False); self.progress_bar.setMaximumWidth(150)
        self.timer_label = QLabel("Meeting time: 00:00:00"); self.timer_label.setStyleSheet("font-weight: 500;")
        status_layout.addWidget(self.status); status_layout.addStretch(); status_layout.addWidget(self.progress_bar); status_layout.addWidget(self.timer_label)
        main_layout.addWidget(status_container)

        # --- Action Buttons ---
        btn_container = QFrame()
        btn_layout = QHBoxLayout(btn_container); btn_layout.setSpacing(16); btn_layout.setContentsMargins(12, 8, 12, 8)
        
        # Define icons (embedded or from theme)
        play_icon = QIcon.fromTheme("media-playback-start") 
        pause_icon = QIcon.fromTheme("media-playback-stop")
        save_icon = QIcon.fromTheme("document-save")
        cloud_icon = QIcon.fromTheme("view-grid")
        
        # Create icon buttons with text
        self.start_btn = self.create_action_button("Start", play_icon, "#4CAF50", tooltip="Start monitoring transcript")
        self.stop_btn = self.create_action_button("Stop", pause_icon, "#F44336", tooltip="Stop monitoring")
        self.save_btn = self.create_action_button("Save", save_icon, "#2196F3", tooltip="Save transcript and insights")
        self.cloud_btn = self.create_action_button("Word Cloud", cloud_icon, "#9C27B0", tooltip="Generate word cloud visualization")
        
        # Connect signals
        self.start_btn.clicked.connect(self.start)
        self.stop_btn.clicked.connect(self.stop)
        self.save_btn.clicked.connect(self.save)
        self.cloud_btn.clicked.connect(self.show_wordcloud)
        
        self.stop_btn.setEnabled(False) # Initially disabled
        
        # Add buttons to layout
        btn_layout.addWidget(self.start_btn)
        btn_layout.addWidget(self.stop_btn)
        btn_layout.addWidget(self.save_btn)
        btn_layout.addWidget(self.cloud_btn)
        
        main_layout.addWidget(btn_container)

        # --- Model/Prompt Selector ---
        model_container = QFrame(); model_container.setFrameShape(QFrame.StyledPanel)
        model_layout = QHBoxLayout(model_container); model_layout.setContentsMargins(12, 8, 12, 8)
        model_layout.addWidget(QLabel("AI Model:")); self.model_selector = QComboBox(); self.model_selector.addItems(AVAILABLE_MODELS); self.model_selector.setCurrentText(settings.get("model", DEFAULT_MODEL)); self.model_selector.setMinimumWidth(150); model_layout.addWidget(self.model_selector); model_layout.addSpacing(20)
        model_layout.addWidget(QLabel("Prompt:")); self.prompt_selector = QComboBox(); self.prompt_selector.addItems([p["name"] for p in settings.get("prompts", [])]); self.prompt_selector.setCurrentText(settings.get("active_prompt", "")); self.prompt_selector.currentTextChanged.connect(self.change_prompt); self.prompt_selector.setMinimumWidth(200); model_layout.addWidget(self.prompt_selector);
        main_layout.addWidget(model_container)

        # --- Main Content Area ---
        content_layout = QVBoxLayout(); content_layout.setSpacing(10)

        # Meeting Topic Header
        meeting_header = QFrame(); meeting_header.setFrameShape(QFrame.StyledPanel)
        meeting_layout = QHBoxLayout(meeting_header); meeting_layout.setContentsMargins(12, 8, 12, 8)
        topic_label = QLabel("<b>Meeting Topic:</b>"); topic_label.setStyleSheet("font-weight: 600;")
        self.topic_value = QLabel("Not detected yet"); self.topic_value.setStyleSheet("font-style: italic;")
        meeting_layout.addWidget(topic_label); meeting_layout.addWidget(self.topic_value); meeting_layout.addStretch()
        content_layout.addWidget(meeting_header)

        # Groq Output Panel (Splitter)
        self.groq_panel = QSplitter(Qt.Horizontal)
        # Left: Suggested Responses
        suggested_frame = self.create_styled_frame("<b>Suggested Responses</b>")
        self.suggested_response_box = CustomTextEdit()
        suggested_layout = suggested_frame.layout(); suggested_layout.addWidget(self.suggested_response_box)
        self.groq_panel.addWidget(suggested_frame)
        # Right: Insights
        insights_frame = self.create_styled_frame("<b>Key Insights & Analysis</b>")
        self.insights_box = CustomTextEdit() # Replaces fixed3's summary_box
        insights_layout = insights_frame.layout(); insights_layout.addWidget(self.insights_box)
        self.groq_panel.addWidget(insights_frame)
        self.groq_panel.setSizes([int(self.width() * 0.5), int(self.width() * 0.5)]) # Initial split
        content_layout.addWidget(self.groq_panel, 1) # Stretch factor 1

        # Transcript Panel
        transcript_frame = self.create_styled_frame("<b>Live Transcript</b>", add_search=True) # Add search to header
        self.transcript_box = CustomTextEdit()
        transcript_layout = transcript_frame.layout(); transcript_layout.addWidget(self.transcript_box)
        content_layout.addWidget(transcript_frame, 1) # Stretch factor 1

        main_layout.addLayout(content_layout)

    def create_action_button(self, text, icon, color, tooltip=None):
        """Create a modern button with both icon and text."""
        button = QPushButton()
        
        # Create a layout for the button content
        content_layout = QHBoxLayout(button)
        content_layout.setContentsMargins(10, 6, 10, 6)
        content_layout.setSpacing(8)
        
        # Add icon
        if icon and not icon.isNull():
            icon_label = QLabel()
            icon_label.setPixmap(icon.pixmap(24, 24))
            content_layout.addWidget(icon_label)
        
        # Add text
        text_label = QLabel(text)
        text_label.setStyleSheet("font-weight: 600; color: white; font-size: 13px;")
        content_layout.addWidget(text_label)
        
        # Set button properties
        if tooltip:
            button.setToolTip(tooltip)
        button.setMinimumHeight(40)
        
        # Apply styling
        lighter_color = QColor(color).lighter(110).name()
        darker_color = QColor(color).darker(110).name()
        button.setStyleSheet(f"""
            QPushButton {{
                background-color: {color};
                border: none;
                border-radius: 6px;
                padding: 8px 15px;
            }}
            QPushButton:hover {{
                background-color: {lighter_color};
            }}
            QPushButton:pressed {{
                background-color: {darker_color};
                padding-top: 9px;
                padding-bottom: 7px;
            }}
            QPushButton:disabled {{
                background-color: #555;
                color: #aaa;
            }}
        """)
        
        return button

    def create_styled_frame(self, title_html, add_search=False):
        """Helper to create styled frames with headers."""
        frame = QFrame(); frame.setFrameShape(QFrame.StyledPanel)
        frame.setObjectName("styledFrame") # For specific styling if needed

        layout = QVBoxLayout(frame); layout.setContentsMargins(0, 0, 0, 0); layout.setSpacing(0)

        header = QFrame(); header.setObjectName("styledHeader")
        header_layout = QHBoxLayout(header); header_layout.setContentsMargins(10, 5, 10, 5)
        title_label = QLabel(title_html); title_label.setStyleSheet("font-weight: 600; font-size: 14px;")
        header_layout.addWidget(title_label)

        if add_search:
             header_layout.addStretch()
             search_container = QFrame(); search_container.setObjectName("searchContainer")
             search_layout = QHBoxLayout(search_container); search_layout.setContentsMargins(8, 0, 8, 0); search_layout.setSpacing(5)
             self.search_box = QLineEdit(); self.search_box.setPlaceholderText("Search transcript..."); self.search_box.setMaximumWidth(250); self.search_box.setMinimumHeight(28)
             self.search_box.textChanged.connect(self.search_transcript_text) # Renamed method
             search_icon = QLabel(); search_icon.setPixmap(QIcon.fromTheme("edit-find").pixmap(16, 16))
             search_layout.addWidget(search_icon); search_layout.addWidget(self.search_box)
             header_layout.addWidget(search_container)

        layout.addWidget(header)
        # The actual QTextEdit will be added by the caller
        return frame

    # Corrected get_button_style - PASTE THIS into your script
    def get_button_style(self, button_type):
      """Returns specific stylesheets for buttons based on theme."""
      # Define common properties WITHOUT the selector and braces
      common_properties = "border: none; border-radius: 6px; padding: 8px 20px; font-weight: 600; color: white;"
      # Define disabled style separately
      disabled_style = "QPushButton:disabled { background-color: #555; color: #aaa; }"
      
      colors = {
        "start": ("#4CAF50", "#43A047"), # Green
        "stop": ("#F44336", "#E53935"),  # Red
        "save": ("#2196F3", "#1E88E5"),  # Blue
        "cloud": ("#9C27B0", "#8E24AA")  # Purple
      }
      start_color, end_color = colors.get(button_type, ("#5E5E5E", "#444444")) # Default gray
      
      gradient = f"qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 {start_color}, stop:1 {end_color})"
      # Use try-except for QColor in case a theme value is invalid or QColor is not available temporarily
      try:
        # Ensure QColor is imported if not already (should be from PyQt5.QtGui)
        hover_start = QColor(start_color).lighter(110).name()
        hover_end = QColor(end_color).lighter(110).name()
        pressed_start = QColor(start_color).darker(110).name()
        pressed_end = QColor(end_color).darker(110).name()
      except Exception as e: # Fallback if color conversion fails
        print(f"Warning: Invalid color for button '{button_type}': {start_color}/{end_color}. Using defaults. Error: {e}")
        hover_start, hover_end = "#656565", "#4F4F4F" # Default hover
        pressed_start, pressed_end = "#444444", "#5E5E5E" # Default pressed
      
      # Construct the final stylesheet using the corrected structure
      return f"""
        QPushButton {{
          background-color: {gradient};
          {common_properties}
        }}
        QPushButton:hover {{
          background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 {hover_start}, stop:1 {hover_end});
          /* filter: brightness(110%); /* Optional: filter might cause issues, removed for safety */ */
        }}
        QPushButton:pressed {{
          background-color: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 {pressed_start}, stop:1 {pressed_end});
          padding-top: 9px;
          padding-bottom: 7px;
          /* filter: brightness(90%); /* Optional: filter might cause issues, removed for safety */ */
        }}
        {disabled_style}
      """

    # --- Replace the existing apply_theme method with this corrected version ---
    def apply_theme(self):
      theme_name = settings.get("theme", "macOS")
      if theme_name not in THEMES: theme_name = "macOS"
      theme = THEMES[theme_name]
      is_dark = theme_name in ["dark", "navy", "dracula", "macOSDark"]
      
      palette = QPalette()
      palette.setColor(QPalette.Window, QColor(theme["window"]))
      palette.setColor(QPalette.WindowText, QColor(theme["windowText"]))
      palette.setColor(QPalette.Base, QColor(theme["base"]))
      palette.setColor(QPalette.AlternateBase, QColor(theme["alternateBase"]))
      palette.setColor(QPalette.Text, QColor(theme["text"]))
      palette.setColor(QPalette.Button, QColor(theme["button"]))
      palette.setColor(QPalette.ButtonText, QColor(theme["buttonText"]))
      palette.setColor(QPalette.Highlight, QColor(theme["highlight"]))
      palette.setColor(QPalette.HighlightedText, QColor(theme["highlightedText"]))
      palette.setColor(QPalette.Link, QColor(theme["link"]))
      palette.setColor(QPalette.BrightText, QColor(theme["brightText"]))
      self.setPalette(palette) # Apply palette to main window
      
      # Apply base font
      font = QFont(settings.get("font_family", DEFAULT_FONT_FAMILY), settings.get("font_size", DEFAULT_FONT_SIZE))
      self.setFont(font)
      
      # General Stylesheet adjustments based on theme
      border_color = theme['borderColor']
      window_bg = theme['window'] # Use the theme's window color
      base_bg = theme['base']
      alt_bg = theme['alternateBase']
      text_color = theme['windowText']
      header_bg = alt_bg if is_dark else QColor(alt_bg).lighter(105).name() # Lighter header for light themes
      frame_bg = base_bg if is_dark else QColor(base_bg).darker(102).name() # Slightly darker frame bg for light themes
      
      # Set the stylesheet for the main window and its children
      # Ensure the object name is set in __init__: self.setObjectName("mainWindow")
      self.setStyleSheet(f"""
        QWidget#mainWindow {{
          background-color: {window_bg}; /* Explicitly set main window background */
        }}
    
        /* General styles for all widgets within mainWindow */
        QWidget {{
          font-family: '{settings.get("font_family", DEFAULT_FONT_FAMILY)}';
          font-size: {settings.get("font_size", DEFAULT_FONT_SIZE)}pt;
          color: {text_color};
        }}
        #url_container {{
          background-color: {alt_bg};
          border-radius: 8px;
          padding: 0px;
        }}
        QLineEdit {{
          background-color: transparent; /* Inside styled container */
          border: none;
          padding: 6px;
          color: {text_color};
        }}
        QFrame#styledFrame {{ /* Target styled frames */
          border: 1px solid {border_color};
          border-radius: 8px;
          background-color: {frame_bg};
        }}
        QFrame#styledHeader {{ /* Target styled headers */
          background-color: {header_bg};
          border-top-left-radius: 8px;
          border-top-right-radius: 8px;
          border-bottom: 1px solid {border_color};
          padding: 8px;
          margin: 0px;
        }}
        QFrame#searchContainer {{
          background-color: {alt_bg if is_dark else QColor(alt_bg).lighter(105).name()};
          border-radius: 6px;
          padding: 0px;
        }}
        QSplitter::handle {{
          background-color: {border_color};
          height: 5px; /* Make handle more visible */
          margin: 2px 0;
        }}
        QProgressBar {{
          border: 1px solid {border_color};
          border-radius: 4px;
          text-align: center;
          color: {text_color};
        }}
        QProgressBar::chunk {{
          background-color: {theme['highlight']};
          border-radius: 3px;
          margin: 1px;
        }}
        QTextEdit {{
          background-color: {base_bg}; /* Ensure text edits match base */
          color: {text_color};
          border: none; /* Remove border if frame provides it */
        }}
      """) # Make sure closing """ is correct
      
      # Re-apply specific button styles using the separate function
      self.start_btn.setStyleSheet(self.get_button_style("start"))
      self.stop_btn.setStyleSheet(self.get_button_style("stop"))
      self.save_btn.setStyleSheet(self.get_button_style("save"))
      self.cloud_btn.setStyleSheet(self.get_button_style("cloud"))
    # --- End of replacement block ---


    def setup_tray(self): # From updated
        self.tray_icon = QSystemTrayIcon(self)
        print(f"DEBUG: Trying to set tray icon. self.app_icon is valid: {not self.app_icon.isNull()}, Name: {self.app_icon.name()}") # ADD THIS LINE
        self.tray_icon.setIcon(self.app_icon)
        tray_menu = QMenu()
        show_action = QAction("Show", self); show_action.triggered.connect(self.show_normal_window)
        hide_action = QAction("Hide", self); hide_action.triggered.connect(self.hide)
        quit_action = QAction("Quit", self); quit_action.triggered.connect(self.close) # Use close for event handling
        tray_menu.addAction(show_action); tray_menu.addAction(hide_action); tray_menu.addSeparator(); tray_menu.addAction(quit_action)
        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.activated.connect(self.tray_activated)
        self.tray_icon.show()

    def show_normal_window(self):
        """Ensure window is shown normally."""
        self.showNormal()
        self.activateWindow()
        self.raise_()

    def tray_activated(self, reason): # From updated
        if reason == QSystemTrayIcon.DoubleClick:
            if self.isVisible():
                self.hide()
            else:
                self.show_normal_window()

    def setup_hotkey(self): # From updated
      # Define shortcut *before* the try block
      shortcut = settings.get("keyboard_shortcut", "ctrl+i")
      try:
        # Remove previous hotkey first if any
        # This line below will likely still cause the AttributeError without sudo
        keyboard.remove_all_hotkeys()
        
        if shortcut: # Only add if shortcut is defined
          # This line will also likely fail without sudo
          keyboard.add_hotkey(shortcut, self.on_demand_analysis_sync, suppress=False)
          print(f"Hotkey '{shortcut}' registered.")
        else:
          print("No hotkey defined in settings.")
          
      except Exception as e:
        # 'shortcut' is now defined and accessible here
        print(f"Error setting up hotkey '{shortcut}': {e}")
        # Provide a more informative warning about permissions
        QMessageBox.warning(self, "Hotkey Error",
                  f"Could not register global hotkey '{shortcut}'.\n\n"
                  f"This feature often requires administrator (sudo/root) privileges to monitor keyboard events system-wide.\n\n"
                  f"Try running the application with 'sudo python ...' if global hotkey functionality is essential, "
                  f"or disable/change the hotkey in settings.\n\nError details: {e}")

    def on_demand_analysis_sync(self):
        """Synchronous wrapper to call the async analysis function."""
        print("[DEBUG] Hotkey pressed / Analysis triggered")
        # Check if running in GUI thread is necessary; if not, directly schedule
        # If GUI updates are needed *before* async call, use QTimer
        # QTimer.singleShot(0, lambda: asyncio.create_task(self.on_demand_analysis()))
        # Simpler: just run the async function in the existing loop
        if self.running or self.transcript_box.toPlainText(): # Allow analysis even if not running if there's text
             loop = asyncio.get_event_loop()
             if loop.is_running():
                 asyncio.ensure_future(self.on_demand_analysis(), loop=loop)
             else:
                 print("Warning: Event loop not running, cannot trigger async analysis from hotkey.")
                 self.log("Error: Event loop not running.")
        else:
             self.log("Cannot analyze: No transcript available.")


    def change_prompt(self, prompt_name): # From updated
        if not prompt_name: return
        settings["active_prompt"] = prompt_name
        save_settings_to_file(settings) # Save immediately
        self.log(f"Prompt changed to: {prompt_name}")

    def open_settings(self): # From updated
        # Declare 'settings' as global RIGHT AT THE START of the function
        global settings
      
        # Now you can safely use the global 'settings' variable
        dialog = SettingsDialog(self, settings) # Pass current global settings
        if dialog.exec_():
          # No need for another 'global settings' here
          
          updated_settings = dialog.get_updated_settings()
          # Check what actually changed to avoid unnecessary updates
          changed_keys = {k for k in updated_settings if k not in settings or settings[k] != updated_settings[k]}
          
          # Update the global 'settings' variable
          settings = updated_settings
          save_settings_to_file(settings) # Save to file
          
          # Apply necessary changes based on the now updated global 'settings'
          if any(k in changed_keys for k in ['theme', 'font_family', 'font_size', 'line_spacing', 'word_spacing', 'use_rounded_corners']):
            self.apply_theme()
          if 'prompts' in changed_keys or 'active_prompt' in changed_keys:
            self.prompt_selector.clear()
            self.prompt_selector.addItems([p["name"] for p in settings.get("prompts", [])])
            self.prompt_selector.setCurrentText(settings.get("active_prompt", ""))
          if 'model' in changed_keys:
            self.model_selector.setCurrentText(settings.get("model", DEFAULT_MODEL))
          if 'keyboard_shortcut' in changed_keys:
            self.setup_hotkey() # Re-register hotkey
          # Update icons if changed
          if 'app_icon_path' in changed_keys:
            self.app_icon = load_icon_from_path_or_default(settings.get("app_icon_path"), DEFAULT_APP_ICON)
            self.setWindowIcon(self.app_icon)
            self.tray_icon.setIcon(self.app_icon) # Update tray icon too
          if 'generate_icon_path' in changed_keys:
            self.generate_icon = load_icon_from_path_or_default(settings.get("generate_icon_path"), DEFAULT_GENERATE_ICON)
            self.generate_btn.setIcon(self.generate_icon)
          if 'settings_icon_path' in changed_keys:
            self.settings_icon = load_icon_from_path_or_default(settings.get("settings_icon_path"), DEFAULT_SETTINGS_ICON)
            self.settings_btn.setIcon(self.settings_icon)
            
          self.log("Settings updated successfully")
    def show_wordcloud(self): # From updated
        if not self.transcript_box.toPlainText().strip():
            QMessageBox.warning(self, "Word Cloud", "No transcript data available to generate a word cloud.")
            return
        self.refresh_wordcloud() # Generate cloud data
        self.wordcloud_window.show()
        self.wordcloud_window.activateWindow()

    def refresh_wordcloud(self): # From updated
        text = self.transcript_box.toPlainText()
        self.signals.update_wordcloud.emit(text) # Emit signal to update WC window

    def export_wordcloud(self): # From updated
        pixmap = self.wordcloud_window.get_pixmap()
        if not pixmap or pixmap.isNull():
            QMessageBox.warning(self, "Export Error", "No word cloud generated to export.")
            return
        filepath, _ = QFileDialog.getSaveFileName(self, "Save Word Cloud", f"wordcloud_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.png", "PNG Image (*.png)")
        if filepath:
            if pixmap.save(filepath, "PNG"):
                self.log(f"Word cloud saved to {filepath}")
                if QMessageBox.question(self, "Open File", "Word cloud saved. Open it now?", QMessageBox.Yes | QMessageBox.No) == QMessageBox.Yes:
                    webbrowser.open(f"file:///{os.path.abspath(filepath)}") # Use webbrowser for cross-platform opening
            else:
                QMessageBox.critical(self, "Export Error", f"Failed to save word cloud to {filepath}.")


    def log(self, msg, show_status=True): 
        """Log a message to console, Electron (if enabled), and optionally to status bar.
        
        Args:
            msg: The message to log
            show_status: Whether to update status bar (default: True)
        """
        timestamp = datetime.datetime.now().strftime('%H:%M:%S')
        print(f"[{timestamp}] {msg}")
        
        # Only update status for important messages
        if show_status:
            self.signals.update_status.emit(msg)
            
            # Send to Electron if running in that mode
            # Access the electron_bridge through the app instance
            app = QApplication.instance()
            if hasattr(app, 'electron_bridge') and app.electron_bridge and app.electron_bridge.electron_mode:
                app.electron_bridge.send_status(msg)

    def update_status(self, message): # From updated
        self.status.setText(f"Status: {message}")

    def update_progress(self, value): # From updated
        self.progress_bar.setValue(value)
        if value > 0 and value < 100:
            self.progress_bar.setVisible(True)
        else: # Hide when 0 or 100 after a delay
             QTimer.singleShot(500, lambda: self.progress_bar.setVisible(False))
             if value >= 100: self.progress_bar.setValue(0) # Reset after completion

    def update_meeting_timer_label(self, time_str): # Renamed from update_meeting_timer
         self.timer_label.setText(time_str)

    def update_meeting_time(self): # From updated (for timer logic)
        if not self.meeting_start_time or not self.running: return
        now = datetime.datetime.now()
        elapsed = now - self.meeting_start_time
        h, rem = divmod(int(elapsed.total_seconds()), 3600)
        m, s = divmod(rem, 60)
        time_str = f"Meeting: {h:02}:{m:02}:{s:02}"
        self.signals.update_meeting_timer.emit(time_str) # Emit signal

        # Reminder logic
        reminder_interval = settings.get("meeting_reminder_interval", 15) * 60
        if settings.get("meeting_timer_enabled") and reminder_interval > 0 and self.last_reminder_time:
            if (now - self.last_reminder_time).total_seconds() >= reminder_interval:
                 self.tray_icon.showMessage("Meeting Timer", f"Meeting running for {h:02}:{m:02}:{s:02}", self.app_icon, 5000)
                 self.last_reminder_time = now


    # --- Core Actions ---

    @asyncSlot()
    async def start(self):
        print("[DEBUG] Start button was clicked.")
        if self.running: self.log("Already running."); return

        url = self.url_input.text().strip()
        if not url.startswith("https://otter.ai"):
            QMessageBox.warning(self, "Invalid URL", "Please enter a valid Otter.ai meeting link.")
            return

        self.running = True
        self.seen_lines.clear(); self.buffer.clear(); self.summaries.clear()
        self.suggested_responses.clear(); self.insights.clear(); self.key_phrases.clear()
        self.transcript_box.clear(); self.suggested_response_box.clear(); self.insights_box.clear()
        self.meeting_start_time = datetime.datetime.now(); self.last_reminder_time = self.meeting_start_time
        self.meeting_title = "Unknown Meeting"; self.topic_value.setText("<i>Detecting...</i>")
        self.code_word_detected = False

        self.log("Starting...")
        self.update_progress(10)
        self.start_btn.setEnabled(False); self.stop_btn.setEnabled(True)

        # Start meeting timer
        if settings.get("meeting_timer_enabled", True):
            self.meeting_timer.timeout.connect(self.update_meeting_time)
            self.meeting_timer.start(1000) # Update every second

        # Launch browser scraping
        try:
            await self.scrape_otter(url)
        except Exception as e:
             self.log(f"Error during startup: {e}")
             await self.stop() # Ensure cleanup if start fails


    @asyncSlot()
    async def stop(self):
        if not self.running: return
        self.running = False
        self.main_timer.stop() # Stop polling timer
        self.meeting_timer.stop() # Stop meeting timer
        self.log("Stopping...")
        if self.browser:
            try:
                await self.browser.close()
                print("Browser closed.")
            except Exception as e: print(f"Error closing browser: {e}")
            finally: self.browser = None; self.page = None
        self.start_btn.setEnabled(True); self.stop_btn.setEnabled(False)
        self.update_progress(0)
        self.log("Stopped.")

    def save(self): # Using enhanced save from updated
        transcript = self.transcript_box.toPlainText().strip()
        insights_text = self.insights_box.toPlainText().strip() # Get insights text
        if not transcript:
            QMessageBox.warning(self, "Empty", "No transcript to save.")
            return

        # Prepare metadata
        duration_str = ""
        if self.meeting_start_time:
             elapsed = datetime.datetime.now() - self.meeting_start_time
             h, rem = divmod(int(elapsed.total_seconds()), 3600); m, s = divmod(rem, 60)
             duration_str = f"{h:02}:{m:02}:{s:02}"

        metadata = {
            "app_name": APP_NAME, "version": VERSION, "date": datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "duration": duration_str, "title": self.meeting_title, "insights": insights_text
        }

        formats = { # Added HTML and Markdown
            "Text File (*.txt)": lambda path: export_to_txt(transcript + "\n\n--- AI Insights ---\n" + insights_text, path),
            "Word Document (*.docx)": lambda path: export_to_docx(transcript, path, metadata),
            "PDF Document (*.pdf)": lambda path: export_to_pdf(transcript + "\n\n--- AI Insights ---\n" + insights_text, path, metadata),
            "SRT Subtitle (*.srt)": lambda path: export_to_srt(transcript, path),
            "HTML Document (*.html)": lambda path: export_to_html(transcript, path, metadata),
            "Markdown Document (*.md)": lambda path: export_to_markdown(transcript, path, metadata)
        }

        format_str, _ = QInputDialog.getItem(self, "Select Format", "Choose export format:", list(formats.keys()), 0, False)
        if not format_str: return

        extension = format_str.split("*")[1].rstrip(")")
        default_filename = f"Meeting_{self.meeting_title.replace(' ','_') if self.meeting_title != 'Unknown Meeting' else datetime.datetime.now().strftime('%Y%m%d_%H%M')}{extension}"
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Transcript", default_filename, format_str)
        if not file_path: return

        try:
            save_function = formats[format_str]
            save_function(file_path)
            self.log(f"Saved to {file_path}")
            if QMessageBox.question(self, "Open File", "Saved successfully. Open file?", QMessageBox.Yes | QMessageBox.No) == QMessageBox.Yes:
                webbrowser.open(f"file:///{os.path.abspath(filepath)}")
        except Exception as e:
            QMessageBox.critical(self, "Save Error", f"Could not save file:\n{e}")


    # --- Scraping and AI Logic ---

    @asyncSlot()
    async def scrape_otter(self, url): # Using logic primarily from fixed3, with logging/progress from updated
        print("[DEBUG] scrape_otter started")
        self.log("Initializing browser...")
        self.update_progress(20)
        try:
            async with async_playwright() as p:
                print("[DEBUG] launching browser")
                self.browser = await p.chromium.launch(headless=False, timeout=30000) # Keep headless=False for login possibility
                self.page = await self.browser.new_page()
                self.log("Navigating to URL...")
                self.update_progress(30)
                await self.page.goto(url, timeout=90000) # Long timeout

                self.log("Waiting for page content...")
                self.update_progress(40)
                try: # Wait for DOM, network idle is less critical
                    await self.page.wait_for_load_state("domcontentloaded", timeout=45000)
                    print("[DEBUG] DOM loaded.")
                except Exception as e:
                    print(f"[DEBUG] Load state wait timed out, proceeding anyway: {e}")

                self.log("Checking for login/consent...")
                self.update_progress(50)
                await asyncio.sleep(10) # Allow rendering and potential popups

                # Login check (from fixed3)
                login_elements = await self.page.query_selector_all("input[type='password'], input[type='email'], button:has-text('Sign in'), button:has-text('Log in')")
                if login_elements:
                    self.log("[WARNING] Login may be required. Please log in via the browser.")
                    self.update_progress(60)
                    await asyncio.sleep(20) # Give time for manual login

                # Cookie check (from fixed3)
                try:
                    accept_buttons = await self.page.query_selector_all("button:has-text('Accept'), button:has-text('Accept All'), button:has-text('Allow')")
                    if accept_buttons:
                        await accept_buttons[0].click(); await asyncio.sleep(2)
                        print("[DEBUG] Clicked cookie accept.")
                except Exception as e: print(f"[DEBUG] Cookie button error: {e}")

                self.log("Monitoring transcript...")
                self.update_progress(100) # Indicate monitoring started

                # Start polling timer
                poll_interval = settings.get("poll_interval", DEFAULT_POLL_INTERVAL) * 1000
                self.main_timer.timeout.connect(lambda: asyncio.ensure_future(self.poll_transcript(), loop=asyncio.get_event_loop()))
                self.main_timer.start(poll_interval)

                # Keep browser alive check (optional, can be intensive)
                # asyncio.create_task(self.keep_browser_alive())

        except Exception as e:
            self.log(f"Browser launch/navigation error: {e}")
            await self.stop() # Stop cleanly
        # No finally block needed here as stop handles cleanup

    async def poll_transcript(self): # Using polling logic from fixed3 with UI signals
        """Poll for new transcript content from Otter.ai with improved error handling."""
        # Check for valid browser session and connection
        if not self.running or not self.page or not self.browser or not self.browser.is_connected():
            if self.running: 
                self.log("Browser disconnected. Stopping.") 
                await self.stop()
            return
            
        print("[DEBUG] Polling transcript...")
        try:
            # Try specialized extraction first (more reliable if structure is known)
            transcript_lines = await self.extract_otter_transcript_specialized()

            # If specialized fails, try generic selectors with priority to Otter.ai selectors
            if not transcript_lines:
                selectors_to_try = [
                    # Primary Otter.ai selectors - keep these as top priority
                    'div[class*="transcript-snippet__content__body__sentence"]',  # Most important Otter.ai selector
                    '.transcript-snippet__content__body__sentence',
                    '.transcript-snippet__content__body',
                    '.transcript-snippet__content__container span', # Words
                    # Fallbacks for other potential layouts
                    '[data-testid="transcript-paragraph"]', '.ot-transcript__paragraph',
                    '.ot-transcript-content span', '.ot-transcription-content p',
                    '.transcript-paragraph', '.transcript-section p', '.transcript-container p',
                ]
                
                for selector in selectors_to_try:
                    try:
                        lines = await self.page.evaluate(f"""
                            () => Array.from(document.querySelectorAll('{selector}'))
                                    .map(e => e.innerText.trim()).filter(Boolean)
                        """)
                        if lines:
                            transcript_lines = lines
                            print(f"[DEBUG] Found {len(lines)} lines with selector: {selector}")
                            break # Stop after first successful selector
                    except Exception as e:
                        print(f"[DEBUG] Error with selector {selector}: {e}")
                        # Continue trying other selectors

            # If we still don't have transcript lines, log and return
            if not transcript_lines:
                print("[DEBUG] No transcript lines found in this poll.")
                return

            # Process any found transcript lines
            new_content_added = False
            for line in transcript_lines:
                try:
                    line_hash = hash(line) # Simple way to check uniqueness
                    if line_hash not in self.seen_lines:
                        # Basic filtering
                        if (len(line) > 5 and
                            "cookie" not in line.lower() and
                            "consent" not in line.lower() and
                            "privacy" not in line.lower()):

                            timestamp = datetime.datetime.now().strftime("[%H:%M:%S]")
                            display_line = f"{timestamp} {line}"
                            self.signals.append_transcript.emit(display_line) # Update UI
                            self.buffer.append(line) # Add raw line to buffer for AI
                            self.seen_lines.add(line_hash)
                            new_content_added = True

                            # Check for question trigger
                            if "?" in line and any(word in line.lower() for word in ["you", "your"]):
                                asyncio.create_task(self.trigger_insight(line))
                except Exception as line_error:
                    print(f"[ERROR] Error processing line: {line_error}")
                    # Continue with next line

            # Trigger summary if buffer full
            if new_content_added and len(self.buffer) >= settings.get("summary_line_interval", 10):
                await self.trigger_summary() # Use helper

        except Exception as e:
            print(f"[ERROR] Polling failed: {e}")
            
            # Track consecutive errors for recovery
            self.error_count = getattr(self, 'error_count', 0) + 1
            if self.error_count > 3:
                self.log("Multiple polling errors detected. Attempting recovery...")
                self.error_count = 0
                await self.attempt_browser_recovery()
        else:
            # Reset error count on success
            self.error_count = 0
            
    async def attempt_browser_recovery(self):
        """Attempt to recover from browser issues."""
        try:
            self.log("Refreshing browser connection...")
            if self.page:
                await self.page.reload(timeout=30000)
                self.log("Browser refreshed successfully.")
                return True
        except Exception as e:
            self.log(f"Browser recovery failed: {e}")
            await self.stop()  # Stop if recovery fails
            return False


    async def extract_otter_transcript_specialized(self):
        """More targeted extraction for Otter.ai (from fixed3)."""
        print("[DEBUG] Using specialized Otter.ai extraction")
        try:
            # Prioritize sentence elements
            sentences = await self.page.evaluate("""
                () => Array.from(document.querySelectorAll('.transcript-snippet__content__body__sentence'))
                          .map(el => el.innerText.trim()).filter(Boolean)
            """)
            if sentences: return sentences

            # Fallback: Reconstruct from word spans within containers
            word_containers = await self.page.evaluate("""
                () => Array.from(document.querySelectorAll('div[class*="transcript-snippet__content__body"]'))
                          .map(container => {
                                const words = Array.from(container.querySelectorAll('span[class*="word"]'))
                                                .map(span => span.innerText.trim());
                                return words.length > 0 ? words.join(' ') : container.innerText.trim();
                           }).filter(Boolean)
            """)
            if word_containers: return word_containers

             # Fallback: Try speaker blocks (less reliable for full transcript text)
            speaker_blocks = await self.page.evaluate("""
                 () => Array.from(document.querySelectorAll('.transcript-snippet__speaker-container'))
                           .map(block => {
                                const speaker = block.querySelector('.transcript-snippet__speaker__icon')?.getAttribute('aria-label') || 'Speaker';
                                const content = block.nextElementSibling?.innerText?.trim();
                                return content ? `${speaker}: ${content}` : null;
                           }).filter(Boolean)
            """)
            if speaker_blocks: return speaker_blocks # Might only capture parts

            print("[DEBUG] Specialized extraction failed.")
            return []
        except Exception as e:
            print(f"[DEBUG] Error in specialized extraction: {e}")
            return []

    async def trigger_summary(self):
        """Generate summary based on buffer content with improved error handling."""
        # Validate buffer content
        if not self.buffer:
            self.log("No content to summarize.")
            return
            
        # Store a copy of the buffer in case something fails
        buffer_copy = self.buffer.copy()
        buffer_size = len(buffer_copy)
            
        # Start the process
        print(f"[AI] Generating summary for {buffer_size} lines.")
        self.log("AI generating summary...")
        self.update_progress(50)
        
        try:
            # Get the active prompt template
            prompt_content = get_active_prompt_content()
            if not prompt_content:
                self.log("Warning: Using default prompt (no active prompt found)")
                prompt_content = DEFAULT_PROMPT_CONTENT
                
            # Format prompt with transcript content
            full_prompt = f"{prompt_content}\n\nRecent Transcript Lines:\n{chr(10).join(buffer_copy)}"
            
            # Clear buffer before async call to prevent duplicate processing
            self.buffer.clear()
            
            # Call the AI API through a thread to keep UI responsive
            result = await asyncio.to_thread(ask_groq, full_prompt)
            
            # Check for error responses
            if result and result.startswith("[AI Error]"):
                self.log(f"Summary error: {result}")
                self.signals.append_insights.emit(f"AI ANALYSIS ERROR:\n{result}\n\nPlease try again later.")
                return
                
            # Process the successful response
            self.process_groq_response(result)
            self.log("AI summary generated successfully.")
            
        except Exception as e:
            self.log(f"AI summary error: {e}")
            import traceback
            print(f"Summary error details:\n{traceback.format_exc()}")
            
            # Add items back to buffer if we failed early (before API call)
            if self.buffer is not None and len(self.buffer) == 0 and not buffer_copy is None:
                self.buffer.extend(buffer_copy)
                self.log("Restored buffer content after error.")
                
        finally:
            # Always reset progress bar
            self.update_progress(0)


    async def trigger_insight(self, triggering_line):
        """Generate AI insight based on a specific line (e.g., a question) with improved error handling."""
        if not triggering_line or triggering_line.isspace():
            return  # Silently ignore empty triggers
            
        print(f"[AI] Generating insight for line: {triggering_line}")
        self.log("AI generating insight...")
        self.update_progress(50)
        
        try:
            # Get active prompt template
            prompt_content = get_active_prompt_content()
            if not prompt_content:
                self.log("Warning: Using default prompt (no active prompt found)")
                prompt_content = DEFAULT_PROMPT_CONTENT
            
            # Gather context (safely)
            try:
                # Try to get the last 10 lines as context
                context_lines = list(self.seen_lines)[-10:] if self.seen_lines else []
                context = "\n".join(str(line) for line in context_lines)
            except Exception as context_error:
                self.log(f"Error gathering context: {context_error}")
                context = "Context unavailable"
            
            # Build prompt with good error handling
            insight_prompt = (f"{prompt_content}\n\nContext:\n{context}\n\n"
                            f"Specifically respond to or analyze this line: {triggering_line}")
            
            # Call AI API in a separate thread
            result = await asyncio.to_thread(ask_groq, insight_prompt)
            
            # Check for error responses
            if result and result.startswith("[AI Error]"):
                self.log(f"Insight error: {result}")
                self.signals.append_suggested_response.emit(f"AI ANALYSIS ERROR:\n{result}\n\nPlease try again later.")
                return
                
            # Process successful response
            self.process_groq_response(result)
            self.log("AI insight generated successfully.")
            
        except Exception as e:
            self.log(f"AI insight error: {e}")
            import traceback
            print(f"Insight error details:\n{traceback.format_exc()}")
            
        finally:
            # Always reset progress
            self.update_progress(0)

    @asyncSlot()
    async def on_demand_analysis(self):
        """Triggers AI analysis based on the full transcript with improved error handling."""
        # Check if there's transcript content to analyze
        transcript_text = self.transcript_box.toPlainText().strip()
        if not transcript_text:
            self.log("No transcript content available for analysis.")
            return
            
        print("[AI] On-demand analysis triggered.")
        self.log("AI generating comprehensive analysis...")
        self.update_progress(30)
        
        try:
            # Get active prompt template
            prompt_content = get_active_prompt_content()
            if not prompt_content:
                self.log("Warning: Using default prompt (no active prompt found)")
                prompt_content = DEFAULT_PROMPT_CONTENT
                
            # Limit transcript size if extremely large to prevent API issues
            if len(transcript_text) > 25000:  # Approximated reasonable limit
                self.log("Transcript is very large - analyzing most recent portion")
                # Keep the beginning for context, but focus on the recent content
                transcript_lines = transcript_text.split('\n')
                # Keep first few lines for context plus most recent content
                transcript_text = '\n'.join(transcript_lines[:10] + ['...'] + transcript_lines[-300:])
                
            # Prepare prompt for analysis
            full_prompt = f"{prompt_content}\n\nFull Transcript Context:\n{transcript_text}"
            
            # Update UI to show analysis in progress
            self.update_progress(50)
            
            # Call API in separate thread to keep UI responsive
            result = await asyncio.to_thread(ask_groq, full_prompt)
            
            # Check for error responses
            if result and result.startswith("[AI Error]"):
                self.log(f"Analysis error: {result}")
                self.signals.append_insights.emit(f"AI ANALYSIS ERROR:\n{result}\n\nPlease try again later.")
                return
                
            # Process successful response
            self.process_groq_response(result)
            self.log("AI analysis completed successfully.")
            
        except Exception as e:
            self.log(f"AI analysis error: {e}")
            import traceback
            print(f"Analysis error details:\n{traceback.format_exc()}")
            
            # Inform the user about the error
            error_msg = f"AI ANALYSIS ERROR: {str(e)[:100]}..."
            self.signals.append_insights.emit(error_msg)
            
        finally:
            # Always reset progress
            self.update_progress(0)


    def process_groq_response(self, response): # From updated
        """Process the AI response and update specific UI sections."""
        lines = response.split('\n')
        sections = {"SUGGESTED RESPONSE": [], "KEY INSIGHTS": [], "TOPIC SUMMARY": []}
        current_section_key = "KEY INSIGHTS" # Default to insights if no header found

        for line in lines:
            stripped_line = line.strip()
            if not stripped_line: continue

            found_header = False
            for header in sections.keys():
                if stripped_line.startswith(header + ":"):
                    current_section_key = header
                    content = stripped_line[len(header)+1:].strip()
                    if content: sections[current_section_key].append(content)
                    found_header = True
                    break
                elif stripped_line == header + ":": # Header on its own line
                    current_section_key = header
                    found_header = True
                    break

            if not found_header:
                sections[current_section_key].append(stripped_line)

        # Update UI sections
        if sections["SUGGESTED RESPONSE"]:
            self.signals.append_suggested_response.emit("SUGGESTED RESPONSE:\n" + "\n".join(sections["SUGGESTED RESPONSE"]))
        # Combine Insights and Summary into the insights box
        insights_text = ""
        if sections["KEY INSIGHTS"]:
            insights_text += "KEY INSIGHTS:\n" + "\n".join(sections["KEY INSIGHTS"]) + "\n\n"
        if sections["TOPIC SUMMARY"]:
            insights_text += "TOPIC SUMMARY:\n" + "\n".join(sections["TOPIC SUMMARY"])
        if insights_text:
             self.signals.append_insights.emit(insights_text.strip())
        elif not any(sections.values()): # If response had no structure, put it all in insights
             self.signals.append_insights.emit("AI RESPONSE:\n"+ response) # Label it clearly


    def append_transcript(self, text):
        """Append text to transcript with formatting and checks."""
        print(f"DEBUG: append_transcript called with: {text[:30]}...")
        if not text:
            return
            
        # Apply spell checking if enabled
        if settings.get("spellcheck_enabled", False):
            text = correct_spelling(text)
                
        # Apply name anonymization if enabled
        if settings.get("anonymize_transcript", False):
            text, name_mapping = anonymize_names(text, self.known_names)
            # Update known names for future anonymization
            self.known_names = list(name_mapping.keys())
                
        # Check for code word activation
        code_word = settings.get("code_word", DEFAULT_CODE_WORD).lower()
        if code_word and code_word in text.lower() and not self.code_word_detected:
            self.code_word_detected = True
            self.log(f"Code word '{code_word}' detected! Triggering analysis.")
            self.highlight_transcript(code_word, "#FFC107") # Highlight code word briefly
            QTimer.singleShot(0, self.on_demand_analysis) # Trigger analysis
            QTimer.singleShot(3000, lambda: self.highlight_transcript(code_word, None)) # Remove highlight
                
        try:
            # Format based on whether it contains a timestamp
            if text.startswith('[') and ']' in text:
                # Split into timestamp and content
                timestamp_end = text.find(']') + 1
                timestamp = text[:timestamp_end]
                content = text[timestamp_end:].strip()
                
                # Add timestamp with gray color
                self.transcript_box.append_formatted(
                    timestamp, {"color": "#888888", "font_size": 12}
                )
                
                # Add content with normal format
                self.transcript_box.append_formatted(
                    content, {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                )
            else:
                # Insert regular text
                self.transcript_box.append_formatted(
                    text, {"font_size": settings.get("font_size", DEFAULT_FONT_SIZE)}
                )
                
            # Make sure the text is visible by scrolling to the bottom
            self.transcript_box.ensureCursorVisible()
        except Exception as e:
            print(f"ERROR in append_transcript: {e}")
            # Fallback method if the custom method fails
            try:
                self.transcript_box.append(text)
            except Exception as e:
                print(f"Critical error: Could not append text to transcript box: {e}")

    def append_suggested_response(self, text): # From updated
        self.suggested_response_box.clear() # Clear previous suggestions
        self.suggested_response_box.append_formatted(text, {}) # Use default formatting

    def append_insights(self, text): # From updated (merges insights/summary)
         self.insights_box.clear() # Clear previous insights/summary
         # Basic Markdown-like formatting for headers
         lines = text.split('\n')
         for line in lines:
              if line.startswith("SUGGESTED RESPONSE:"): # Shouldn't be here, but handle anyway
                   self.insights_box.append_formatted(line, {"bold": True, "color": "#007AFF", "font_size": settings.get("font_size", DEFAULT_FONT_SIZE) + 1})
              elif line.startswith("KEY INSIGHTS:"):
                   self.insights_box.append_formatted(line, {"bold": True, "color": "#FF9500", "font_size": settings.get("font_size", DEFAULT_FONT_SIZE) + 1})
              elif line.startswith("TOPIC SUMMARY:"):
                   self.insights_box.append_formatted(line, {"bold": True, "color": "#5AC8FA", "font_size": settings.get("font_size", DEFAULT_FONT_SIZE) + 1})
              elif line.startswith("AI RESPONSE:"): # Handle unstructured response header
                  self.insights_box.append_formatted(line, {"bold": True, "color": "#34C759", "font_size": settings.get("font_size", DEFAULT_FONT_SIZE) + 1})
              else:
                   self.insights_box.append_formatted(line, {})


    def highlight_transcript(self, text, color="#FFFF99"): # From updated
        if color:
             self.transcript_box.highlight_text(re.escape(text), color)
        else: # Clear highlight if color is None
             self.transcript_box.remove_highlight_pattern(re.escape(text))

    def search_transcript_text(self, query): # Renamed from search_transcript
        """Highlight search terms in the transcript box."""
        self.transcript_box.clear_highlights() # Clear previous search highlights
        if query:
            # Use a distinct color for search highlighting
            search_color = QColor(settings.get("highlight", "#4DB6AC")).lighter(130).name()
            self.transcript_box.highlight_text(re.escape(query), search_color)


    def update_key_phrases(self, phrases): # From updated
        self.key_phrases = phrases
        # Optionally highlight key phrases (can make transcript busy)
        # key_phrase_color = QColor(settings.get("highlight", "#4DB6AC")).lighter(150).name()
        # for phrase in phrases:
        #     self.transcript_box.highlight_text(re.escape(phrase), key_phrase_color)

    def update_meeting_title(self, title): # From updated
        if title and title != "Unknown Meeting":
            self.meeting_title = title
            self.topic_value.setText(f"<b>{title}</b>") # Make title bold
        else:
            self.topic_value.setText("<i>Unknown Meeting</i>")

    # --- Window Event Handling ---
    def closeEvent(self, event): # Ensure clean shutdown
        print("Close event triggered.")
        # Stop monitoring if running
        if self.running:
             # Need to run async stop in event loop if possible
             # This is tricky on close. Simplest might be just to close browser sync if possible.
             # Or show a message "Please stop monitoring before closing".
             if QMessageBox.question(self, "Confirm Exit", "Monitoring is active. Stop monitoring and exit?", QMessageBox.Yes | QMessageBox.No) == QMessageBox.Yes:
                  # We can't easily run async stop here cleanly.
                  # Keyboard interrupt might be needed for the hotkey thread.
                  # For playwright, force closing might be the only option.
                  print("Force closing browser on exit...")
                  # This might leave zombie processes, not ideal.
                  # A cleaner way involves managing the async loop lifecycle better.
                  event.accept()
                  QApplication.quit() # Force quit
             else:
                  event.ignore()
                  return
        else:
            event.accept() # Accept close if not running

# --------------------------
# Electron Integration Support
# --------------------------
class ElectronBridge:
    """Bridge class for Electron integration with actual IPC implementation."""
    
    def __init__(self, app=None):
        self.app = app
        self.electron_mode = False
        self.check_electron_mode()
        self.message_queue = []
        self.initialized = False
        
    def check_electron_mode(self):
        """Check if we're running in Electron mode."""
        # Check command line arguments first (most reliable)
        if '--electron-mode' in sys.argv:
            self.electron_mode = True
            print("Electron mode enabled via command line.")
            return
            
        # Check environment variables as fallback
        env_markers = ['ELECTRON_RUN_AS_NODE', 'ELECTRON_NO_ASAR', 'ELECTRON_DISABLE_SECURITY_WARNINGS']
        for marker in env_markers:
            if marker in os.environ:
                self.electron_mode = True
                print(f"Detected Electron environment via {marker}.")
                return
        
    def setup_ipc(self):
        """Setup IPC communication with Electron."""
        if not self.electron_mode:
            return
            
        # Set up stdin listening for messages from Electron
        # In a production app, you might want to use a more robust IPC method
        # like sockets, HTTP, or a dedicated IPC library
        
        if sys.platform == 'win32':
            # On Windows, we need to ensure binary mode for stdin/stdout
            import msvcrt
            import os
            msvcrt.setmode(sys.stdin.fileno(), os.O_BINARY)
            msvcrt.setmode(sys.stdout.fileno(), os.O_BINARY)
        
        # Start a thread to read from stdin
        threading.Thread(target=self._stdin_listener, daemon=True).start()
        
        # Send queued messages
        self.initialized = True
        self._send_queued_messages()
        
        print("IPC:{'type':'status','content':'Electron IPC bridge initialized'}")
        
    def _stdin_listener(self):
        """Listen for messages from Electron via stdin."""
        try:
            while True:
                line = sys.stdin.readline().strip()
                if not line:
                    time.sleep(0.1)
                    continue
                    
                # Parse the message
                try:
                    message = json.loads(line)
                    self._handle_message(message)
                except json.JSONDecodeError:
                    print(f"IPC:{{'type':'error','content':'Invalid JSON message from Electron: {line}'}}")
        except Exception as e:
            print(f"IPC:{{'type':'error','content':'Error in IPC listener: {str(e)}'}}")
            
    def _handle_message(self, message):
        """Handle a message from Electron."""
        if not isinstance(message, dict) or 'type' not in message:
            return
            
        msg_type = message.get('type')
        
        if msg_type == 'ping':
            # Simple ping-pong to test connection
            self.send_message('pong', {'timestamp': time.time()})
        elif msg_type == 'command':
            # Execute a command in the app
            # Would need to implement command handling in a real app
            pass
        
    def notify_ready(self):
        """Notify Electron that app is ready."""
        if not self.electron_mode:
            return
            
        self.send_message('ready', {
            'appName': APP_NAME,
            'version': VERSION,
            'timestamp': time.time()
        })
        
    def _send_queued_messages(self):
        """Send any queued messages."""
        if not self.message_queue:
            return
            
        for msg_type, data in self.message_queue:
            self._send_ipc_message(msg_type, data)
            
        self.message_queue.clear()
        
    def send_message(self, message_type, data=None):
        """Send a message to Electron."""
        if not self.electron_mode:
            return
            
        if not self.initialized:
            # Queue message if IPC not yet initialized
            self.message_queue.append((message_type, data))
            return
            
        self._send_ipc_message(message_type, data)
        
    def _send_ipc_message(self, message_type, data=None):
        """Send an IPC message to Electron via stdout."""
        message = {
            'type': message_type,
            'timestamp': time.time()
        }
        
        if data:
            if isinstance(data, dict):
                message.update(data)
            else:
                message['content'] = data
                
        # Format: IPC:{json_message}
        # This format is recognized by the Electron process
        try:
            json_message = json.dumps(message)
            print(f"IPC:{json_message}")
            sys.stdout.flush()  # Ensure message is sent immediately
        except Exception as e:
            print(f"Error sending IPC message: {e}")
            
    def log_to_electron(self, message):
        """Log a message to the Electron console."""
        self.send_message('log', {'content': message})
        
    def send_status(self, status_message):
        """Send a status update to Electron."""
        self.send_message('status', {'content': status_message})
        
    def send_error(self, error_message):
        """Send an error message to Electron."""
        self.send_message('error', {'content': error_message})

# --------------------------
# MAIN LAUNCHER (Using structure from fixed3, but simplified for GUI only)
# --------------------------
def main():
    # Create Qt application
    app = QApplication(sys.argv)
    
    # Set application info
    app.setApplicationName(APP_NAME)
    app.setApplicationVersion(VERSION)
    
    # Setup Electron bridge if needed
    app.electron_bridge = ElectronBridge(app)
    app.electron_bridge.setup_ipc()
    
    # Setup event loop for async operations
    loop = QEventLoop(app)
    asyncio.set_event_loop(loop)

    # Create main window
    window = MeetingAssistant()
    
    # In Electron mode, we might want to delay showing the window
    if app.electron_bridge.electron_mode:
        app.electron_bridge.send_status("Application initialized and ready")
        # Let Electron know we're ready
        app.electron_bridge.notify_ready()
    else:
        # In standalone mode, show the window immediately
        window.show()
    
    # Run event loop
    with loop:
        sys.exit(loop.run_forever())

# --------------------------
# Electron Packaging Support
# --------------------------
def create_electron_files():
    """Create necessary Electron files if they don't exist."""
    electron_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "electron")
    os.makedirs(electron_dir, exist_ok=True)
    
    # Package.json
    package_json = {
        "name": "meeting-assistant-plus-electron",
        "version": VERSION,
        "description": "Meeting Assistant Plus - Electron Wrapper",
        "main": "main.js",
        "scripts": {
            "start": "electron .",
            "build": "electron-builder"
        },
        "author": "MeetingAssistant+ Team",
        "dependencies": {
            "electron": "^29.0.0"
        },
        "devDependencies": {
            "electron-builder": "^24.9.1"
        },
        "build": {
            "appId": "com.meetingassistantplus.app",
            "productName": APP_NAME,
            "files": ["*.js", "*.html", "*.css", "assets/**/*", "dist/**/*"],
            "extraResources": ["../MeetingAssistantPlus_merged.py", "../settings.json"],
            "mac": {
                "category": "public.app-category.productivity"
            },
            "win": {
                "target": ["nsis", "portable"]
            },
            "linux": {
                "target": ["AppImage", "deb"],
                "category": "Office"
            }
        }
    }
    
    # Main.js (Electron entry point)
    main_js = """const { app, BrowserWindow, ipcMain, Menu, Tray, dialog, shell } = require('electron');
const path = require('path');
const { spawn } = require('child_process');
const fs = require('fs');
const isDev = !app.isPackaged;

let mainWindow;
let pyProc = null;
let tray = null;
let pyPort = null;
let appIcon = null;
let pythonReady = false;
let launchTimeout = null;

// Prevent multiple instances
const gotTheLock = app.requestSingleInstanceLock();
if (!gotTheLock) {
    app.quit();
    return;
}

// Find the python executable
function getPythonPath() {
    // Check for Python executable in common locations or use platform default
    const pythonPaths = [
        'python3',
        'python',
        path.join(app.getAppPath(), 'python', 'python'),
        'py'
    ];
    
    // On Windows, prioritize 'python.exe'
    if (process.platform === 'win32') {
        pythonPaths.unshift('python.exe');
    }
    
    // In a production app, you would check each path with fs.existsSync or try-catch spawn
    // For simplicity, we'll just return the first one
    return pythonPaths[0];
}

function createWindow() {
    // Load app icon
    appIcon = path.join(__dirname, 'assets', 'icon.png');
    
    // Create the browser window
    mainWindow = new BrowserWindow({
        width: 1200,
        height: 800,
        minWidth: 800,
        minHeight: 600,
        webPreferences: {
            nodeIntegration: true,
            contextIsolation: false,
            preload: path.join(__dirname, 'preload.js')
        },
        icon: appIcon,
        show: false, // Don't show until Python is ready
        backgroundColor: '#f5f5f5'
    });

    // Load the index.html file
    mainWindow.loadFile('index.html');
    
    // Open DevTools in development mode
    if (isDev) {
        mainWindow.webContents.openDevTools();
    }
    
    // Set up window events
    mainWindow.on('closed', function() {
        mainWindow = null;
        closePythonProcess();
    });
    
    mainWindow.on('close', function(e) {
        if (app.quitting) {
            mainWindow = null;
        } else {
            e.preventDefault();
            mainWindow.hide();
        }
    });
    
    // Start the Python process
    startPythonProcess();
    
    // Set a timeout for Python startup
    launchTimeout = setTimeout(() => {
        if (!pythonReady) {
            dialog.showErrorBox(
                'Launch Error',
                'The application backend failed to start within a reasonable time. Please check your Python installation.'
            );
            app.quit();
        }
    }, 60000); // 60 second timeout
    
    setupTray();
}

function setupTray() {
    tray = new Tray(appIcon);
    const contextMenu = Menu.buildFromTemplate([
        { label: 'Show', click: () => mainWindow.show() },
        { label: 'Hide', click: () => mainWindow.hide() },
        { type: 'separator' },
        { label: 'Quit', click: () => { app.quitting = true; app.quit(); } }
    ]);
    tray.setToolTip('Meeting Assistant Plus');
    tray.setContextMenu(contextMenu);
    
    tray.on('click', () => {
        if (mainWindow.isVisible()) {
            mainWindow.hide();
        } else {
            mainWindow.show();
        }
    });
}

function closePythonProcess() {
    if (pyProc) {
        try {
            // Try to gracefully terminate first
            pyProc.kill('SIGTERM');
            
            // Force kill after a timeout if needed
            setTimeout(() => {
                if (pyProc) {
                    console.log('Forcing Python process termination');
                    pyProc.kill('SIGKILL');
                }
            }, 5000);
        } catch (e) {
            console.error('Error closing Python process:', e);
        }
    }
}

function parseIPCMessage(data) {
    // Try to parse Python outputs as IPC messages
    const lines = data.toString().split('\\n');
    for (const line of lines) {
        if (line.startsWith('IPC:')) {
            try {
                const jsonStr = line.substring(4).trim();
                const message = JSON.parse(jsonStr);
                
                // Handle various message types
                if (message.type === 'ready') {
                    pythonReady = true;
                    clearTimeout(launchTimeout);
                    mainWindow.show();
                    
                    // Send a message to the renderer
                    mainWindow.webContents.send('python-message', JSON.stringify({
                        type: 'ready',
                        content: 'Application is now running'
                    }));
                }
                else if (message.type === 'status') {
                    // Pass status messages to the renderer
                    mainWindow.webContents.send('python-message', JSON.stringify(message));
                }
                else if (message.type === 'error') {
                    dialog.showErrorBox('Application Error', message.content);
                }
            } catch (e) {
                console.error('Error parsing IPC message:', e);
            }
        }
    }
}

function startPythonProcess() {
    const pythonExecutable = getPythonPath();
    
    // Find the script path
    const script = isDev 
        ? path.join(__dirname, '..', 'MeetingAssistantPlus_merged.py')
        : path.join(process.resourcesPath, 'MeetingAssistantPlus_merged.py');
    
    // Check if script exists
    if (!fs.existsSync(script)) {
        dialog.showErrorBox(
            'Configuration Error',
            `Could not find Python script at: ${script}`
        );
        app.quit();
        return;
    }
    
    // Build command line arguments
    const args = [script, '--electron-mode'];
    
    console.log(`Starting Python process: ${pythonExecutable} ${args.join(' ')}`);
    
    try {
        pyProc = spawn(pythonExecutable, args, {
            stdio: ['pipe', 'pipe', 'pipe']
        });
        
        if (pyProc != null) {
            console.log('Python process started with PID:', pyProc.pid);
            
            pyProc.stdout.on('data', (data) => {
                console.log(`Python stdout: ${data}`);
                parseIPCMessage(data);
            });
            
            pyProc.stderr.on('data', (data) => {
                console.error(`Python stderr: ${data}`);
                
                // Send error messages to the renderer
                if (mainWindow && mainWindow.webContents) {
                    mainWindow.webContents.send('python-message', JSON.stringify({
                        type: 'status',
                        content: `Error: ${data}`
                    }));
                }
            });
            
            pyProc.on('close', (code) => {
                console.log(`Python process exited with code ${code}`);
                pyProc = null;
                
                // Show error if Python process exits unexpectedly
                if (code !== 0 && !app.quitting) {
                    dialog.showErrorBox(
                        'Backend Error',
                        `The application backend exited unexpectedly with code ${code}.`
                    );
                    app.quit();
                }
            });
        }
    } catch (error) {
        console.error('Failed to start Python process:', error);
        dialog.showErrorBox(
            'Startup Error',
            `Failed to start the application backend: ${error.message}`
        );
        app.quit();
    }
}

// Handle IPC messages from renderer
ipcMain.on('renderer-ready', () => {
    console.log('Renderer process is ready');
});

// Handle Python commands from renderer
ipcMain.on('python-command', (event, command) => {
    if (pyProc && pyProc.stdin) {
        pyProc.stdin.write(command + '\\n');
    }
});

// App lifecycle events
app.on('ready', createWindow);

app.on('window-all-closed', function() {
    if (process.platform !== 'darwin') {
        app.quit();
    }
});

app.on('before-quit', () => {
    app.quitting = true;
});

app.on('activate', function() {
    if (mainWindow === null) {
        createWindow();
    } else {
        mainWindow.show();
    }
});

// Handle second instance attempts
app.on('second-instance', () => {
    // Someone tried to run a second instance, focus our window instead
    if (mainWindow) {
        if (mainWindow.isMinimized()) mainWindow.restore();
        mainWindow.focus();
        mainWindow.show();
    }
});
"""
    
    # Simple HTML page
    index_html = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Meeting Assistant Plus</title>
    <link rel="stylesheet" href="styles.css">
</head>
<body>
    <div id="loading">
        <div class="spinner"></div>
        <h2>Starting Meeting Assistant Plus...</h2>
        <p>Please wait while the application initializes</p>
    </div>
    
    <script src="renderer.js"></script>
</body>
</html>
"""
    
    # CSS styles for Electron
    styles_css = """body {
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, 'Open Sans', 'Helvetica Neue', sans-serif;
    margin: 0;
    padding: 0;
    background-color: #f5f5f5;
    color: #333;
    display: flex;
    justify-content: center;
    align-items: center;
    height: 100vh;
    overflow: hidden;
}

#loading {
    text-align: center;
    padding: 30px;
    border-radius: 10px;
    background-color: white;
    box-shadow: 0 4px 15px rgba(0, 0, 0, 0.1);
    width: 400px;
}

.spinner {
    width: 50px;
    height: 50px;
    margin: 20px auto;
    border: 5px solid #f3f3f3;
    border-top: 5px solid #4DB6AC;
    border-radius: 50%;
    animation: spin 1s linear infinite;
}

h2 {
    margin-top: 0;
    color: #4DB6AC;
}

p {
    color: #666;
    margin-bottom: 5px;
}

.status {
    margin-top: 20px;
    padding: 10px;
    background-color: #f0f0f0;
    border-radius: 5px;
    font-family: monospace;
    max-height: 100px;
    overflow-y: auto;
    text-align: left;
}

@keyframes spin {
    0% { transform: rotate(0deg); }
    100% { transform: rotate(360deg); }
}
"""

    # Renderer.js for Electron UI
    renderer_js = """// This file handles the renderer process (UI)
document.addEventListener('DOMContentLoaded', () => {
    const loadingDiv = document.getElementById('loading');
    const statusDiv = document.createElement('div');
    statusDiv.className = 'status';
    statusDiv.innerHTML = 'Initializing application...';
    loadingDiv.appendChild(statusDiv);
    
    // Create a listener for IPC messages from the main process
    const { ipcRenderer } = require('electron');
    
    ipcRenderer.on('python-message', (event, data) => {
        try {
            const message = JSON.parse(data);
            
            if (message.type === 'status') {
                statusDiv.innerHTML = message.content;
            } else if (message.type === 'ready') {
                // App is ready - hide the loading screen (would normally transition to the app)
                document.body.innerHTML = `
                    <div style="text-align: center; padding: 30px;">
                        <h2>Meeting Assistant Plus</h2>
                        <p>The application is now running in its own window.</p>
                    </div>
                `;
            }
        } catch (e) {
            statusDiv.innerHTML = `Error processing message: ${e.message}`;
        }
    });
    
    // Let the main process know the UI is ready
    ipcRenderer.send('renderer-ready');
});
"""

    # Preload script for Electron (security best practice)
    preload_js = """// Preload script with secure context bridge
const { contextBridge, ipcRenderer } = require('electron');

// Expose a limited API to the renderer process
contextBridge.exposeInMainWorld('electronAPI', {
    // Send messages to the main process
    sendMessage: (channel, data) => {
        ipcRenderer.send(channel, data);
    },
    
    // Receive messages from the main process
    onMessage: (channel, callback) => {
        ipcRenderer.on(channel, (event, ...args) => callback(...args));
    }
});
"""

    # Create assets directory for icons
    assets_dir = os.path.join(electron_dir, "assets")
    os.makedirs(assets_dir, exist_ok=True)
    
    # Write files if they don't exist
    package_json_path = os.path.join(electron_dir, "package.json")
    if not os.path.exists(package_json_path):
        with open(package_json_path, "w") as f:
            json.dump(package_json, f, indent=2)
            
    main_js_path = os.path.join(electron_dir, "main.js")
    if not os.path.exists(main_js_path):
        with open(main_js_path, "w") as f:
            f.write(main_js)
            
    index_html_path = os.path.join(electron_dir, "index.html")
    if not os.path.exists(index_html_path):
        with open(index_html_path, "w") as f:
            f.write(index_html)
            
    styles_css_path = os.path.join(electron_dir, "styles.css")
    if not os.path.exists(styles_css_path):
        with open(styles_css_path, "w") as f:
            f.write(styles_css)
            
    renderer_js_path = os.path.join(electron_dir, "renderer.js")
    if not os.path.exists(renderer_js_path):
        with open(renderer_js_path, "w") as f:
            f.write(renderer_js)
            
    preload_js_path = os.path.join(electron_dir, "preload.js")
    if not os.path.exists(preload_js_path):
        with open(preload_js_path, "w") as f:
            f.write(preload_js)
    
    # Create a basic README for the Electron project
    readme_md = """# Meeting Assistant Plus - Electron App

This is the Electron wrapper for Meeting Assistant Plus, allowing it to run as a cross-platform desktop application.

## Development

1. Install dependencies:
   ```
   npm install
   ```

2. Run the app in development mode:
   ```
   npm start
   ```

## Building

To build distributable packages for your platform:
```
npm run build
```

This will create executables in the `dist` folder.

## Requirements

- Node.js 14+
- Python 3.7+
- Required Python packages (see main application requirements)
"""
    
    readme_path = os.path.join(electron_dir, "README.md")
    if not os.path.exists(readme_path):
        with open(readme_path, "w") as f:
            f.write(readme_md)
            
    print(f"Electron files created in {electron_dir}")
    return electron_dir

# --------------------------
# Python Packaging Support for Electron
# --------------------------
def create_python_bundle():
    """Create a self-contained Python bundle for Electron distribution."""
    import shutil
    
    print("Creating Python bundle for Electron distribution...")
    
    # Create a spec file for PyInstaller
    spec_content = """# -*- mode: python ; coding: utf-8 -*-

block_cipher = None

a = Analysis(
    ['MeetingAssistantPlus_merged.py'],
    pathex=[],
    binaries=[],
    datas=[],
    hiddenimports=[
        'groq', 'qasync', 'playwright', 'docx', 'fpdf', 'wordcloud', 'keyboard', 
        'markdown', 'PyQt5', 'PIL'
    ],
    hookspath=[],
    hooksconfig={},
    runtime_hooks=[],
    excludes=[],
    win_no_prefer_redirects=False,
    win_private_assemblies=False,
    cipher=block_cipher,
    noarchive=False,
)

pyz = PYZ(a.pure, a.zipped_data, cipher=block_cipher)

exe = EXE(
    pyz,
    a.scripts,
    [],
    exclude_binaries=True,
    name='MeetingAssistantPlus',
    debug=False,
    bootloader_ignore_signals=False,
    strip=False,
    upx=True,
    console=True,
    disable_windowed_traceback=False,
    argv_emulation=True,
    target_arch=None,
    codesign_identity=None,
    entitlements_file=None,
)

coll = COLLECT(
    exe,
    a.binaries,
    a.zipfiles,
    a.datas,
    strip=False,
    upx=True,
    upx_exclude=[],
    name='MeetingAssistantPlus',
)
"""
    
    # Write spec file
    with open("MeetingAssistantPlus.spec", "w") as f:
        f.write(spec_content)
    
    # Create a setup script
    setup_script = """#!/bin/bash
# Setup script for bundling MeetingAssistantPlus for Electron

echo "Setting up Python bundle for Electron distribution..."

# Check if PyInstaller is installed
if ! pip show pyinstaller > /dev/null 2>&1; then
    echo "Installing PyInstaller..."
    pip install pyinstaller
fi

# Run PyInstaller
echo "Running PyInstaller..."
pyinstaller --clean MeetingAssistantPlus.spec

# Copy bundle to Electron folder
if [ -d "electron" ]; then
    echo "Copying Python bundle to Electron folder..."
    mkdir -p electron/python
    cp -r dist/MeetingAssistantPlus/* electron/python/
    echo "Python bundle created successfully in electron/python/"
else
    echo "Error: Electron directory not found. Run with --create-electron first."
    exit 1
fi

echo "Setup complete!"
"""
    
    # Write setup script
    setup_script_path = "setup_electron_bundle.sh"
    with open(setup_script_path, "w") as f:
        f.write(setup_script)
    
    # Make the script executable on Unix systems
    if sys.platform != 'win32':
        os.chmod(setup_script_path, 0o755)
    
    # Create a Windows batch script as well
    win_setup_script = """@echo off
echo Setting up Python bundle for Electron distribution...

REM Check if PyInstaller is installed
pip show pyinstaller > nul 2>&1
if %ERRORLEVEL% NEQ 0 (
    echo Installing PyInstaller...
    pip install pyinstaller
)

REM Run PyInstaller
echo Running PyInstaller...
pyinstaller --clean MeetingAssistantPlus.spec

REM Copy bundle to Electron folder
if exist electron (
    echo Copying Python bundle to Electron folder...
    if not exist electron\\python mkdir electron\\python
    xcopy /E /I /Y dist\\MeetingAssistantPlus electron\\python\\
    echo Python bundle created successfully in electron\\python\\
) else (
    echo Error: Electron directory not found. Run with --create-electron first.
    exit /b 1
)

echo Setup complete!
"""
    
    # Write Windows setup script
    with open("setup_electron_bundle.bat", "w") as f:
        f.write(win_setup_script)
    
    print("Python bundling scripts created:")
    print("- MeetingAssistantPlus.spec: PyInstaller specification")
    print("- setup_electron_bundle.sh: Unix setup script")
    print("- setup_electron_bundle.bat: Windows setup script")
    print("\nTo create the Python bundle:")
    print("1. On Unix/macOS: ./setup_electron_bundle.sh")
    print("2. On Windows: setup_electron_bundle.bat")

if __name__ == "__main__":
    # Process command line arguments
    if "--create-electron" in sys.argv:
        create_electron_files()
        print("Electron files created. To build the app:")
        print("1. cd electron")
        print("2. npm install")
        print("3. npm run build")
        sys.exit(0)
    elif "--create-bundle" in sys.argv:
        create_python_bundle()
        sys.exit(0)
    elif "--help" in sys.argv or "-h" in sys.argv:
        print(f"{APP_NAME} v{VERSION}")
        print("\nUsage options:")
        print("  --create-electron  : Create Electron project files")
        print("  --create-bundle    : Create Python bundle for Electron")
        print("  --electron-mode    : Run in Electron integration mode")
        print("  --help, -h         : Show this help message")
        sys.exit(0)
        
    # Normal application start
    main()